"""
International Transfer Pricing Comparability Database
====================================================
A Streamlit-based frontend application for the International Tax Cooperation
Framework Convention's global transfer pricing comparability analysis database.

Helps developing economy tax officials access comparable data and conduct
transfer pricing analyses under Articles 5 (Fair Distribution of Taxing Rights)
and 11 (Capacity Building).

Pages:
    1. Global Data Portal - Overview with world map and KPIs
    2. Comparability Analysis Wizard - 5-step guided analysis
    3. Extractive Industry Pricing - Commodity pricing workbench
    4. Country Policy Lookup - National TP regulation lookup
    5. Data Contribution Hub - Data upload and management
    6. MAP/APA Case Tracker - Dispute resolution tracking
    7. Digital Economy Dashboard - Cross-border services analysis

Author: Senior Streamlit Developer
Date: 2025-06-24
"""

from __future__ import annotations

import io
import os
import random
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional

import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor

# ============== Page Configuration ==============
st.set_page_config(
    page_title="国际转让定价对比数据库 / ITP Comparability Database",
    page_icon=":earth_americas:",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ============== Theme / Styling ==============
PRIMARY_COLOR = "#5B92E5"  # UN Blue
PRIMARY_DARK = "#3A6FCC"
PRIMARY_LIGHT = "#7AB0F0"
ACCENT_GREEN = "#27AE60"
ACCENT_RED = "#E74C3C"
ACCENT_YELLOW = "#F39C12"
ACCENT_PURPLE = "#8E44AD"
ACCENT_TEAL = "#1ABC9C"
BG_LIGHT = "#F7F9FC"
BG_CARD = "#FFFFFF"
TEXT_DARK = "#2C3E50"
TEXT_MUTED = "#7F8C8D"
BORDER_LIGHT = "#E8ECF0"

st.markdown(
    f"""
    <style>
    /* ===== Global Typography & Background ===== */
    .stApp {{
        background-color: {BG_LIGHT};
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Helvetica Neue', 'Noto Sans SC', sans-serif;
        color: {TEXT_DARK};
    }}
    /* Reduce default spacing for tighter layout */
    .stApp > header {{ background-color: transparent; }}
    .block-container {{ padding-top: 1.5rem; padding-bottom: 2rem; max-width: 1400px; }}

    /* ===== Page Main Header — gradient banner ===== */
    .main-header {{
        font-size: 2rem;
        font-weight: 800;
        color: #FFFFFF;
        background: linear-gradient(135deg, {PRIMARY_COLOR} 0%, {PRIMARY_DARK} 100%);
        padding: 1.2rem 1.8rem;
        border-radius: 14px;
        margin-bottom: 0.3rem;
        box-shadow: 0 4px 15px rgba(91, 146, 229, 0.3);
        letter-spacing: 0.5px;
    }}
    .sub-header {{
        font-size: 1rem;
        color: {TEXT_MUTED};
        margin-bottom: 1.2rem;
        padding-left: 0.3rem;
    }}

    /* ===== KPI Cards — modern with hover lift ===== */
    .kpi-card {{
        background: {BG_CARD};
        border-radius: 14px;
        padding: 1.3rem 1.5rem;
        border-top: 4px solid {PRIMARY_COLOR};
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
        transition: transform 0.25s ease, box-shadow 0.25s ease;
        position: relative;
        overflow: hidden;
    }}
    .kpi-card::before {{
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0;
        height: 4px;
        background: linear-gradient(90deg, {PRIMARY_COLOR}, {ACCENT_TEAL});
    }}
    .kpi-card:hover {{
        transform: translateY(-4px);
        box-shadow: 0 8px 25px rgba(0,0,0,0.12);
    }}
    .kpi-value {{
        font-size: 1.9rem;
        font-weight: 800;
        color: {PRIMARY_COLOR};
        line-height: 1.2;
    }}
    .kpi-label {{
        font-size: 0.78rem;
        color: {TEXT_MUTED};
        text-transform: uppercase;
        letter-spacing: 0.8px;
        margin-top: 0.3rem;
    }}

    /* ===== Info Bar — softer, rounded ===== */
    .info-bar {{
        background: linear-gradient(135deg, #EBF3FC 0%, #F5F9FF 100%);
        border-left: 4px solid {PRIMARY_COLOR};
        padding: 1.1rem 1.3rem;
        border-radius: 10px;
        margin: 0.8rem 0;
        font-size: 0.9rem;
        line-height: 1.6;
        box-shadow: 0 1px 4px rgba(0,0,0,0.04);
    }}

    /* ===== Status Colors ===== */
    .metric-good {{ color: {ACCENT_GREEN}; font-weight: 700; }}
    .metric-warning {{ color: {ACCENT_YELLOW}; font-weight: 700; }}
    .metric-danger {{ color: {ACCENT_RED}; font-weight: 700; }}

    /* ===== Update Feed Items ===== */
    .update-item {{
        padding: 0.7rem 0.8rem;
        border-bottom: 1px solid {BORDER_LIGHT};
        font-size: 0.85rem;
        transition: background 0.2s;
        border-radius: 6px;
    }}
    .update-item:hover {{
        background-color: #F0F4FF;
    }}

    /* ===== Sidebar Title ===== */
    .sidebar-title {{
        font-size: 1.15rem;
        font-weight: 700;
        color: {PRIMARY_COLOR};
        margin-bottom: 0.4rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid {PRIMARY_COLOR};
    }}

    /* ===== Method Badge ===== */
    .method-badge {{
        background: linear-gradient(135deg, {PRIMARY_COLOR}, {PRIMARY_DARK});
        color: white;
        padding: 0.35rem 1rem;
        border-radius: 20px;
        font-size: 0.8rem;
        font-weight: 600;
        display: inline-block;
    }}

    /* ===== Step Navigator — visual connected steps ===== */
    .step-nav-container {{
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin: 1rem 0 1.5rem 0;
        position: relative;
    }}
    .step-nav-item {{
        display: flex;
        flex-direction: column;
        align-items: center;
        position: relative;
        z-index: 1;
    }}
    .step-nav-circle {{
        width: 40px;
        height: 40px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: 700;
        font-size: 0.9rem;
        transition: all 0.3s;
    }}
    .step-nav-circle.active {{
        background: linear-gradient(135deg, {PRIMARY_COLOR}, {PRIMARY_DARK});
        color: white;
        box-shadow: 0 3px 12px rgba(91,146,229,0.4);
        transform: scale(1.1);
    }}
    .step-nav-circle.completed {{
        background: {ACCENT_GREEN};
        color: white;
    }}
    .step-nav-circle.pending {{
        background: #E0E6ED;
        color: {TEXT_MUTED};
    }}
    .step-nav-label {{
        font-size: 0.72rem;
        margin-top: 0.4rem;
        text-align: center;
        max-width: 120px;
        line-height: 1.3;
    }}

    /* ===== Buttons — modern, rounded ===== */
    div.stButton > button:first-child {{
        background: linear-gradient(135deg, {PRIMARY_COLOR}, {PRIMARY_DARK});
        color: white;
        border: none;
        border-radius: 10px;
        padding: 0.5rem 1.5rem;
        font-weight: 600;
        transition: all 0.25s ease;
        box-shadow: 0 2px 8px rgba(91,146,229,0.25);
    }}
    div.stButton > button:first-child:hover {{
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(91,146,229,0.4);
    }}
    div.stButton > button:first-child:active {{
        transform: translateY(0);
    }}
    /* Nav button styling: compact, text-left */
    div.stButton > button[kind="secondary"] {{
        background-color: {BG_CARD};
        color: {TEXT_DARK};
        border: 1px solid {BORDER_LIGHT};
        border-radius: 8px;
        padding: 0.45rem 0.8rem;
        font-size: 0.82rem;
        text-align: left;
        font-weight: 500;
        transition: all 0.2s;
    }}
    div.stButton > button[kind="secondary"]:hover {{
        background-color: #EBF3FC;
        border-color: {PRIMARY_COLOR};
        color: {PRIMARY_COLOR};
        transform: translateX(3px);
    }}

    /* ===== Dataframe / Table styling ===== */
    .stDataFrame {{
        border-radius: 10px;
        overflow: hidden;
        box-shadow: 0 1px 6px rgba(0,0,0,0.06);
    }}
    .stDataFrame table {{
        border-radius: 10px;
    }}
    .stDataFrame thead th {{
        background-color: {PRIMARY_COLOR} !important;
        color: white !important;
        font-weight: 600 !important;
    }}
    .stDataFrame tbody tr:nth-child(even) {{
        background-color: #F7F9FC;
    }}
    .stDataFrame tbody tr:hover {{
        background-color: #EBF3FC !important;
    }}

    /* ===== Expander styling ===== */
    .streamlit-expanderHeader {{
        font-weight: 600;
        font-size: 0.95rem;
        border-radius: 8px;
        transition: background 0.2s;
    }}
    .streamlit-expanderHeader:hover {{
        background-color: #F0F4FF;
    }}

    /* ===== Tabs styling ===== */
    .stTabs > div[role="tablist"] {{
        gap: 4px;
    }}
    .stTabs > div[role="tablist"] > button {{
        border-radius: 8px 8px 0 0;
        font-weight: 600;
        font-size: 0.88rem;
        padding: 0.5rem 1rem;
    }}
    .stTabs > div[role="tablist"] > button[aria-selected="true"] {{
        color: {PRIMARY_COLOR};
        border-bottom: 3px solid {PRIMARY_COLOR};
    }}

    /* ===== Metric cards spacing ===== */
    [data-testid="stMetric"] {{
        background: {BG_CARD};
        border-radius: 10px;
        padding: 0.8rem 1rem;
        box-shadow: 0 1px 4px rgba(0,0,0,0.05);
        border-left: 3px solid {PRIMARY_LIGHT};
    }}
    [data-testid="stMetric"]:hover {{
        box-shadow: 0 3px 10px rgba(0,0,0,0.08);
    }}

    /* ===== Selectbox / Input styling ===== */
    .stSelectbox > div > div {{
        border-radius: 8px;
    }}
    .stNumberInput > div > input {{
        border-radius: 8px;
    }}

    /* ===== Section divider ===== */
    hr {{
        border: none;
        border-top: 1px solid {BORDER_LIGHT};
        margin: 1.5rem 0;
    }}

    /* ===== Scrollbar styling ===== */
    ::-webkit-scrollbar {{
        width: 8px;
        height: 8px;
    }}
    ::-webkit-scrollbar-track {{
        background: {BG_LIGHT};
    }}
    ::-webkit-scrollbar-thumb {{
        background: #C4CDD5;
        border-radius: 4px;
    }}
    ::-webkit-scrollbar-thumb:hover {{
        background: {PRIMARY_LIGHT};
    }}

    /* ===== Sidebar ===== */
    section[data-testid="stSidebar"] {{
        background-color: {BG_CARD};
        border-right: 1px solid {BORDER_LIGHT};
    }}
    section[data-testid="stSidebar"] .stMarkdown {{
        font-size: 0.9rem;
    }}

    /* ===== Alert boxes ===== */
    .stAlert {{
        border-radius: 10px;
    }}

    /* ===== Download button ===== */
    .stDownloadButton > button {{
        background: linear-gradient(135deg, {ACCENT_GREEN}, #219A52) !important;
        border: none !important;
        border-radius: 10px !important;
        font-weight: 600 !important;
        box-shadow: 0 2px 8px rgba(39,174,96,0.25) !important;
    }}
    .stDownloadButton > button:hover {{
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(39,174,96,0.35) !important;
    }}
    </style>
    """,
    unsafe_allow_html=True,
)

# ============== Database Connection ==============
# Load config from db_config.py; fall back to inline defaults
try:
    from db_config import DB_CONFIG
except ImportError:
    DB_CONFIG = {
        "host": "localhost",
        "port": "5432",
        "database": "tp_comparability_db",
        "user": "postgres",
        "password": "your_password_here",
    }


@st.cache_resource
def get_db_engine() -> Optional[Any]:
    """Create SQLAlchemy engine for PostgreSQL connection.

    Uses connection pooling optimized for Supabase free tier (max 60 connections).
    Returns None if connection fails, in which case mock data is used.
    """
    try:
        from sqlalchemy import create_engine, text
        from urllib.parse import quote_plus

        # URL-encode password to handle special characters
        encoded_password = quote_plus(str(DB_CONFIG['password']))
        encoded_user = quote_plus(str(DB_CONFIG['user']))

        conn_str = (
            f"postgresql://{encoded_user}:{encoded_password}"
            f"@{DB_CONFIG['host']}:{DB_CONFIG['port']}/{DB_CONFIG['database']}"
        )
        engine = create_engine(
            conn_str,
            pool_pre_ping=True,          # Check connection before use
            pool_size=5,                  # Small pool for free tier
            max_overflow=3,               # Allow 3 extra connections
            pool_timeout=30,              # Wait max 30s for connection
            pool_recycle=1800,            # Recycle connections every 30 min
            connect_args={
                "connect_timeout": 10,    # 10s connection timeout
                "application_name": "tp_database_app",
            },
        )
        # Quick connectivity test
        with engine.connect() as conn:
            conn.execute(text("SELECT 1"))
        return engine
    except Exception as e:
        st.session_state.setdefault("db_error", str(e))
        return None


def get_db_connection() -> Optional[Any]:
    """Return a DB connection or None if unavailable."""
    engine = get_db_engine()
    if engine is None:
        return None
    try:
        return engine.connect()
    except Exception as e:
        st.session_state.setdefault("db_error", str(e))
        return None


def db_query(sql: str, params: Optional[dict] = None) -> Optional[pd.DataFrame]:
    """Execute a SQL query and return a DataFrame.

    Args:
        sql: SQL query string (use named params with :name for user input)
        params: Optional dict of query parameters for parameterized queries

    Returns None if DB unavailable. Errors are logged to session_state.
    """
    engine = get_db_engine()
    if engine is None:
        return None
    try:
        from sqlalchemy import text
        with engine.connect() as conn:
            if params:
                return pd.read_sql(text(sql), conn, params=params)
            else:
                return pd.read_sql(text(sql), conn)
    except Exception as e:
        import logging
        logging.warning(f"DB query failed: {e}")
        st.session_state.setdefault("db_error", str(e)[:200])
        return None


# ============== Data Loaders (real DB with mock fallback) ==============


@st.cache_data(ttl=300, max_entries=10)
def load_countries_data_v2() -> pd.DataFrame:
    """Load country-level data for the global map and KPIs from the database."""
    df = db_query("""
        SELECT
            c.country_name AS country,
            c.country_code AS iso_code,
            COALESCE(cc.comparables_count, 0) AS comparables_count,
            COALESCE(cc.last_updated, c.updated_at::date::text) AS last_updated,
            COALESCE(cc.industries_count, 0) AS industries_count,
            c.region_name AS region,
            CASE WHEN c.income_level IN ('Low income', 'Lower middle income', 'Upper middle income')
                 THEN TRUE ELSE FALSE END AS developing,
            COALESCE(cp.map_cases, 0) AS map_cases
        FROM countries c
        LEFT JOIN (
            SELECT country_code, COUNT(*) AS comparables_count,
                   MAX(data_as_of_date::text) AS last_updated,
                   COUNT(DISTINCT industry_code_isic) AS industries_count
            FROM comparable_companies
            GROUP BY country_code
        ) cc ON cc.country_code = c.country_code
        LEFT JOIN (
            SELECT country_code, map_cases_filed AS map_cases
            FROM country_policies
            WHERE policy_period_year = (SELECT MAX(policy_period_year) FROM country_policies)
        ) cp ON cp.country_code = c.country_code
        ORDER BY c.country_name
    """)
    if df is not None and not df.empty:
        # Taiwan is an inseparable part of China — mirror CHN data for TWN
        # Keep both rows for map coloring (Plotly needs both CHN and TWN iso codes)
        # But display as one entry in dropdowns
        chn_row = df[df["iso_code"] == "CHN"]
        if not chn_row.empty:
            twn_row = chn_row.iloc[0].copy()
            twn_row["iso_code"] = "TWN"
            twn_row["country"] = "Taiwan (China)"
            df = df[df["iso_code"] != "TWN"]  # remove existing TWN if any
            df = pd.concat([df, pd.DataFrame([twn_row])], ignore_index=True)
        # Normalize country names: treat "Taiwan (China)" as "China" for selection
        df["display_name"] = df["country"].replace({"Taiwan (China)": "China"})
        return df
    # Fallback: use comprehensive UN member states data
    from un_countries_data import generate_countries_data
    return generate_countries_data()


@st.cache_data(ttl=3600)
def load_mock_companies() -> pd.DataFrame:
    """Load comparable companies data from the database."""
    df = db_query("""
        SELECT
            cc.company_id,
            cc.company_name,
            co.country_name AS country,
            co.country_code,
            cc.industry_code_isic,
            cc.industry_description AS industry,
            cc.revenue / 1000000.0 AS revenue_musd,
            cc.net_worth / 1000000.0 AS net_assets_musd,
            cc.operating_margin,
            cc.return_on_net_worth AS roe,
            cc.roa,
            CASE WHEN cc.cogs > 0 THEN cc.gross_profit / cc.cogs END AS cost_plus_markup,
            CASE WHEN cc.operating_expenses > 0 THEN cc.gross_profit / cc.operating_expenses END AS berry_ratio,
            CASE WHEN cc.net_worth > 0 THEN (cc.revenue - cc.cogs - cc.operating_expenses) / cc.net_worth END AS roce,
            CASE WHEN cc.revenue > 0 THEN (cc.revenue - cc.cogs - cc.operating_expenses - cc.interest_expense) / cc.revenue END AS net_profit_margin,
            cc.related_party_transaction_ratio AS related_party_pct,
            cc.fiscal_year,
            cc.functional_complexity_score,
            cc.functions_performed AS functions,
            cc.risks_assumed AS risks,
            cc.data_tier,
            cc.data_source,
            CASE WHEN cc.risks_assumed LIKE '%Inventory%' THEN TRUE ELSE FALSE END AS has_inventory_risk,
            CASE WHEN cc.risks_assumed LIKE '%Market%' THEN TRUE ELSE FALSE END AS has_market_risk
        FROM comparable_companies cc
        LEFT JOIN countries co ON co.country_code = cc.country_code
        ORDER BY cc.company_id
    """)
    if df is not None and not df.empty:
        return df
    # Fallback: use comprehensive UN member states data
    from un_countries_data import generate_companies_data
    return generate_companies_data(800)


@st.cache_data(ttl=3600)
def load_commodity_data() -> pd.DataFrame:
    """Load commodity price data from the database."""
    df = db_query("""
        SELECT
            date_of_quote::text AS date,
            commodity_type AS commodity,
            spot_price AS price_usd,
            exchange_source
        FROM commodity_prices
        ORDER BY date_of_quote DESC
        LIMIT 500
    """)
    if df is not None and not df.empty:
        return df
    # Fallback mock data
    today = datetime.now()
    dates = [(today - timedelta(days=i)).strftime("%Y-%m-%d") for i in range(30)]
    np.random.seed(123)
    data = []
    base_prices = {"Copper": 8500, "Iron Ore": 120, "Gold": 1950}
    for commodity, base in base_prices.items():
        for d in dates:
            noise = np.random.normal(0, base * 0.02)
            data.append({"date": d, "commodity": commodity, "price_usd": round(base + noise, 2), "exchange_source": "Mock"})
    return pd.DataFrame(data)


@st.cache_data(ttl=3600)
def load_country_policies() -> pd.DataFrame:
    """Load country policy data from the database."""
    df = db_query("""
        SELECT
            co.country_name AS country,
            cp.arm_length_principle_enshrined AS has_tp_regulations,
            CASE
                WHEN cp.master_file_required = TRUE AND cp.local_file_required = TRUE AND cp.cbc_report_required = TRUE
                THEN 'Three-tier (Master + Local + CbCR)'
                WHEN cp.local_file_required = TRUE AND cp.cbc_report_required = TRUE
                THEN 'Local File + CbCR'
                WHEN cp.master_file_required = TRUE AND cp.local_file_required = TRUE
                THEN 'Master + Local File'
                WHEN cp.local_file_required = TRUE
                THEN 'Local File only'
                ELSE 'Not specified'
            END AS documentation_req,
            CASE WHEN array_length(cp.safe_harbour_type, 1) IS NOT NULL THEN TRUE ELSE FALSE END AS safe_harbor,
            COALESCE(cp.total_treaties_in_force, 0) AS treaty_count,
            COALESCE(cp.map_cases_filed, 0) AS map_cases_2024,
            CASE WHEN array_length(cp.apa_types_available, 1) IS NOT NULL THEN TRUE ELSE FALSE END AS apa_available,
            CASE WHEN array_length(cp.information_exchange_mechanism, 1) IS NOT NULL
                 THEN array_to_string(cp.information_exchange_mechanism, ', ')
                 ELSE 'N/A' END AS exchange_mechanism
        FROM country_policies cp
        JOIN countries co ON co.country_code = cp.country_code
        WHERE cp.policy_period_year = (SELECT MAX(policy_period_year) FROM country_policies)
        ORDER BY co.country_name
    """)
    if df is not None and not df.empty:
        return df
    # Fallback mock data
    data = [
        {"country": "Kenya", "has_tp_regulations": True, "documentation_req": "Three-tier (Master, Local, CbCR)",
         "safe_harbor": False, "treaty_count": 15, "map_cases_2024": 3, "apa_available": True, "exchange_mechanism": "EOIR + CbC"},
        {"country": "Nigeria", "has_tp_regulations": True, "documentation_req": "Local File + CbCR",
         "safe_harbor": True, "treaty_count": 18, "map_cases_2024": 5, "apa_available": True, "exchange_mechanism": "EOIR"},
        {"country": "South Africa", "has_tp_regulations": True, "documentation_req": "Three-tier",
         "safe_harbor": True, "treaty_count": 30, "map_cases_2024": 8, "apa_available": True, "exchange_mechanism": "EOIR + CbC + spontaneous"},
        {"country": "India", "has_tp_regulations": True, "documentation_req": "Three-tier + Form 3CEB",
         "safe_harbor": True, "treaty_count": 45, "map_cases_2024": 12, "apa_available": True, "exchange_mechanism": "EOIR + CbC"},
        {"country": "China", "has_tp_regulations": True, "documentation_req": "Three-tier + Special File",
         "safe_harbor": True, "treaty_count": 55, "map_cases_2024": 15, "apa_available": True, "exchange_mechanism": "EOIR + CbC"},
        {"country": "Brazil", "has_tp_regulations": True, "documentation_req": "Local File + CbCR",
         "safe_harbor": True, "treaty_count": 35, "map_cases_2024": 9, "apa_available": True, "exchange_mechanism": "EOIR"},
        {"country": "Germany", "has_tp_regulations": True, "documentation_req": "Three-tier",
         "safe_harbor": False, "treaty_count": 60, "map_cases_2024": 11, "apa_available": True, "exchange_mechanism": "EOIR + CbC + spontaneous"},
        {"country": "USA", "has_tp_regulations": True, "documentation_req": "Three-tier + 5471/8865",
         "safe_harbor": False, "treaty_count": 65, "map_cases_2024": 18, "apa_available": True, "exchange_mechanism": "EOIR + CbC + spontaneous"},
        {"country": "Indonesia", "has_tp_regulations": True, "documentation_req": "Three-tier",
         "safe_harbor": True, "treaty_count": 40, "map_cases_2024": 4, "apa_available": True, "exchange_mechanism": "EOIR + CbC"},
        {"country": "Vietnam", "has_tp_regulations": True, "documentation_req": "Local File + CbCR",
         "safe_harbor": True, "treaty_count": 22, "map_cases_2024": 3, "apa_available": True, "exchange_mechanism": "EOIR"},
        {"country": "Mexico", "has_tp_regulations": True, "documentation_req": "Three-tier",
         "safe_harbor": True, "treaty_count": 28, "map_cases_2024": 5, "apa_available": True, "exchange_mechanism": "EOIR + CbC"},
        {"country": "Chile", "has_tp_regulations": True, "documentation_req": "Local File + CbCR",
         "safe_harbor": True, "treaty_count": 20, "map_cases_2024": 2, "apa_available": True, "exchange_mechanism": "EOIR"},
    ]
    return pd.DataFrame(data)


@st.cache_data(ttl=3600)
def load_map_cases() -> pd.DataFrame:
    """Load MAP/APA case data from the database."""
    df = db_query("""
        SELECT
            dr.case_id,
            dr.dispute_type AS case_type,
            dr.taxpayer_name AS taxpayer,
            ci.country_name AS country_a,
            cc.country_name AS country_b,
            dr.issue_type AS issue,
            dr.case_status AS status,
            dr.date_filed::text AS start_date,
            COALESCE(dr.dispute_resolution_timeline, 0) * 30 AS duration_days,
            COALESCE(dr.adjustment_amount, 0) / 1000000.0 AS amount_disputed_musd,
            dr.date_resolved::text AS resolution_date
        FROM dispute_resolution dr
        LEFT JOIN countries ci ON ci.country_code = dr.initiating_jurisdiction
        LEFT JOIN countries cc ON cc.country_code = dr.counterparty_jurisdiction
        ORDER BY dr.date_filed DESC
    """)
    if df is not None and not df.empty:
        # Normalize status values for consistency
        status_map = {
            'Closed': 'Resolved',
            'Agreed': 'Resolved',
            'Implemented': 'Resolved',
            'Under negotiation': 'Active',
            'Pending': 'Pending',
            'Arbitration pending': 'Active',
        }
        df['status'] = df['status'].replace(status_map).fillna('Pending')
        df['amount_disputed_musd'] = df['amount_disputed_musd'].fillna(0)
        df['duration_days'] = df['duration_days'].fillna(0)
        return df
    # Fallback: use comprehensive UN member states data
    from un_countries_data import generate_map_cases
    return generate_map_cases(120)


@st.cache_data(ttl=3600)
def load_digital_comparables() -> pd.DataFrame:
    """Load digital economy comparable data from the database."""
    df = db_query("""
        SELECT
            de.digital_id AS company_id,
            cc.company_name,
            co.country_name AS country,
            de.digital_platform_type AS service_type,
            de.digital_revenue / 1000000.0 AS revenue_musd,
            CASE WHEN de.digital_revenue > 0
                 THEN (de.digital_revenue - de.digital_cost_of_sales - de.digital_operating_expenses) / de.digital_revenue
                 ELSE NULL END AS operating_margin,
            CASE WHEN de.digital_cost_of_sales > 0
                 THEN (de.digital_revenue - de.digital_cost_of_sales) / de.digital_cost_of_sales
                 ELSE NULL END AS cost_plus_markup,
            CASE WHEN de.digital_revenue > 0
                 THEN (de.digital_revenue - de.digital_cost_of_sales - de.digital_operating_expenses) / de.digital_revenue
                 ELSE NULL END AS tnmm_profit_margin,
            CASE WHEN (de.digital_revenue + de.non_digital_revenue) > 0
                 THEN de.digital_revenue / (de.digital_revenue + de.non_digital_revenue)
                 ELSE NULL END AS digital_intensity,
            de.automation_level,
            de.active_user_count / 1000000.0 AS users_millions,
            de.fiscal_year
        FROM digital_economy de
        LEFT JOIN comparable_companies cc ON cc.company_id = de.company_id
        LEFT JOIN countries co ON co.country_code = de.user_jurisdiction
        ORDER BY de.digital_id
    """)
    if df is not None and not df.empty:
        return df
    # Fallback: use comprehensive UN member states data
    from un_countries_data import generate_digital_comparables
    return generate_digital_comparables(150)


@st.cache_data(ttl=3600)
def load_recent_updates() -> list:
    """Load bilingual recent data contribution feed."""
    return [
        {"time": "2小时前 / 2h ago", "message": "肯尼亚上传了12家制造业可比公司 / Kenya uploaded 12 manufacturing comparables", "type": "upload"},
        {"time": "5小时前 / 5h ago", "message": "印度更新了8家软件公司的2025年Q1财务数据 / India updated Q1 2025 financial data for 8 software companies", "type": "update"},
        {"time": "1天前 / 1d ago", "message": "巴西贡献了采掘业铁矿石定价数据 / Brazil contributed extractive industry pricing data (iron ore)", "type": "upload"},
        {"time": "1天前 / 1d ago", "message": "南非批准了2项新的双边APA / South Africa approved 2 new bilateral APAs", "type": "policy"},
        {"time": "2天前 / 2d ago", "message": "尼日利亚添加了2024年MAP解决文档 / Nigeria added MAP resolution documentation for 2024", "type": "update"},
        {"time": "3天前 / 3d ago", "message": "越南上传了数字服务可比数据（SaaS行业）/ Vietnam uploaded digital services comparables (SaaS sector)", "type": "upload"},
        {"time": "3天前 / 3d ago", "message": "墨西哥更新了maquiladora安全港门槛 / Mexico updated safe harbor thresholds for maquiladoras", "type": "policy"},
        {"time": "4天前 / 4d ago", "message": "加纳贡献了5家农产品加工可比公司 / Ghana contributed 5 new agricultural processing comparables", "type": "upload"},
        {"time": "5天前 / 5d ago", "message": "中国更新了跨国企业CbC报告指引 / China updated CbC reporting guidance for MNEs", "type": "policy"},
        {"time": "1周前 / 1w ago", "message": "印度尼西亚上传了15家矿业可比公司 / Indonesia uploaded 15 new mining sector comparables", "type": "upload"},
    ]


# ============== Helper Functions ==============

def generate_import_template() -> pd.DataFrame:
    """Generate a standardized import template DataFrame for comparable company data."""
    return pd.DataFrame({
        'company_id': ['CC0001', 'CC0002', ''],
        'company_name': ['Company A (anonymized)', 'Company B (anonymized)', ''],
        'country_code': ['CHN', 'BRA', ''],
        'industry_code_isic': ['C26', 'C20', ''],
        'industry_code_naics': ['334413', '325199', ''],
        'legal_form': ['Public', 'Private', ''],
        'ownership_type': ['Independent', 'Subsidiary', ''],
        'revenue': [500000000, 120000000, ''],
        'cogs': [350000000, 80000000, ''],
        'operating_expenses': [80000000, 25000000, ''],
        'net_worth': [200000000, 60000000, ''],
        'fiscal_year': [2024, 2024, ''],
        'operating_margin': [14.0, 12.5, ''],
        'return_on_net_worth': [25.0, 20.0, ''],
        'data_source': ['Company Annual Reports', 'Tax Authority Assessment', ''],
        'data_tier': ['Tier1', 'Tier2', ''],
        'functional_complexity_score': [7, 5, ''],
        'functions_performed': ['Manufacturing; Distribution', 'Manufacturing', ''],
        'risks_assumed': ['Market Risk; Inventory Risk', 'Market Risk', ''],
    })


def get_color_for_count(count: int) -> str:
    """Return color code based on comparables count."""
    if count >= 50:
        return ACCENT_GREEN
    elif count >= 10:
        return ACCENT_YELLOW
    else:
        return ACCENT_RED


def calculate_wca(
    receivable_days: float,
    payable_days: float,
    inventory_days: float,
    interest_rate: float,
    annual_revenue: float,
    comp_ar_days: float = 0.0,
    comp_ap_days: float = 0.0,
    comp_inv_days: float = 0.0,
) -> tuple:
    """Calculate Working Capital Adjustment (WCA).

    WCA adjusts for the difference in working capital cycles between the tested party
    and the comparable companies.

    Formula:
        Tested Party WC Cycle = AR_Days - AP_Days + Inv_Days
        Comparable WC Cycle   = Comp_AR - Comp_AP + Comp_Inv
        WCA = (Tested_Cycle - Comparable_Cycle) × (Rate / 100) × (Revenue / 365)

    Args:
        receivable_days: Tested party's AR days
        payable_days: Tested party's AP days
        inventory_days: Tested party's inventory days
        interest_rate: Applicable interest rate (%)
        annual_revenue: Annual revenue (in millions)
        comp_ar_days: Comparable's AR days (default 0 = no adjustment)
        comp_ap_days: Comparable's AP days
        comp_inv_days: Comparable's inventory days

    Returns:
        tuple: (wca_amount, tested_cycle, comparable_cycle, cycle_diff)
    """
    if annual_revenue <= 0 or interest_rate <= 0:
        return 0.0, 0.0, 0.0, 0.0
    tested_cycle = receivable_days - payable_days + inventory_days
    comparable_cycle = comp_ar_days - comp_ap_days + comp_inv_days
    cycle_diff = tested_cycle - comparable_cycle
    wca = cycle_diff * (interest_rate / 100) * (annual_revenue / 365)
    return round(wca, 4), round(tested_cycle, 1), round(comparable_cycle, 1), round(cycle_diff, 1)


def generate_word_report(
    transaction_desc: dict,
    method_data: dict,
    screening_params: dict,
    adjustment_results: dict,
    al_range: dict,
    selected_method: str,
) -> bytes:
    """Generate a Word document report for the comparability analysis.

    Args:
        transaction_desc: Step 1 transaction description
        method_data: Step 2 method selection data
        screening_params: Step 3 screening parameters
        adjustment_results: Step 4 adjustment results
        al_range: Step 5 arm's length range
        selected_method: Selected transfer pricing method

    Returns:
        bytes: The DOCX file content as bytes
    """
    doc = Document()

    # Title
    title = doc.add_heading("Transfer Pricing Comparability Analysis Report", 0)
    title.alignment = 1  # Center

    # Metadata
    doc.add_paragraph(f"Report Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Generated by: ITP Comparability Database")
    doc.add_paragraph(f"Reference: UN TP Manual Ch.3 & OECD TPG Ch.1-3")
    doc.add_paragraph("")

    # Section 1: Transaction Description
    doc.add_heading("1. Controlled Transaction Description & Functional Analysis", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for key, value in transaction_desc.items():
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    # Section 2: Method Selection
    doc.add_heading("2. TP Method Selection", level=1)
    table_method = doc.add_table(rows=0, cols=2)
    table_method.style = "Light Grid Accent 1"
    for key, value in method_data.items():
        row = table_method.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    # Section 3: Screening Parameters
    doc.add_heading("3. Comparable Company Screening", level=1)
    table2 = doc.add_table(rows=0, cols=2)
    table2.style = "Light Grid Accent 1"
    for key, value in screening_params.items():
        row = table2.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    # Section 4: Adjustments
    doc.add_heading("4. Comparability Adjustments", level=1)
    table3 = doc.add_table(rows=0, cols=2)
    table3.style = "Light Grid Accent 1"
    for key, value in adjustment_results.items():
        row = table3.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    # Section 5: Arm's Length Range
    doc.add_heading("5. Arm's Length Range Determination", level=1)
    table4 = doc.add_table(rows=0, cols=2)
    table4.style = "Light Grid Accent 1"
    for key, value in al_range.items():
        row = table4.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    # Method (selected, not recommended)
    doc.add_heading("6. Selected Transfer Pricing Method", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Selected Method: {selected_method}").bold = True

    # Disclaimer
    doc.add_paragraph("")
    doc.add_heading("Disclaimer", level=2)
    doc.add_paragraph(
        "This report was generated by the ITP Comparability Database for informational purposes. "
        "The analysis should be reviewed by qualified transfer pricing professionals before use "
        "in any tax filing or dispute resolution proceeding.\n\n"
        "IMPORTANT: This tool does NOT constitute tax advice and cannot replace the judgment of a qualified "
        "transfer pricing advisor. The results are based on the data and parameters entered by the user, "
        "and the tool's methodology has limitations. Users should always verify results against current "
        "OECD Transfer Pricing Guidelines, UN Practical Manual on Transfer Pricing, and local tax regulations."
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def status_color(status: str) -> str:
    """Return color for MAP/APA case status."""
    colors = {
        "Active": ACCENT_YELLOW,
        "Resolved": ACCENT_GREEN,
        "Pending": PRIMARY_COLOR,
        "Closed - No agreement": ACCENT_RED,
    }
    return colors.get(status, "#999999")


# ============== Page 1: Global Data Portal ==============

def global_data_portal() -> None:
    """Page 1: Global Data Portal - Overview dashboard with world map and KPIs."""
    st.markdown("<div class='main-header'>🌐 全球数据总览 / Global Data Portal</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='sub-header'>国际转让定价数据库 — 全球可比性数据实时概览 / International Transfer Pricing Database</div>",
        unsafe_allow_html=True,
    )

    # Disclaimer banner — compact, styled
    st.markdown(
        f"""
        <div style='
            background: linear-gradient(135deg, #FFF3CD 0%, #FFF8E1 100%);
            border-left: 4px solid {ACCENT_YELLOW};
            padding: 0.8rem 1.2rem;
            border-radius: 10px;
            margin-bottom: 0.8rem;
            font-size: 0.83rem;
            line-height: 1.5;
        '>
            <b>⚠️ 数据声明 / Data Disclaimer:</b> 本数据库中的数据为演示用模拟数据。实际分析应使用经审计的真实财务数据。
            / Data is simulated for demonstration. Use audited real-world data for actual analysis.<br>
            <b>⛔ 免责声明 / Legal Disclaimer:</b> 本工具不构成税务建议，不能替代专业税务顾问的判断。
            / This tool does NOT constitute tax advice. Always verify against current OECD TPG, UN TP Manual, and local tax regulations.
        </div>
        """,
        unsafe_allow_html=True,
    )

    # 功能说明栏
    st.markdown("""
    <div class='info-bar'>
    💡 <b>本页功能说明 / Page Guide:</b><br>
    <b>中文：</b>这是数据库的「首页」，展示全球各经济体可比数据的总体情况。<br>
    &nbsp;&nbsp;• <b>世界地图</b>：颜色越绿表示该国有更多可比公司数据，红色表示数据较少<br>
    &nbsp;&nbsp;• <b>关键指标卡片</b>：显示全球数据总量、覆盖经济体数等核心数字<br>
    &nbsp;&nbsp;• <b>数据分布</b>：按地区和收入水平分组展示数据覆盖情况<br>
    <b>English:</b> This is the database homepage, showing the overall comparability data coverage across countries.<br>
    &nbsp;&nbsp;• <b>World Map</b>: Greener = more comparable company data; Red = limited data<br>
    &nbsp;&nbsp;• <b>KPI Cards</b>: Key metrics at a glance<br>
    <br>
    <i>依据 / Reference: 联合国国际税收合作框架公约第5条（公平分配征税权）和第11条（能力建设）<br>
    UN International Tax Cooperation Framework Convention, Art. 5 (Fair Distribution of Taxing Rights) & Art. 11 (Capacity Building)</i>
    </div>
    """, unsafe_allow_html=True)

    # Search + Help
    c1, c2 = st.columns([4, 1])
    with c1:
        search_term = st.text_input("搜索经济体、行业或公司 / Search", placeholder="例如：中国 制造业 / e.g., China manufacturing...")
    with c2:
        st.write("")
        st.write("")
        if st.button("❓ 帮助", key="help_global", help="点击查看使用指南"):
            st.info("""
            **全球数据总览 — 使用指南**

            - **地图**：鼠标悬停在国家上可查看可比公司数量、最后更新时间、覆盖行业数
            - **颜色含义**：🟢 绿色（>50家可比公司）、🟡 黄色（10-50家）、🔴 红色（<10家）
            - **指标卡片**：一目了然地看到全球数据规模
            - **最新动态**：各成员国最近的数据贡献记录

            本页面支持联合国国际税收合作框架公约第5条和第11条的实施。
            """)

    # Load data
    df = load_countries_data_v2()

    # Search filter
    if search_term:
        mask = (
            df["country"].str.contains(search_term, case=False, na=False) |
            df["region"].str.contains(search_term, case=False, na=False)
        )
        df = df[mask]

    # KPI Cards
    st.markdown("---")
    total_comparables = int(df["comparables_count"].sum())
    total_countries = len(df)
    total_industries = int(df["industries_count"].sum())
    active_map_cases = int(df["map_cases"].sum())
    dev_countries = len(df[df["developing"] == True])
    dev_coverage = round(dev_countries / total_countries * 100, 1) if total_countries > 0 else 0

    kpi_cols = st.columns(4)
    kpis = [
        ("可比公司 / Global Comparables", f"{total_comparables:,}", "Total comparable companies", "🏢"),
        ("覆盖经济体 / Economies", f"{total_countries}", "Economies covered", "🌍"),
        ("行业 / Industries", f"{total_industries}", "Industry segments", "🏭"),
        ("MAP案件 / Active MAP", f"{active_map_cases}", "Mutual Agreement Procedures", "⚖️"),
    ]
    for col, (label, value, help_text, icon) in zip(kpi_cols, kpis):
        with col:
            st.markdown(
                f"""
                <div class='kpi-card' title='{help_text}'>
                    <div style='font-size:1.6rem;margin-bottom:0.2rem'>{icon}</div>
                    <div class='kpi-value'>{value}</div>
                    <div class='kpi-label'>{label}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    st.markdown("---")

    # World Map
    map_col, feed_col = st.columns([3, 1])

    with map_col:
        st.subheader("全球数据覆盖地图 / Global Data Coverage Map")

        # Color mapping
        df["color_label"] = df["comparables_count"].apply(
            lambda x: "多(>50) / High" if x >= 50 else ("中(10-50) / Medium" if x >= 10 else "少(<10) / Low")
        )
        color_map = {"多(>50) / High": ACCENT_GREEN, "中(10-50) / Medium": ACCENT_YELLOW, "少(<10) / Low": ACCENT_RED}

        fig = px.choropleth(
            df,
            locations="iso_code",
            color="color_label",
            color_discrete_map=color_map,
            hover_name="country",
            hover_data={
                "iso_code": False,
                "comparables_count": True,
                "last_updated": True,
                "industries_count": True,
                "map_cases": True,
                "region": True,
            },
            labels={
                "comparables_count": "可比公司数 / Companies",
                "last_updated": "最后更新 / Last Updated",
                "industries_count": "行业数 / Industries",
                "map_cases": "MAP案件 / MAP Cases",
                "region": "地区 / Region",
                "color_label": "数据量 / Data Level",
            },
            projection="natural earth",
            height=500,
        )
        fig.update_layout(
            geo=dict(showframe=False, showcoastlines=True),
            margin=dict(l=0, r=0, t=20, b=0),
            legend=dict(orientation="h", yanchor="bottom", y=-0.1, xanchor="center", x=0.5),
        )
        st.plotly_chart(fig, use_container_width=True)

    with feed_col:
        st.subheader("最新数据动态 / Recent Data Updates")
        updates = load_recent_updates()
        for upd in updates[:7]:
            icon = {"upload": "📤", "update": "🔄", "policy": "📋"}.get(upd["type"], "•")
            st.markdown(
                f"<div class='update-item'><small><b>{icon} {upd['time']}</b><br/>{upd['message']}</small></div>",
                unsafe_allow_html=True,
            )

    # ===== Country Detail Panel =====
    st.markdown("---")
    st.subheader("🌍 经济体详情 / Economy Detail")
    # Use display_name for deduplicated list (China = CHN + TWN as one entry)
    if "display_name" in df.columns:
        display_names = sorted(df["display_name"].unique().tolist())
    else:
        display_names = sorted(df["country"].unique().tolist())
    country_options = ["— 选择经济体 / Select an economy —"] + display_names
    sel_country = st.selectbox("选择经济体查看详情 / Select economy", country_options,
                               help="点击选择国家，查看可比公司、政策等详情")
    if sel_country != country_options[0]:
        # Match by display_name — get all rows for this country (CHN + TWN)
        if "display_name" in df.columns:
            country_rows = df[df["display_name"] == sel_country]
        else:
            country_rows = df[df["country"] == sel_country]
        country_code = country_rows["iso_code"].iloc[0] if not country_rows.empty else ""
        render_country_detail(sel_country, country_code)

    # ===== Export =====
    st.markdown("---")
    export_col1, export_col2 = st.columns([1, 4])
    with export_col1:
        csv_portal = df.to_csv(index=False).encode('utf-8-sig')
        st.download_button(
            label="📥 导出国家数据CSV / Export",
            data=csv_portal,
            file_name=f"global_portal_data_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
        )

    # Info bar
    st.markdown("---")
    st.markdown(
        f"""
        <div class='info-bar'>
            <b>关于本数据库 / About this Database</b><br/>
            国际转让定价对比数据库支持<b>联合国国际税收合作框架公约</b>的实施。
            帮助发展中国家获取高质量可比数据进行转让定价分析，支持<b>第5条</b>（公平分配征税权）和<b>第11条</b>（能力建设）。
            各成员国在自愿基础上贡献匿名化的可比公司数据，确保数据主权的同时促进国际税收合作。<br/><br/>
            The ITP Comparability Database supports the implementation of the
            <b>UN International Tax Cooperation Framework Convention</b>. It enables developing economies
            to access quality comparability data for transfer pricing analyses, supporting
            <b>Article 5</b> (Fair Distribution of Taxing Rights) and <b>Article 11</b> (Capacity Building).
        </div>
        """,
        unsafe_allow_html=True,
    )

    # Regional breakdown
    st.subheader("地区数据分布 / Regional Data Distribution")
    region_summary = df.groupby("region").agg(
        countries=("country", "count"),
        total_comparables=("comparables_count", "sum"),
        avg_comparables=("comparables_count", "mean"),
    ).round(1).reset_index()
    region_summary.columns = ["地区 / Region", "经济体 / Economies", "Total Comparables", "平均/经济体 / Avg per Economy"]

    fig_bar = px.bar(
        region_summary,
        x="地区 / Region",
        y="Total Comparables",
        color="经济体 / Economies",
        text="Total Comparables",
        color_continuous_scale="Blues",
        height=350,
    )
    fig_bar.update_layout(margin=dict(l=0, r=0, t=30, b=0))
    st.plotly_chart(fig_bar, use_container_width=True)


# ============== Page 2: Comparability Analysis Wizard ==============

def comparability_wizard() -> None:
    """Page 2: Comparability Analysis Wizard - 5-step guided TP analysis."""
    st.markdown("<div class='main-header'>🔍 可比性分析向导 / Comparability Analysis Wizard</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='sub-header'>逐步引导完成转让定价可比性分析 / Step-by-step Transfer Pricing Comparability Analysis</div>",
        unsafe_allow_html=True,
    )

    # 功能说明栏
    st.markdown("""
    <div class='info-bar'>
    💡 <b>本页功能说明 / Page Guide:</b><br>
    <b>中文：</b>这是数据库的「核心功能」，通过5个步骤帮你完成专业的转让定价分析。<br>
    &nbsp;&nbsp;• <b>第1步</b>：描述关联交易 + 功能分析(FAR)：交易类型、参与方、履行的功能、承担的风险、使用的资产<br>
    &nbsp;&nbsp;• <b>第2步</b>：选择转让定价方法：系统根据FAR分析智能推荐，用户确认<br>
    &nbsp;&nbsp;• <b>第3步</b>：筛选可比公司：根据所选方法确定筛选条件，从数据库找可比<br>
    &nbsp;&nbsp;• <b>第4步</b>：可比性调整：对筛选出的可比公司进行调整（如营运资金调整）<br>
    &nbsp;&nbsp;• <b>第5步</b>：确定独立交易区间：计算四分位区间，判断被测试方结果是否合规<br>
    <b>English:</b> 5-step TP analysis following UN TP Manual Ch.3 & OECD TPG standard workflow.<br>
    &nbsp;&nbsp;• <b>Step 1</b>: Describe transaction + Functional Analysis (FAR: Functions, Assets, Risks)<br>
    &nbsp;&nbsp;• <b>Step 2</b>: Select TP Method (system recommends based on FAR)<br>
    &nbsp;&nbsp;• <b>Step 3</b>: Screen Comparables (filters driven by method choice)<br>
    &nbsp;&nbsp;• <b>Step 4</b>: Comparability Adjustments (e.g., Working Capital Adjustment)<br>
    &nbsp;&nbsp;• <b>Step 5</b>: Determine Arm's Length Range (quartile analysis)<br>
    <br>
    <i>依据 / Reference: UN TP Manual Ch.3 (Comparability Analysis) & OECD TPG Ch.1-3<br>
    标准流程：交易描述→功能分析(FAR)→方法选择→识别可比→可比性调整→独立交易区间</i>
    </div>
    """, unsafe_allow_html=True)

    # Help button
    col_help, _ = st.columns([1, 5])
    with col_help:
        if st.button("❓ 帮助 / Help", key="help_wizard", help="分析流程指南 / Analysis guide"):
            st.info("""
            **可比性分析向导 — 使用指南 / Wizard Guide**

            本向导按照UN TP Manual和OECD TPG标准流程设计：
            This wizard follows the UN TP Manual and OECD TPG standard workflow:

            **第1步 / Step 1**: 描述关联交易 + 功能分析(FAR) / Describe transaction + Functional Analysis
            **第2步 / Step 2**: 选择转让定价方法 / Select TP Method
            **第3步 / Step 3**: 筛选可比公司 / Screen Comparables
            **第4步 / Step 4**: 可比性调整 / Comparability Adjustments
            **第5步 / Step 5**: 确定独立交易区间 / Determine Arm's Length Range

            💡 方法选择必须在筛选可比之前——因为不同方法需要不同类型的可比数据。
            Method selection must precede comparables screening — different methods require different types of comparable data.
            """)

    # Initialize session state
    if "wiz_step1" not in st.session_state:
        st.session_state.wiz_step1 = {}
    if "wiz_step2" not in st.session_state:
        st.session_state.wiz_step2 = {}
    if "wiz_step3" not in st.session_state:
        st.session_state.wiz_step3 = {}
    if "wiz_step4" not in st.session_state:
        st.session_state.wiz_step4 = {}
    if "wiz_current_tab" not in st.session_state:
        st.session_state.wiz_current_tab = 0

    tab_labels = [
        "① 交易+FAR / Transaction",
        "② 选择方法 / Method",
        "③ 筛选可比 / Screen",
        "④ 可比性调整 / Adjust",
        "⑤ 独立交易区间 / Range",
    ]
    current_step = st.session_state.get("wiz_current_tab", 0)

    # Progress bar
    progress_pct = int((current_step + 1) / 5 * 100)
    st.progress(progress_pct / 100, text=f"分析进度 / Progress: {progress_pct}% (第{current_step + 1}步 / Step {current_step + 1} of 5)")

    # Visual step navigator — circle-based connected steps
    step_nums = ["①", "②", "③", "④", "⑤"]
    step_short = ["交易+FAR", "选方法", "筛可比", "做调整", "定区间"]
    step_short_en = ["Transaction", "Method", "Screen", "Adjust", "Range"]
    nav_html = '<div class="step-nav-container">'
    for i in range(5):
        if i < current_step:
            cls = "completed"
        elif i == current_step:
            cls = "active"
        else:
            cls = "pending"
        nav_html += f'''
        <div class="step-nav-item">
            <div class="step-nav-circle {cls}">{step_nums[i] if cls != 'completed' else '✓'}</div>
            <div class="step-nav-label"><b>{step_short[i]}</b><br/>{step_short_en[i]}</div>
        </div>'''
        if i < 4:
            nav_html += f'<div style="flex:1;height:2px;background:{"#27AE60" if i < current_step else "#E0E6ED"};margin:0 4px;margin-top:-20px;transition:all 0.3s;"></div>'
    nav_html += '</div>'
    st.markdown(nav_html, unsafe_allow_html=True)

    # Also provide clickable buttons for navigation
    step_cols = st.columns(5)
    for i, (col, label) in enumerate(zip(step_cols, tab_labels)):
        with col:
            if i == current_step:
                st.markdown(f"**👉 {label}**")
            else:
                if st.button(label, key=f"goto_step_{i}", help=f"跳转到 {label}"):
                    st.session_state.wiz_current_tab = i
                    st.rerun()

    # ========== Step 1: Transaction Description + FAR ==========
    if current_step == 0:
        st.subheader("第1步：描述关联交易 + 功能分析(FAR) / Step 1: Transaction + Functional Analysis")
        st.markdown("""
        📌 **这步做什么 / What this step does:** 描述你要分析的关联交易，并完成功能分析(FAR)。
        功能分析是转让定价分析的基础——确定各方实际做了什么、承担了什么风险、使用了什么资产。

        🎯 **通俗理解 / In plain terms:** 就像写一个"故事摘要"——谁和谁交易、交易什么、各自负责干什么活、各自承担什么风险。
        """)

        # Tested party selection guidance (OECD TPG Ch.3.60)
        st.markdown("#### 🤖 被测试方选择指引 / Tested Party Selection Guidance")
        st.caption("依据 / Ref: OECD TPG Ch.3.60 — 被测试方应是功能较简单、不拥有独特无形资产的一方")
        tested_party_complexity = st.selectbox(
            "被测试方特征 / Tested Party Characteristics",
            ["请选择 / Select...", "功能简单（常规制造/分销/服务）/ Routine functions",
             "功能复杂（含研发/品牌/核心无形资产）/ Complex functions with intangibles"],
            key="tested_party_guide",
        )
        if tested_party_complexity.startswith("功能复杂"):
            st.warning("⚠️ OECD TPG Ch.3.60: 被测试方通常应选择功能较简单的一方。"
                       "如果被测试方拥有高价值无形资产，建议改选另一方作为被测试方，或考虑使用利润分割法(PSM)。")

        tx_type = st.selectbox(
            "交易类型 / Transaction Type",
            ["Sale of Goods", "Provision of Services", "Transfer of Intangibles",
             "Financial Transactions", "Digital Services", "Cross-Border Services"],
            format_func=lambda x: {
                "Sale of Goods": "货物销售 / Sale of Goods",
                "Provision of Services": "提供服务 / Provision of Services",
                "Transfer of Intangibles": "无形资产转让 / Transfer of Intangibles",
                "Financial Transactions": "金融交易 / Financial Transactions",
                "Digital Services": "数字服务 / Digital Services",
                "Cross-Border Services": "跨境服务 / Cross-Border Services",
            }.get(x, x),
            help="选择最符合你要分析的关联交易的类型",
        )
        tested_party = st.selectbox(
            "被测试方 / Tested Party",
            ["Local Entity", "Foreign Entity"],
            format_func=lambda x: "本地实体 / Local Entity" if x == "Local Entity" else "外国实体 / Foreign Entity",
            help="被测试方是其利润率将被检验的实体。通常选择功能较简单、风险较低的一方作为被测试方。",
        )

        # FAR Analysis — Functions
        st.markdown("##### 📋 功能分析(FAR) / Functional Analysis")
        st.markdown("**功能 (F) — 履行的功能 / Functions Performed**")
        func_cols = st.columns(3)
        functions = []
        with func_cols[0]:
            if st.checkbox("研发 / R&D", help="研发活动"): functions.append("R&D")
            if st.checkbox("制造 / Manufacturing", help="生产或制造货物"): functions.append("Manufacturing")
        with func_cols[1]:
            if st.checkbox("分销 / Distribution", help="销售和分销功能"): functions.append("Distribution")
            if st.checkbox("营销 / Marketing", help="营销、品牌推广活动"): functions.append("Marketing")
        with func_cols[2]:
            if st.checkbox("数字平台运营 / Digital Platform", help="运营数字平台或市场"): functions.append("Digital Platform Operation")

        # FAR Analysis — Risks
        st.markdown("**风险 (R) — 承担的风险 / Risks Assumed**")
        risk_cols = st.columns(3)
        risks = []
        with risk_cols[0]:
            if st.checkbox("市场风险 / Market Risk", help="市场需求波动风险"): risks.append("Market Risk")
            if st.checkbox("信用风险 / Credit Risk", help="客户不付款风险"): risks.append("Credit Risk")
        with risk_cols[1]:
            if st.checkbox("库存风险 / Inventory Risk", help="库存过时或价格变动风险"): risks.append("Inventory Risk")
            if st.checkbox("汇率风险 / FX Risk", help="外汇汇率波动风险"): risks.append("FX Risk")
        with risk_cols[2]:
            if st.checkbox("数字服务风险 / Digital Service Risk", help="数字服务特有风险"): risks.append("Digital Service Risk")

        # FAR Analysis — Assets
        st.markdown("**资产 (A) — 使用的资产 / Assets Used**")
        asset_cols = st.columns(3)
        assets = []
        with asset_cols[0]:
            if st.checkbox("固定资产 / Fixed Assets", help="厂房、设备等"): assets.append("Fixed Assets")
            if st.checkbox("无形资产 / Intangibles", help="专利、商标、技术等"): assets.append("Intangibles")
        with asset_cols[1]:
            if st.checkbox("金融资产 / Financial Assets", help="应收账款、投资等"): assets.append("Financial Assets")
            if st.checkbox("存货 / Inventory", help="原材料、在产品、产成品"): assets.append("Inventory")
        with asset_cols[2]:
            if st.checkbox("人力资本 / Human Capital", help="关键技术人员、管理团队"): assets.append("Human Capital")

        # Additional comparability factors (OECD TPG Ch.3.46-3.63)
        st.markdown("##### 📋 其他可比性因素 / Other Comparability Factors")
        st.caption("依据 / Ref: OECD TPG Ch.3.46-3.63 — 除FAR外，还需考虑合同条款、经济环境和商业策略")

        other_cols = st.columns(3)
        with other_cols[0]:
            st.markdown("**合同条款 / Contractual Terms**")
            contractual = st.text_area(
                "合同条款概述 / Contract Terms",
                placeholder="如：定价方式、付款条件、保修期限、变更条款...",
                height=80,
                help="OECD TPG Ch.3.55-3.57: 合同条款规定了风险和收益的分配，是可比性分析的重要因素",
            )
        with other_cols[1]:
            st.markdown("**经济环境 / Economic Circumstances**")
            economic = st.text_area(
                "经济环境概述 / Economic Circumstances",
                placeholder="如：市场规模、竞争程度、监管成本、地理位置...",
                height=80,
                help="OECD TPG Ch.3.58-3.61: 不同市场的经济环境差异可能影响定价",
            )
        with other_cols[2]:
            st.markdown("**商业策略 / Business Strategies**")
            business = st.text_area(
                "商业策略概述 / Business Strategies",
                placeholder="如：市场渗透策略、创新驱动、成本领先...",
                height=80,
                help="OECD TPG Ch.3.62-3.63: 企业的商业策略影响其定价决策",
            )

        value_creation_location = st.text_input(
            "价值创造地 / Value Creation Location",
            placeholder="例如：中国（研发）；德国（制造）",
            help="说明该交易中价值主要在哪里创造",
        )

        if st.button("保存第1步，进入第2步 / Save & Next", key="save_step1", type="primary"):
            st.session_state.wiz_step1 = {
                "Transaction Type": tx_type,
                "Tested Party": tested_party,
                "Functions": ", ".join(functions) if functions else "None selected",
                "Risks": ", ".join(risks) if risks else "None selected",
                "Assets": ", ".join(assets) if assets else "None selected",
                "Contractual Terms": contractual or "Not specified",
                "Economic Circumstances": economic or "Not specified",
                "Business Strategies": business or "Not specified",
                "Value Creation Location": value_creation_location or "Not specified",
            }
            st.session_state.wiz_current_tab = 1
            st.rerun()

    # ========== Step 2: Method Selection ==========
    if current_step == 1:
        st.subheader("第2步：选择转让定价方法 / Step 2: Select TP Method")
        st.markdown("""
        📌 **这步做什么 / What this step does:** 根据第1步的功能分析(FAR)结果，选择最合适的转让定价方法。
        方法选择决定了第3步要筛选什么类型的可比数据。

        🎯 **通俗理解 / In plain terms:** 不同交易适用不同"定价方法"——就像评估房子用"可比成交价法"、评估公司用"市盈率法"。
        选错方法会导致后续分析全部偏离。
        """)

        # Back button
        if st.button("⬅️ 上一步 / Previous", key="back_to_1"):
            st.session_state.wiz_current_tab = 0
            st.rerun()

        # Get FAR data from Step 1
        step1_data = st.session_state.get("wiz_step1", {})
        tx_type = step1_data.get("Transaction Type", "Sale of Goods")
        functions_str = step1_data.get("Functions", "")
        risks_str = step1_data.get("Risks", "")

        st.info(f"📊 第1步摘要 / Step 1 Summary:\n交易类型: {tx_type} | 功能: {functions_str} | 风险: {risks_str}")

        # All 5 TP methods
        tp_methods = {
            "CUP": {
                "name_cn": "可比非受控价格法 (CUP)",
                "name_en": "Comparable Uncontrolled Price",
                "desc": "将关联交易价格与独立交易价格直接对比。最直接但需要高质量可比价格数据。",
                "pli": "价格 / Price",
                "best_for": "有公开市场价格的商品交易、金融交易",
            },
            "RPM": {
                "name_cn": "再销售价格法 (RPM)",
                "name_en": "Resale Price Method",
                "desc": "从独立买方的再销售价格中减去合理毛利，得出关联交易价格。",
                "pli": "毛利率 / Gross Margin",
                "best_for": "分销商（买进卖出、不改变产品形态）",
            },
            "CPM": {
                "name_cn": "成本加成法 (CPM)",
                "name_en": "Cost Plus Method",
                "desc": "在合理成本基础上加上合理加成，得出关联交易价格。",
                "pli": "成本加成率 / Cost Plus Markup",
                "best_for": "制造商（合同制造）、服务提供商",
            },
            "TNMM": {
                "name_cn": "交易净利润法 (TNMM)",
                "name_en": "Transactional Net Margin Method",
                "desc": "比较关联交易的净利润率与独立交易的净利润率。最常用，对差异容忍度最高。",
                "pli": "营业利润率(OM) / Berry Ratio / ROCE",
                "best_for": "大多数关联交易（尤其当CUP/RPM/CPM不可用时）",
            },
            "PSM": {
                "name_cn": "利润分割法 (PSM)",
                "name_en": "Profit Split Method",
                "desc": "将关联交易合并利润按各方贡献分配。最复杂但最公平。\n两种分析法：①贡献分析法(Contribution Analysis) — 按各方相对贡献分配总利润；②剩余利润分析法(Residual Analysis) — 先按常规功能分配常规利润(TNMM/CPM)，再按独特无形资产贡献分割剩余利润 (OECD TPG Ch.2.C)",
                "pli": "利润分割比例 / Profit Split Ratio",
                "best_for": "双方都拥有高价值无形资产、高度整合的交易",
            },
        }

        # Smart recommendation
        st.markdown("#### 🤖 智能推荐 / Smart Recommendation")
        recommended = []
        reasons = []
        if tx_type == "Financial Transactions":
            recommended.append("CUP"); reasons.append("金融交易有公开市场利率可比 / Financial transactions have public benchmarks")
        if tx_type == "Sale of Goods" and "Distribution" in functions_str:
            recommended.append("RPM"); reasons.append("分销商适用再销售价格法 / Distributors suit RPM (gross margin)")
        if tx_type == "Sale of Goods" and "Manufacturing" in functions_str:
            recommended.append("CPM"); reasons.append("合同制造适用成本加成法 / Contract manufacturing suits CPM")
        if "Transfer of Intangibles" in tx_type:
            recommended.append("PSM"); reasons.append("无形资产涉及双方共创价值 / Intangibles involve joint value creation")
        if tx_type in ("Provision of Services", "Digital Services", "Cross-Border Services"):
            recommended.append("TNMM"); reasons.append("服务/数字交易通常用TNMM / Service/digital transactions suit TNMM")
        if not recommended:
            recommended.append("TNMM"); reasons.append("TNMM最通用 / TNMM is most versatile")

        for m, r in zip(recommended, reasons):
            info = tp_methods[m]
            st.success(f"✅ **{info['name_cn']} / {info['name_en']}** — {r}")

        st.markdown("#### 📋 所有5种方法 / All 5 Methods (请选择一个)")
        method_choice = st.radio(
            "选择方法 / Select Method",
            list(tp_methods.keys()),
            format_func=lambda k: f"{tp_methods[k]['name_cn']} / {tp_methods[k]['name_en']}",
            index=list(tp_methods.keys()).index(recommended[0]) if recommended else 3,
        )

        # Show method details
        info = tp_methods[method_choice]
        st.markdown(f"""
        <div class='info-bar'>
        <b>{info['name_cn']} / {info['name_en']}</b><br>
        {info['desc']}<br><br>
        <b>利润水平指标(PLI)：</b> {info['pli']}<br>
        <b>最适用于：</b> {info['best_for']}
        </div>
        """, unsafe_allow_html=True)

        st.info("💡 方法选择依据：UN TP Manual Ch.3 & OECD TPG Ch.2。选择方法时应考虑：\n"
                "1) 关联交易的性质；2) 功能分析(FAR)结果；3) 可比数据的可用性；4) 数据的可靠性。\n\n"
                "📌 **OECD TPG Ch.2.1 — 最适当方法原则 / Most Appropriate Method Rule:**\n"
                "不是选\"能用的\"方法，而是选\"最可靠的\"方法。应比较各方法的优缺点，"
                "选择能对关联交易得出最可靠独立交易结果的方法。\n"
                "Not just any applicable method, but the one that produces the most reliable arm's length result.")

        if st.button("保存第2步，进入第3步 / Save & Next", key="save_step2", type="primary"):
            st.session_state.wiz_step2 = {
                "Selected Method": method_choice,
                "Method Name": info['name_en'],
                "PLI": info['pli'],
            }
            st.session_state.wiz_current_tab = 2
            st.rerun()

    # ========== Step 3: Screen Comparables ==========
    if current_step == 2:
        st.subheader("第3步：筛选可比公司 / Step 3: Screen Comparables")
        st.markdown("""
        📌 **这步做什么 / What this step does:** 按照实际案件分析流程设置筛选条件，逐步缩小可比范围。

        🎯 **通俗理解 / In plain terms:** 从数据库中找出和你分析的公司"相似"的企业——同年份、同区域、同行业、
        规模相近、以独立交易为主。这些就是你的"参照物"。

        📋 **筛选顺序 / Screening Order:** ①年度 → ②区域 → ③独立性 → ④数据质量 → ⑤定量指标 → ⑥产品功能
        """)

        # Back button
        if st.button("⬅️ 上一步 / Previous", key="back_to_2"):
            st.session_state.wiz_current_tab = 1
            st.rerun()

        # Show selected method context
        step2_data = st.session_state.get("wiz_step2", {})
        selected_method = step2_data.get("Selected Method", "TNMM")
        method_pli = step2_data.get("PLI", "营业利润率(OM)")
        st.info(f"📊 第2步选择的方法: **{step2_data.get('Method Name', 'TNMM')}** | "
                f"对应PLI: **{method_pli}**\n"
                f"💡 筛选条件已根据所选方法调整。{selected_method}法需要可比公司的{method_pli}数据。")

        companies_df = load_mock_companies()

        # Prepare ISIC section mapping
        if 'industry_code_isic' in companies_df.columns:
            companies_df['isic_section'] = companies_df['industry_code_isic'].astype(str).str[0]
        section_names = {
            'A': 'A-农林牧渔 / Agriculture', 'B': 'B-采矿 / Mining',
            'C': 'C-制造 / Manufacturing', 'D': 'D-电力燃气 / Energy',
            'E': 'E-水务 / Water', 'F': 'F-建筑 / Construction',
            'G': 'G-批发零售 / Wholesale & Retail', 'H': 'H-运输仓储 / Transport',
            'I': 'I-住宿餐饮 / Hospitality', 'J': 'J-信息技术 / IT',
            'K': 'K-金融保险 / Financial', 'L': 'L-房地产 / Real Estate',
            'M': 'M-专业服务 / Professional', 'N': 'N-辅助服务 / Support',
        }

        # ===== 3.1 对应年度 / Fiscal Year =====
        with st.expander("📅 ① 对应年度 / Fiscal Year", expanded=True):
            st.caption("选择与被测试方一致的财政年度。建议使用3年数据平滑波动 (OECD TPG Ch.3.79)。")
            available_years = sorted(companies_df["fiscal_year"].dropna().unique().tolist())
            fy_range = st.slider(
                "财政年度范围 / Fiscal Year Range",
                int(min(available_years)) if available_years else 2022,
                int(max(available_years)) if available_years else 2026,
                (2022, 2024),
                help="💡 选取与被测试方相同年度的可比数据。OECD建议使用3年平均以平滑异常波动。"
            )

        # ===== 3.2 区域 / Region =====
        with st.expander("🌍 ② 区域 / Region & Economy"):
            st.caption("选择可比公司所在的经济体。优先选择同一地理区域或类似经济环境的经济体。")
            sel_countries = st.multiselect(
                "经济体范围 / Economies",
                sorted(companies_df["country"].unique().tolist()),
                default=[],
                help="💡 不选则包含所有经济体。选择同一区域或经济发展水平相近的经济体可提高可比性。\n依据：OECD TPG Ch.3.55 经济环境可比性因素",
            )

        # ===== 3.3 独立性 / Independence =====
        with st.expander("🔗 ③ 独立性筛选 / Independence Test"):
            st.caption("确保可比公司以独立交易为主，排除高度关联化的企业。")
            rp_max = st.slider(
                "关联交易占比上限 / Max Related-Party Revenue Ratio",
                0.0, 1.0, 0.30, format="%.0f%%",
                help="💡 关联交易占比>30%的公司通常不适合作为独立可比。建议默认30%上限。\n依据：UN TP Manual Ch.3 — 可比公司应以独立交易为主",
            )
            if 'ownership_type' in companies_df.columns:
                ownership_filter = st.multiselect(
                    "所有权类型 / Ownership Type",
                    ["Independent", "Subsidiary", "Public", "Private"],
                    default=["Independent", "Public"],
                    help="独立公司可比性更强",
                )
            else:
                ownership_filter = []

        # ===== 3.4 数据完整性与异常 / Data Quality =====
        with st.expander("✅ ④ 数据完整性与异常 / Data Quality & Outliers"):
            st.caption("确保数据质量可靠，剔除异常值。")
            exclude_loss = st.checkbox(
                "剔除亏损公司 / Exclude loss-making companies",
                value=True,
                help="💡 亏损公司通常不作为可比对象——亏损可能由非正常因素导致（管理层决策失误、一次性减值等），不具备可比性。\n依据：UN TP Manual Ch.3 / OECD TPG Ch.3",
            )
            require_audited = st.checkbox(
                "仅保留审计数据 / Audited data only",
                value=True,
                help="优先使用经审计的财务数据",
            )
            min_data_tier = st.selectbox(
                "最低数据密级 / Minimum Data Tier",
                ["Tier1", "Tier2", "Tier3"],
                index=1,
                help="💡 Tier1=审计财报（最可靠）, Tier2=税务评估数据, Tier3=估算数据\n建议至少Tier2以上",
            )

        # ===== 3.5 定量指标 / Quantitative Indicators =====
        with st.expander("📊 ⑤ 定量指标 / Quantitative Indicators"):
            st.caption("按收入规模、资产规模等定量指标进一步缩小范围。")
            qi1, qi2 = st.columns(2)
            with qi1:
                rev_min, rev_max = st.slider(
                    "收入范围 / Revenue Range (M$)",
                    float(companies_df["revenue_musd"].min()),
                    float(companies_df["revenue_musd"].max()),
                    (5.0, 500.0),
                    help="与被测试方收入规模相近的公司更具可比性",
                )
            with qi2:
                asset_min, asset_max = st.slider(
                    "净资产范围 / Net Assets Range (M$)",
                    float(companies_df["net_assets_musd"].min()),
                    float(companies_df["net_assets_musd"].max()),
                    (1.0, 750.0),
                    help="资产规模相近的公司运营模式更具可比性",
                )

            # Function-specific quantitative indicators
            st.markdown("**功能相关指标 / Function-Specific Indicators:**")
            func_ind1, func_ind2 = st.columns(2)
            with func_ind1:
                rd_filter = st.checkbox(
                    "研发功能筛选 / R&D Function Filter",
                    help="仅包含有研发活动的公司（适用于技术密集型交易）",
                )
                if rd_filter:
                    rd_intensity_min = st.slider("最低研发强度 / Min R&D Intensity (%)", 0.0, 20.0, 1.0, 0.5)
                sg_filter = st.checkbox(
                    "销管费功能筛选 / SG&A Function Filter",
                    help="仅包含有销售管理活动的公司（适用于分销/服务交易）",
                )
                if sg_filter:
                    sga_intensity_min = st.slider("最低销管费强度 / Min SG&A Intensity (%)", 0.0, 50.0, 5.0, 1.0)
            with func_ind2:
                inv_filter = st.checkbox(
                    "必须承担库存风险 / Must have Inventory Risk",
                    help="仅包含承担库存风险的公司",
                )
                mkt_filter = st.checkbox(
                    "必须承担市场风险 / Must have Market Risk",
                    help="仅包含承担市场风险的公司",
                )

        # ===== 3.6 产品和功能 / Product & Functions =====
        with st.expander("🏭 ⑥ 产品和功能 / Product & Functions"):
            st.caption("按ISIC行业大类和功能组合筛选，确保产品/服务类型可比。")
            if 'isic_section' in companies_df.columns:
                sections = sorted(companies_df['isic_section'].dropna().unique().tolist())
                section_labels = [section_names.get(s, f'{s}-其他 / Other') for s in sections]
                sel_section_idx = st.selectbox(
                    "行业大类 / Industry Sector (ISIC)",
                    range(len(sections)),
                    format_func=lambda i: section_labels[i],
                    help="按ISIC行业大类筛选 — 选择与被测试方相同或相近的行业",
                )
                sel_industry = sections[sel_section_idx]
            else:
                sel_industry = st.selectbox(
                    "行业 / Industry",
                    ["全部 / All"] + sorted(companies_df["industry"].dropna().unique().tolist()),
                    help="选择可比公司所在的行业",
                )

            func_match = st.multiselect(
                "要求的功能组合 / Required Functions",
                ["R&D", "Manufacturing", "Distribution", "Marketing", "Digital Platform"],
                default=[],
                help="💡 可比公司必须至少履行其中一个功能。\n例如：分析分销商时选择「Distribution」；分析合同制造商时选择「Manufacturing」。",
            )

        # ===== Apply all filters =====
        filtered = companies_df.copy()

        # 3.1 Year filter
        filtered = filtered[(filtered["fiscal_year"] >= fy_range[0]) & (filtered["fiscal_year"] <= fy_range[1])]

        # 3.2 Region filter
        if sel_countries:
            filtered = filtered[filtered["country"].isin(sel_countries)]

        # 3.3 Independence filter
        filtered = filtered[filtered["related_party_pct"] <= rp_max]
        if ownership_filter and 'ownership_type' in filtered.columns:
            filtered = filtered[filtered['ownership_type'].isin(ownership_filter)]

        # 3.4 Data quality filters
        if exclude_loss and "operating_margin" in filtered.columns:
            filtered = filtered[filtered["operating_margin"] >= 0]
        if min_data_tier and 'data_tier' in filtered.columns:
            tier_order = {"Tier1": 1, "Tier2": 2, "Tier3": 3}
            min_tier_num = tier_order.get(min_data_tier, 2)
            filtered = filtered[filtered['data_tier'].apply(lambda x: tier_order.get(x, 3) <= min_tier_num)]

        # 3.5 Quantitative filters
        filtered = filtered[
            (filtered["revenue_musd"] >= rev_min) &
            (filtered["revenue_musd"] <= rev_max) &
            (filtered["net_assets_musd"] >= asset_min) &
            (filtered["net_assets_musd"] <= asset_max)
        ]
        if inv_filter:
            filtered = filtered[filtered["has_inventory_risk"] == True]
        if mkt_filter and 'has_market_risk' in filtered.columns:
            filtered = filtered[filtered["has_market_risk"] == True]

        # 3.6 Product & functions
        if sel_industry and sel_industry != "全部 / All" and 'isic_section' in filtered.columns:
            filtered = filtered[filtered['isic_section'] == sel_industry]
        if func_match:
            filtered = filtered[filtered["functions"].apply(lambda x: any(f in str(x) for f in func_match))]

        # Results
        st.markdown("---")
        result_col1, result_col2, result_col3 = st.columns([2, 1, 1])
        with result_col1:
            st.markdown(f"### 🔍 筛选结果：**{len(filtered)} 家可比公司**")
        with result_col2:
            if len(filtered) > 0:
                st.metric("覆盖经济体 / Economies", f"{filtered['country'].nunique()}")
        with result_col3:
            if len(filtered) > 0:
                st.metric("覆盖年度 / Years", f"{filtered['fiscal_year'].nunique()}")

        if len(filtered) == 0:
            st.warning("⚠️ 没有公司符合筛选条件，请放宽一些筛选标准。/ No companies match — try broadening filters.")
        else:
            # Small sample warning
            if len(filtered) < 7:
                st.warning(f"⚠️ **小样本警告 / Small Sample**: 仅 {len(filtered)} 家，OECD建议≥7家。"
                           f"考虑放宽筛选条件。/ OECD TPG Ch.3.55 recommends ≥7 comparables.")

            # Display results
            display_cols = ["company_name", "country", "industry", "revenue_musd",
                           "operating_margin", "berry_ratio", "roce", "related_party_pct",
                           "fiscal_year", "data_tier"]
            display_cols = [c for c in display_cols if c in filtered.columns]
            st.dataframe(filtered[display_cols], use_container_width=True, hide_index=True)

            # Export
            wiz_export_col1, _ = st.columns([1, 4])
            with wiz_export_col1:
                csv_wiz = filtered[display_cols].to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    label="📥 导出筛选结果 / Export CSV",
                    data=csv_wiz,
                    file_name=f"screened_companies_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv",
                )

            if st.button("保存第3步，进入第4步 / Save & Next", key="save_step3", type="primary"):
                st.session_state.wiz_step3 = {
                    "Industry": sel_industry,
                    "经济体 / Economies": ", ".join(sel_countries) if sel_countries else "All",
                    "FY Range": f"{fy_range[0]}-{fy_range[1]}",
                    "Revenue Range": f"${rev_min:.1f}M - ${rev_max:.1f}M",
                    "Net Assets Range": f"${asset_min:.1f}M - ${asset_max:.1f}M",
                    "Max Related Party %": f"{rp_max*100:.0f}%",
                    "Selected Companies": len(filtered),
                }
                st.session_state.wiz_current_tab = 3
                st.rerun()

    # ========== Step 4: Comparability Adjustments ==========
    if current_step == 3:
        st.subheader("第4步：可比性调整 / Step 4: Comparability Adjustments")
        st.markdown("""
        📌 **这步做什么 / What this step does:** 对筛选出的可比公司进行调整，使其与被测试方更具可比性。

        🎯 **通俗理解 / In plain terms:** 可比公司和你分析的公司不可能完全一样——有的收账快、有的存货多。
        这些差异会影响利润率，需要"调整"后才能公平比较。最常用的是**营运资金调整(WCA)**。
        """)

        # Back button
        if st.button("⬅️ 上一步 / Previous", key="back_to_3"):
            st.session_state.wiz_current_tab = 2
            st.rerun()

        # ---- PSM Profit Split Calculator (if PSM was selected) ----
        step2_data_psm = st.session_state.get("wiz_step2", {})
        if step2_data_psm.get("Selected Method") == "PSM":
            st.markdown("##### 🔀 利润分割法(PSM)计算器 / Profit Split Calculator")
            st.caption("依据 / Ref: OECD TPG Ch.2.C — 两步法：先分配常规利润，再分割剩余利润")

            psm_col1, psm_col2 = st.columns(2)
            with psm_col1:
                total_combined_profit = st.number_input("合并关联利润 / Combined Related-Party Profit (M$)",
                                                         min_value=-10000.0, max_value=100000.0, value=50.0, step=1.0,
                                                         key="psm_combined")
                routine_profit_a = st.number_input("A方常规利润 (TNMM/CPM) / Party A Routine Profit (M$)",
                                                    min_value=-10000.0, max_value=100000.0, value=10.0, step=0.5,
                                                    key="psm_routine_a")
                routine_profit_b = st.number_input("B方常规利润 (TNMM/CPM) / Party B Routine Profit (M$)",
                                                    min_value=-10000.0, max_value=100000.0, value=8.0, step=0.5,
                                                    key="psm_routine_b")
            with psm_col2:
                # Contribution analysis
                contrib_a = st.slider("A方独特贡献权重 / Party A Unique Contribution (%)",
                                       0, 100, 50, 5, key="psm_contrib_a",
                                       help="基于功能分析(FAR)中各方对独特无形资产的贡献")
                contrib_b = 100 - contrib_a
                st.metric("B方独特贡献权重 / Party B Contribution", f"{contrib_b}%")

            # Calculate residual profit
            routine_total = routine_profit_a + routine_profit_b
            residual_profit = total_combined_profit - routine_total
            residual_a = residual_profit * (contrib_a / 100)
            residual_b = residual_profit * (contrib_b / 100)
            total_a = routine_profit_a + residual_a
            total_b = routine_profit_b + residual_b

            st.markdown("---")
            psm_result_cols = st.columns(4)
            with psm_result_cols[0]:
                st.metric("常规利润合计 / Routine Total", f"${routine_total:.1f}M")
            with psm_result_cols[1]:
                color_r = "#27AE60" if residual_profit >= 0 else "#E74C3C"
                st.metric("剩余利润 / Residual Profit", f"${residual_profit:.1f}M")
            with psm_result_cols[2]:
                st.metric("A方总利润 / Party A Total", f"${total_a:.1f}M")
            with psm_result_cols[3]:
                st.metric("B方总利润 / Party B Total", f"${total_b:.1f}M")

            if residual_profit < 0:
                st.warning("⚠️ 剩余利润为负，说明常规利润分配过高或关联交易整体亏损。需重新评估常规利润计算。")

            st.markdown("---")

        adj_type = st.selectbox(
            "调整类型",
            ["Working Capital Adjustment (WCA)", "Risk Adjustment", "Asset Intensity Adjustment",
             "Functional Adjustment", "Custom Adjustment"],
            format_func=lambda x: {
                "Working Capital Adjustment (WCA)": "营运资金调整 (WCA)",
                "Risk Adjustment": "风险调整",
                "Asset Intensity Adjustment": "资产密集度调整",
                "Functional Adjustment": "功能调整",
                "Custom Adjustment": "自定义调整",
            }.get(x, x),
            help="选择要应用的调整类型",
        )

        if adj_type == "Working Capital Adjustment (WCA)":
            st.markdown("#### 营运资金调整计算器")
            st.markdown("""
            计算被测试方与可比方营运资金周期的**差异**，并将差异金额调整为利润率调整。

            💡 **原理 / Principle:** 如果被测试方收账慢、存货多，意味着更多资金被占用，需要相应调整利润率。
            """)
            st.caption("依据 / Ref: OECD TPG Ch.3.55-3.57 — 营运资金调整是最常见的可比性调整")

            wca_col1, wca_col2 = st.columns(2)
            with wca_col1:
                st.markdown("**🏢 被测试方 / Tested Party**")
                ar_days = st.number_input("应收账款天数 / AR Days", min_value=0.0, max_value=365.0, value=60.0, step=1.0,
                                         help="被测试方应收账款平均回收期", key="wca_ar")
                ap_days = st.number_input("应付账款天数 / AP Days", min_value=0.0, max_value=365.0, value=45.0, step=1.0,
                                         help="被测试方应付账款平均付款期", key="wca_ap")
                inv_days = st.number_input("存货天数 / Inventory Days", min_value=0.0, max_value=365.0, value=30.0, step=1.0,
                                          help="被测试方存货平均持有天数", key="wca_inv")
            with wca_col2:
                st.markdown("**📊 可比方均值 / Comparable Average**")
                comp_ar = st.number_input("可比方应收账款天数 / Comp AR Days", min_value=0.0, max_value=365.0, value=50.0, step=1.0,
                                         help="可比公司应收账款天数平均值", key="wca_comp_ar")
                comp_ap = st.number_input("可比方应付账款天数 / Comp AP Days", min_value=0.0, max_value=365.0, value=40.0, step=1.0,
                                         help="可比公司应付账款天数平均值", key="wca_comp_ap")
                comp_inv = st.number_input("可比方存货天数 / Comp Inv Days", min_value=0.0, max_value=365.0, value=25.0, step=1.0,
                                          help="可比公司存货天数平均值", key="wca_comp_inv")

            int_rate = st.number_input("适用利率 (%) / Interest Rate", min_value=0.0, max_value=50.0, value=5.0, step=0.1,
                                       help="用于折现营运资金差异的年化利率", key="wca_rate")
            annual_rev = st.number_input("被测试方年收入（百万美元）/ Annual Revenue (M$)", min_value=0.0, max_value=10000.0, value=100.0, step=1.0,
                                         help="被测试方的年收入（百万美元）", key="wca_rev")

            wca_result, tested_cycle, comp_cycle, cycle_diff = calculate_wca(
                ar_days, ap_days, inv_days, int_rate, annual_rev,
                comp_ar, comp_ap, comp_inv
            )

            st.markdown("---")
            result_cols = st.columns(4)
            with result_cols[0]:
                st.metric("被测试方周期 / Tested Cycle", f"{tested_cycle:.1f} 天")
            with result_cols[1]:
                st.metric("可比方周期 / Comp Cycle", f"{comp_cycle:.1f} 天")
            with result_cols[2]:
                diff_color = "#E74C3C" if cycle_diff > 0 else "#27AE60"
                st.metric("周期差异 / Difference", f"{cycle_diff:.1f} 天")
            with result_cols[3]:
                st.metric("WCA 调整额 / WCA Amount", f"${wca_result:.4f}M")

            pct_of_revenue = (wca_result / annual_rev * 100) if annual_rev > 0 else 0
            st.info(f"💡 WCA占收入比例: **{pct_of_revenue:.4f}%** — 此调整将在第5步自动应用到被测试方利润率")

            if st.button("保存第4步，进入第5步 / Save & Next", key="save_step4", type="primary"):
                st.session_state.wiz_step4 = {
                    "Adjustment Type": adj_type,
                    "Tested AR Days": ar_days, "Tested AP Days": ap_days, "Tested Inv Days": inv_days,
                    "Comp AR Days": comp_ar, "Comp AP Days": comp_ap, "Comp Inv Days": comp_inv,
                    "Interest Rate": f"{int_rate}%",
                    "Annual Revenue": f"${annual_rev}M",
                    "WCA Amount": f"${wca_result:.4f}M",
                    "Cycle Difference": f"{cycle_diff:.1f} days",
                }
                st.session_state.wiz_current_tab = 4
                st.rerun()

        else:
            # ---- FIX #4: Non-WCA adjustments with actual calculation logic ----
            st.info(f"**{adj_type}** — 请根据功能分析(FAR)结果输入调整参数。")
            if adj_type == "Risk Adjustment":
                st.caption("OECD TPG Ch.3.50-3.54: 风险调整用于补偿被测试方与可比公司承担的风险差异。")
                risk_type = st.selectbox("风险类型 / Risk Type",
                    ["市场风险 / Market Risk", "信用风险 / Credit Risk", "外汇风险 / FX Risk", "存货风险 / Inventory Risk"],
                    key="risk_type")
                risk_premium = st.slider("风险溢价 / Risk Premium (%)", 0.0, 10.0, 1.0, 0.1,
                    help="被测试方额外承担的风险需要相应的溢价补偿", key="risk_prem")
                adj_magnitude = risk_premium
                st.metric("风险调整幅度 / Risk Adjustment", f"{adj_magnitude:.2f}%")
            elif adj_type == "Asset Intensity Adjustment":
                st.caption("OECD TPG Ch.3.47-3.49: 资产密集度差异影响回报率，需进行调整。")
                asset_diff = st.slider("资产密集度差异 / Asset Intensity Difference (%)", -50.0, 50.0, 0.0, 1.0,
                    help="被测试方与可比公司平均资产密集度的差异（正=被测试方更高）", key="asset_diff")
                adj_magnitude = -asset_diff * 0.1  # Simplified: 10% of diff as adjustment
                st.metric("资产调整幅度 / Asset Adjustment", f"{adj_magnitude:.2f}%")
            elif adj_type == "Functional Adjustment":
                st.caption("OECD TPG Ch.3.46-3.49: 功能差异调整用于补偿被测试方与可比公司功能复杂度差异。")
                func_diff = st.slider("功能复杂度差异 / Functional Complexity Difference", -5, 5, 0, 1,
                    help="正=被测试方功能更复杂，需更高回报；负=更简单", key="func_diff")
                adj_magnitude = float(func_diff) * 0.5
                st.metric("功能调整幅度 / Functional Adjustment", f"{adj_magnitude:.2f}%")
            else:
                adj_magnitude = st.slider("调整幅度 (%) / Adjustment (%)", -20.0, 20.0, 0.0, 0.5,
                                         help="手动调整百分比 / Manual adjustment", key="custom_adj")
                st.metric("应用调整 / Applied Adjustment", f"{adj_magnitude:.2f}%")

            # Save non-WCA adjustment to session state
            if st.button("保存第4步，进入第5步 / Save & Next", key="save_step4_other", type="primary"):
                st.session_state.wiz_step4 = {
                    "Adjustment Type": adj_type,
                    "Adjustment Magnitude": f"{adj_magnitude:.2f}%",
                }
                st.session_state.wiz_current_tab = 4
                st.rerun()

    # ========== Step 5: Arm's Length Range ==========
    if current_step == 4:
        st.subheader("第5步：独立交易区间确定 / Step 5: Arm's Length Range Determination")
        st.markdown("""
        📌 **这步做什么 / What this step does:** 用筛选并调整后的可比公司数据计算独立交易区间（四分位区间），
        然后判断被测试方的利润率是否落在区间内。

        🎯 **通俗理解 / In plain terms:** 把所有可比公司的利润率排个队，取中间25%-75%的范围作为"合理区间"。
        被测试方的利润率落在这个区间内=合规；低于下限=可能利润转移；高于上限=也需调查原因。
        """)

        # Back button
        if st.button("⬅️ 上一步 / Previous", key="back_to_4"):
            st.session_state.wiz_current_tab = 3
            st.rerun()

        # ---- FIX #1: Use filtered comparables from Step 3 ----
        companies_df = load_mock_companies()
        step3_data = st.session_state.get("wiz_step3", {})
        step2_data = st.session_state.get("wiz_step2", {})
        step4_data = st.session_state.get("wiz_step4", {})

        selected_method = step2_data.get("Selected Method", "TNMM")

        # Re-apply Step 3 filters to get the actual comparable set
        if step3_data:
            sel_industry = step3_data.get("Industry", "")
            sel_countries_str = step3_data.get("经济体 / Economies", "All")
            fy_range_str = step3_data.get("FY Range", "2022-2024")
            rev_range_str = step3_data.get("Revenue Range", "$5.0M - $500.0M")
            asset_range_str = step3_data.get("Net Assets Range", "$1.0M - $750.0M")
            rp_max_str = step3_data.get("Max Related Party %", "30%")

            filtered = companies_df.copy()
            # Recreate isic_section the same way Step 3 does
            if 'industry_code_isic' in filtered.columns and 'isic_section' not in filtered.columns:
                filtered['isic_section'] = filtered['industry_code_isic'].astype(str).str[0]
            elif 'isic_section' in filtered.columns:
                # Normalize to first character for consistency
                filtered['isic_section'] = filtered['isic_section'].astype(str).str[0]
            # Industry filter — match by first character of isic_section
            if sel_industry and sel_industry != "All" and 'isic_section' in filtered.columns:
                filtered = filtered[filtered['isic_section'] == sel_industry]
            # Country filter
            if sel_countries_str and sel_countries_str != "All":
                countries_list = [c.strip() for c in sel_countries_str.split(",")]
                filtered = filtered[filtered["country"].isin(countries_list)]
            # FY filter
            try:
                fy_parts = fy_range_str.split("-")
                fy_lo, fy_hi = int(fy_parts[0].strip()), int(fy_parts[1].strip())
                filtered = filtered[(filtered["fiscal_year"] >= fy_lo) & (filtered["fiscal_year"] <= fy_hi)]
            except Exception:
                pass
            # Revenue filter
            try:
                rev_parts = rev_range_str.replace("$", "").replace("M", "").split("-")
                rev_lo, rev_hi = float(rev_parts[0].strip()), float(rev_parts[1].strip())
                filtered = filtered[(filtered["revenue_musd"] >= rev_lo) & (filtered["revenue_musd"] <= rev_hi)]
            except Exception:
                pass
            # Net assets filter
            try:
                asset_parts = asset_range_str.replace("$", "").replace("M", "").split("-")
                asset_lo, asset_hi = float(asset_parts[0].strip()), float(asset_parts[1].strip())
                filtered = filtered[(filtered["net_assets_musd"] >= asset_lo) & (filtered["net_assets_musd"] <= asset_hi)]
            except Exception:
                pass
            # Related party ratio
            try:
                rp_max_val = float(rp_max_str.replace("%", "")) / 100
                filtered = filtered[filtered["related_party_pct"] <= rp_max_val]
            except Exception:
                pass
            # Exclude loss-making
            if "operating_margin" in filtered.columns:
                filtered = filtered[filtered["operating_margin"] >= 0]
        else:
            filtered = companies_df.copy()

        st.info(f"📊 使用第3步筛选结果：**{len(filtered)} 家可比公司** | "
                f"方法: **{step2_data.get('Method Name', 'TNMM')}** | "
                f"PLI: **{step2_data.get('PLI', '营业利润率(OM)')}**")

        # ---- FIX #2: Use correct PLI based on method selection ----
        pli_column_map = {
            "CUP": ("operating_margin", "价格 / Price (此处以毛利率替代演示)"),
            "RPM": ("gross_margin", "毛利率 / Gross Margin"),
            "CPM": ("cost_plus_markup", "成本加成率 / Cost Plus Markup"),
            "TNMM": ("operating_margin", "营业利润率 / Operating Margin"),
            "PSM": ("operating_margin", "利润分割比例 (以OM为基础) / Profit Split (OM-based)"),
        }
        pli_col, pli_label = pli_column_map.get(selected_method, ("operating_margin", "营业利润率 / Operating Margin"))

        # Use appropriate column or fall back to OM
        if pli_col not in filtered.columns:
            pli_col = "operating_margin"
        pli_data = filtered[pli_col].dropna()
        pli_values = pli_data.values * 100  # Convert to percentage

        # ---- FIX #6: Small sample warning (OECD TPG Ch.3.55) ----
        if len(pli_values) > 0 and len(pli_values) < 7:
            st.warning(
                f"⚠️ **小样本警告 / Small Sample Warning** (OECD TPG Ch.3.55):\n"
                f"当前仅有 {len(pli_values)} 家可比公司。OECD建议至少7家以保证四分位区间的统计可靠性。\n"
                f"建议放宽筛选条件以增加可比样本，或考虑使用全部可比公司的中位数而非四分位区间。\n\n"
                f"Only {len(pli_values)} comparables. OECD recommends ≥7 for reliable interquartile range. "
                f"Consider broadening filters or using median."
            )

        # ---- FIX #7: Outlier detection (1.5×IQR rule) ----
        if len(pli_values) >= 4:
            q1_raw = np.percentile(pli_values, 25)
            q3_raw = np.percentile(pli_values, 75)
            iqr_raw = q3_raw - q1_raw
            lower_fence = q1_raw - 1.5 * iqr_raw
            upper_fence = q3_raw + 1.5 * iqr_raw
            outliers_mask = (pli_values < lower_fence) | (pli_values > upper_fence)
            n_outliers = int(outliers_mask.sum())
            if n_outliers > 0:
                st.warning(
                    f"⚠️ **检测到 {n_outliers} 个异常值 / {n_outliers} outliers detected** "
                    f"(超出1.5×IQR范围: [{lower_fence:.2f}%, {upper_fence:.2f}%])\n"
                    f"建议检查这些公司是否真正可比，或考虑剔除。OECD TPG Ch.3.83建议对异常值进行调查。"
                )
                with st.expander("查看异常值详情 / View Outlier Details"):
                    outlier_df = filtered.iloc[np.where(outliers_mask)[0]][
                        ["company_name", "country", "industry", "revenue_musd", pli_col, "fiscal_year"]
                    ].copy()
                    outlier_df[pli_col] = (outlier_df[pli_col] * 100).round(2).astype(str) + "%"
                    st.dataframe(outlier_df, use_container_width=True, hide_index=True)
                    if st.checkbox("剔除异常值后重新计算 / Recalculate excluding outliers", value=False, key="remove_outliers"):
                        pli_values = pli_values[~outliers_mask]
                        st.success(f"已剔除 {n_outliers} 个异常值，剩余 {len(pli_values)} 家可比公司。")

        if len(pli_values) > 0:
            stats = {
                "Min": np.min(pli_values),
                "Q1 (25th)": np.percentile(pli_values, 25),
                "Median": np.median(pli_values),
                "Q3 (75th)": np.percentile(pli_values, 75),
                "Max": np.max(pli_values),
                "Mean": np.mean(pli_values),
                "Count": len(pli_values),
            }

            # Stats display
            stat_cols = st.columns(len(stats))
            for col, (key, val) in zip(stat_cols, stats.items()):
                with col:
                    if key == "Count":
                        st.metric(key, f"{int(val)}")
                    else:
                        st.metric(key, f"{val:.2f}%")

            # Box plot
            fig_box = go.Figure()
            fig_box.add_trace(go.Box(
                y=pli_values,
                name=f"{pli_label}",
                boxmean=True,
                marker_color=PRIMARY_COLOR,
                line_color=PRIMARY_COLOR,
            ))
            fig_box.update_layout(
                title=f"Arm's Length Range — {pli_label}",
                yaxis_title=f"{pli_label} (%)",
                height=400,
                showlegend=False,
            )
            st.plotly_chart(fig_box, use_container_width=True)

            # ---- FIX #5: Multi-year averaging guidance (OECD TPG Ch.3.79) ----
            st.markdown("##### 📅 多年平均分析 / Multi-Year Averaging")
            st.caption("依据 / Ref: OECD TPG Ch.3.79 — 建议使用3年平均数据平滑异常波动")
            if len(filtered) > 0 and "fiscal_year" in filtered.columns:
                yearly_stats = filtered.groupby("fiscal_year")[pli_col].agg(['mean', 'median', 'count']).reset_index()
                yearly_stats['mean'] = (yearly_stats['mean'] * 100).round(2)
                yearly_stats['median'] = (yearly_stats['median'] * 100).round(2)
                yearly_stats.columns = ["年度 / Fiscal Year", "平均值 / Mean (%)", "中位数 / Median (%)", "公司数 / Count"]
                st.dataframe(yearly_stats, use_container_width=True, hide_index=True)
                if len(yearly_stats) >= 3:
                    three_yr_avg = yearly_stats["平均值 / Mean (%)"].tail(3).mean()
                    three_yr_median = yearly_stats["中位数 / Median (%)"].tail(3).mean()
                    st.info(f"💡 **近3年平均值: {three_yr_avg:.2f}% | 近3年中位数: {three_yr_median:.2f}%** "
                            f"(OECD TPG Ch.3.79 推荐使用多年平均以减少年度波动影响)")
                else:
                    st.info("💡 数据年度不足3年，建议补充更多年度数据以提高分析可靠性。")

            # Interpretation
            st.markdown("#### 🤖 自动解读 / Auto Interpretation")
            median_val = stats["Median"]
            iqr = stats["Q3 (75th)"] - stats["Q1 (25th)"]
            if median_val > 15:
                interpretation = (
                    f"中位{pli_label} **{median_val:.2f}%** 表明该行业利润率较高。"
                    f"四分位距 (IQR) 为 **{iqr:.2f}%**，说明可比公司之间差异适中。"
                    f"被测试方的利润率应落在 **{stats['Q1 (25th)']:.2f}% - {stats['Q3 (75th)']:.2f}%**"
                    f"（四分位区间）内才算符合独立交易原则。"
                )
            elif median_val > 5:
                interpretation = (
                    f"中位{pli_label} **{median_val:.2f}%** 符合行业常态。"
                    f"四分位距 **{iqr:.2f}%** 提供了合理的独立交易区间。"
                    f"落在 **{stats['Q1 (25th)']:.2f}% - {stats['Q3 (75th)']:.2f}%** 内的结果通常可接受。"
                )
            else:
                interpretation = (
                    f"中位{pli_label} **{median_val:.2f}%** 较低，"
                    f"可能反映薄利商业模式或竞争激烈的市场环境。"
                    f"建议仔细分析功能差异。"
                    f"可接受区间为 **{stats['Q1 (25th)']:.2f}% - {stats['Q3 (75th)']:.2f}%**。"
                )
            st.info(interpretation)

            # ---- FIX #3: Apply WCA from Step 4 ----
            wca_adjustment_pct = 0.0
            if step4_data and "WCA Amount" in step4_data and "Annual Revenue" in step4_data:
                try:
                    wca_amt = float(step4_data["WCA Amount"].replace("$", "").replace("M", ""))
                    ann_rev = float(step4_data["Annual Revenue"].replace("$", "").replace("M", ""))
                    if ann_rev > 0:
                        wca_adjustment_pct = (wca_amt / ann_rev) * 100
                        st.markdown(f"##### 📐 已应用第4步调整 / Applied Step 4 Adjustment")
                        st.info(f"营运资金调整(WCA): ${wca_amt:.4f}M → 占收入 {wca_adjustment_pct:.4f}%\n"
                                f"💡 被测试方利润率应相应调整：实际利润率 + WCA调整比例 = 调整后利润率")
                except Exception:
                    pass

            # ===== 被测试方结果输入 + 合规判定 =====
            st.markdown("---")
            st.markdown("#### 🎯 被测试方结果输入 + 合规判定 / Tested Party Result & Compliance")
            st.markdown(f"输入被测试方的实际{pli_label}，系统自动判定是否落在独立交易区间内。")

            tested_margin_raw = st.number_input(
                f"被测试方实际{pli_label} (%) / Tested Party Result (%)",
                min_value=-50.0, max_value=100.0, value=float(stats["Median"]), step=0.1,
                help=f"输入被测试方的实际{pli_label}（%）。系统将自动判断是否落在独立交易四分位区间内。",
            )

            # Apply WCA adjustment to tested party result
            tested_margin = tested_margin_raw + wca_adjustment_pct
            if wca_adjustment_pct != 0:
                st.caption(f"📊 调整后利润率 = 原始 {tested_margin_raw:.2f}% + WCA {wca_adjustment_pct:.4f}% = **{tested_margin:.2f}%**")

            q1 = stats["Q1 (25th)"]
            q3 = stats["Q3 (75th)"]
            median_val_test = stats["Median"]

            # Compliance determination
            # ---- FIX #10: Remove duplicate method recommendation, add OECD Ch.3.65 guidance ----
            if tested_margin < q1:
                compliance_status = "⚠️ 低于独立交易区间下限 / Below Arm's Length Range"
                compliance_color = "#E74C3C"
                compliance_advice = (
                    f"被测试方利润率 {tested_margin:.2f}% 低于四分位下限 {q1:.2f}%。"
                    "可能存在利润转移风险，建议：\n"
                    "1. 检查功能分析(FAR)是否准确\n"
                    "2. 确认可比性调整是否充分\n"
                    "3. 考虑是否需要扩大可比样本\n"
                    "4. 如确属异常，可能需要启动转让定价调整\n\n"
                    f"📌 OECD TPG Ch.3.65: 税务机关可将被测试方利润率调整至区间内任一点，"
                    f"通常使用中位数 {median_val_test:.2f}%。"
                )
            elif tested_margin > q3:
                compliance_status = "⚠️ 高于独立交易区间上限 / Above Arm's Length Range"
                compliance_color = "#F39C12"
                compliance_advice = (
                    f"被测试方利润率 {tested_margin:.2f}% 高于四分位上限 {q3:.2f}%。"
                    "虽然偏高对税务机关有利，但仍需调查原因：\n"
                    "1. 确认被测试方功能是否比可比公司更复杂\n"
                    "2. 检查是否存在一次性收益\n"
                    "3. 确认可比样本选取是否合理\n\n"
                    f"📌 OECD TPG Ch.3.65: 如确需调整，可调整至中位数 {median_val_test:.2f}%。"
                )
            else:
                compliance_status = "✅ 落在独立交易区间内 / Within Arm's Length Range"
                compliance_color = "#27AE60"
                compliance_advice = (
                    f"被测试方利润率 {tested_margin:.2f}% 落在四分位区间 {q1:.2f}% - {q3:.2f}% 内。"
                    "符合独立交易原则，转让定价安排合规。"
                )

            # Display compliance result — modern card style
            st.markdown(f"""
            <div style='
                background: linear-gradient(135deg, {compliance_color}11 0%, {compliance_color}08 100%);
                border: 2px solid {compliance_color};
                border-radius: 14px;
                padding: 1.8rem;
                text-align: center;
                margin: 1rem 0;
                box-shadow: 0 4px 15px {compliance_color}22;
            '>
                <div style='font-size: 1.5rem; margin-bottom: 0.5rem;'>{'🎯' if '✅' in compliance_status else '⚠️'}</div>
                <h2 style='color: {compliance_color}; margin: 0; font-size: 1.3rem;'>{compliance_status}</h2>
                <div style='
                    display: flex; justify-content: center; gap: 2rem;
                    margin-top: 1rem; flex-wrap: wrap;
                '>
                    <div>
                        <div style='font-size:0.75rem;color:#7F8C8D;text-transform:uppercase;letter-spacing:1px'>被测试方 / Tested Party</div>
                        <div style='font-size:1.5rem;font-weight:800;color:{compliance_color}'>{tested_margin:.2f}%</div>
                    </div>
                    <div>
                        <div style='font-size:0.75rem;color:#7F8C8D;text-transform:uppercase;letter-spacing:1px'>独立交易区间 / AL Range</div>
                        <div style='font-size:1.5rem;font-weight:800;color:{PRIMARY_COLOR}'>{q1:.2f}% - {q3:.2f}%</div>
                    </div>
                    <div>
                        <div style='font-size:0.75rem;color:#7F8C8D;text-transform:uppercase;letter-spacing:1px'>中位数 / Median</div>
                        <div style='font-size:1.5rem;font-weight:800;color:{TEXT_DARK}'>{median_val_test:.2f}%</div>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)

            st.info(f"💡 **分析建议 / Analysis Advice:**\n{compliance_advice}")

            # Visual: Add tested party line to box plot
            fig_box.add_hline(y=tested_margin, line_dash="dot", line_color=compliance_color,
                              line_width=3, annotation_text=f"被测试方 {tested_margin:.1f}%")

            # Show selected method (from Step 2, not a new recommendation)
            st.markdown("#### 📋 已选方法 / Selected Method")
            st.markdown(f"<span class='method-badge'>{step2_data.get('Method Name', 'TNMM')} — {step2_data.get('PLI', '')}</span>", unsafe_allow_html=True)

            # Export report
            st.markdown("---")
            if st.button("📤 导出分析报告 (.docx)", key="export_report"):
                tx_desc = st.session_state.get("wiz_step1", {})
                method_data = step2_data
                screen_params = step3_data
                adj_results = step4_data
                al_range_stats = {k: f"{v:.2f}%" if isinstance(v, float) else str(v) for k, v in stats.items()}
                al_range_stats["PLI Used"] = pli_label
                al_range_stats["WCA Adjustment"] = f"{wca_adjustment_pct:.4f}%"
                al_range_stats["Tested Party (Adjusted)"] = f"{tested_margin:.2f}%"
                al_range_stats["Compliance Status"] = compliance_status

                doc_bytes = generate_word_report(tx_desc, method_data, screen_params, adj_results, al_range_stats, step2_data.get('Method Name', 'TNMM'))
                st.download_button(
                    label="📥 下载Word报告",
                    data=doc_bytes,
                    file_name=f"TP_Analysis_Report_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                st.success("报告生成成功！")
        else:
            st.warning("数据不足，无法确定独立交易区间。请先完成第1-3步筛选。")


# ============== Page 3: Extractive Industry Pricing ==============

def extractive_pricing() -> None:
    """Page 3: Extractive Industry Pricing Workbench - Commodity pricing tools."""
    st.markdown("<div class='main-header'>⛏️ 采掘业定价工作台 / Extractive Industry Pricing</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='sub-header'>为采掘业提供大宗商品定价分析工具 / Commodity Pricing Tools for the Extractive Sector</div>",
        unsafe_allow_html=True,
    )

    st.markdown("""
    <div class='info-bar'>
    💡 <b>本页功能说明 / Page Guide:</b><br>
    <b>中文：</b>专门针对采掘业（采矿、石油、天然气）的转让定价分析工具。<br>
    &nbsp;&nbsp;• <b>大宗商品价格查询</b>：查看各类矿产、能源的国际市场价格走势<br>
    &nbsp;&nbsp;• <b>可比价格分析</b>：将关联交易价格与市场公开价格对比<br>
    <b>English:</b> Transfer pricing analysis tools for the extractive sector (mining, oil & gas).<br>
    &nbsp;&nbsp;• <b>Commodity Prices</b>: View international market price trends<br>
    &nbsp;&nbsp;• <b>Comparable Price Analysis</b>: Compare related-party prices with market prices<br>
    <br>
    <i>依据 / Reference: 联合国国际税收合作框架公约第6条（税基侵蚀）及联合国采掘业税收指南<br>
    UN Framework Convention Art. 6 (Base Erosion) & UN Extractive Industries Taxation Guide</i>
    </div>
    """, unsafe_allow_html=True)

    # Help
    if st.button("Help", key="help_extractive", help="Guidance on extractive industry pricing"):
        st.info("""
        **采掘业定价工作台 — 使用帮助 / Extractive Industry Pricing Help**

        本工作台为采掘业提供大宗商品定价分析工具：
        This workbench provides tools for pricing commodities in the extractive sector:

        - **交易所价格 / Exchange Prices**: 主要大宗商品的市场参考价格 / Reference market prices for major commodities
        - **CUP计算器 / CUP Calculator**: 计算矿产销售的独立交易价格 / Calculate the Comparable Uncontrolled Price for mineral sales
        - **可比交易 / Comparable Transactions**: 近期可比交易参考表 / Reference table of recent comparable transactions

        当有公认交易所报价时，CUP法通常是采掘业最适当的方法。
        The CUP method is typically the most appropriate method for commodity transactions
        where quoted prices are available on recognized exchanges.
        """)

    # Commodity selector
    commodity = st.selectbox(
        "选择商品 / Select Commodity",
        ["Copper Concentrate / 铜精矿", "Iron Ore (Fe 62%) / 铁矿石", "Gold Dore (Au 85%) / 金多尔"],
        help="选择要分析的大宗商品类型 / Select the commodity type for pricing analysis",
    )

    # Exchange prices
    st.markdown("---")
    st.subheader("交易所参考价格（30日趋势）/ Reference Exchange Prices (30-Day Trend)")

    commodity_data = load_commodity_data()
    # Match commodity by keyword
    if "Copper" in commodity or "铜" in commodity:
        selected_commodity = "Copper"
    elif "Iron" in commodity or "铁" in commodity:
        selected_commodity = "Iron Ore"
    elif "Gold" in commodity or "金" in commodity:
        selected_commodity = "Gold"
    else:
        selected_commodity = "Copper"

    comm_data = commodity_data[commodity_data["commodity"] == selected_commodity].sort_values("date")

    if len(comm_data) > 0:
        latest_price = comm_data.iloc[-1]["price_usd"]
        prev_price = comm_data.iloc[-2]["price_usd"] if len(comm_data) > 1 else latest_price
        price_change = latest_price - prev_price

        kpi_cols = st.columns(4)
        with kpi_cols[0]:
            unit = "USD/tonne" if selected_commodity in ["Copper", "Iron Ore"] else "USD/oz"
            st.metric(f"最新价格 / Latest Price ({unit})", f"${latest_price:,.2f}", f"${price_change:,.2f}")
        with kpi_cols[1]:
            avg_price = comm_data["price_usd"].mean()
            st.metric("30日均 / 30-Day Avg", f"${avg_price:,.2f}")
        with kpi_cols[2]:
            high_price = comm_data["price_usd"].max()
            st.metric("30日最高 / 30-Day High", f"${high_price:,.2f}")
        with kpi_cols[3]:
            low_price = comm_data["price_usd"].min()
            st.metric("30日最低 / 30-Day Low", f"${low_price:,.2f}")

        # Price trend chart
        fig_trend = go.Figure()
        fig_trend.add_trace(go.Scatter(
            x=comm_data["date"], y=comm_data["price_usd"],
            mode="lines+markers", name=selected_commodity,
            line=dict(color=PRIMARY_COLOR, width=2),
            marker=dict(size=6),
        ))
        fig_trend.update_layout(
            title=f"{selected_commodity} 价格趋势（近30天）/ Price Trend",
            xaxis_title="日期 / Date",
            yaxis_title=f"价格 / Price ({unit})",
            height=400,
            margin=dict(l=0, r=0, t=40, b=0),
        )
        st.plotly_chart(fig_trend, use_container_width=True)

    # CUP Calculator
    st.markdown("---")
    st.subheader("CUP计算器 — 矿产销售定价 / CUP Calculator - Mineral Sale Pricing")
    st.markdown("""
    使用可比非受控价格法(CUP)计算矿产销售的独立交易价格。
    Calculate the arm's length price for mineral sales using the Comparable Uncontrolled Price method.

    📌 **参考交易所 / Reference Exchanges:**
    - 铜 / Copper: **LME (London Metal Exchange)** 或 SHFE (上海期货交易所)
    - 铁矿石 / Iron Ore: **Platts / The Steel Index (TSI)** 或 青岛港CBF指数
    - 黄金 / Gold: **LBMA (London Bullion Market Association)** 或 COMEX

    💡 CUP法在采掘业中是最可靠的方法，因为大宗商品有公开透明的市场价格。
    税务机关验证CUP时会检查：①价格来源是否权威；②品位/质量调整是否合理；③计价货币和 Incoterms 是否一致。
    The CUP method is the most reliable for extractive industries due to transparent commodity prices.
    Tax authorities verify: ① price source authority; ② grade/quality adjustments; ③ currency and Incoterms consistency.
    """)

    calc_col1, calc_col2 = st.columns(2)
    with calc_col1:
        spot_price = st.number_input("现货/参考价格 / Spot Price (USD)", min_value=0.0, value=float(latest_price if len(comm_data) > 0 else 0), step=10.0,
                                    help="交易所当前现货价格或参考价 / Current spot price or reference price on the exchange")
        grade_pct = st.number_input("品位/化验值 / Grade/Assay (%)", min_value=0.0, max_value=100.0, value=30.0 if "Copper" in commodity or "铜" in commodity else (62.0 if "Iron" in commodity or "铁" in commodity else 85.0), step=0.1,
                                   help="矿石品位或化验百分比 / Mineral grade or assay percentage")
    with calc_col2:
        tc_rc = st.number_input("加工精炼费 / TC/RC (USD/tonne)", min_value=0.0, value=100.0 if "Copper" in commodity or "铜" in commodity else 0.0, step=1.0,
                               help="处理和精炼费用(适用于铜精矿) / Treatment and refining charges (for copper concentrate)")
        freight = st.number_input("运费及保险 / Freight & Insurance (USD/tonne)", min_value=0.0, value=25.0, step=1.0,
                                 help="运输和保险成本 / Transportation and insurance costs")

    # Calculate
    if "Copper" in commodity or "铜" in commodity:
        payable_metal = spot_price * (grade_pct / 100)
        deductions = tc_rc + freight
        net_price = payable_metal - deductions
        payable_formula = f"现货价 × 品位% = ${spot_price:,.2f} × {grade_pct}% = ${payable_metal:,.2f}/吨\nSpot Price × Grade% = ${spot_price:,.2f} × {grade_pct}% = ${payable_metal:,.2f}/tonne"
    elif "Iron" in commodity or "铁" in commodity:
        fe_adjustment = (grade_pct - 62) * 0.01 * spot_price  # Simplified
        net_price = spot_price + fe_adjustment - freight
        payable_formula = f"基准价 + 铁品位调整 - 运费\nBase Price + Fe Adjustment - Freight"
    else:  # Gold
        payable_metal = spot_price * (grade_pct / 100)
        deductions = freight
        net_price = payable_metal - deductions
        payable_formula = f"现货价 × 含金量% = ${spot_price:,.2f} × {grade_pct}%\nSpot Price × Gold Content% = ${spot_price:,.2f} × {grade_pct}%"

    st.markdown("---")
    result_cols = st.columns(3)
    with result_cols[0]:
        unit_label = "/tonne" if "Gold" not in commodity and "金" not in commodity else "/oz"
        st.metric("计算净价 / Net Price", f"${net_price:,.2f}{unit_label}")
    with result_cols[1]:
        st.metric("扣减项 / Deductions", f"${deductions:,.2f}")
    with result_cols[2]:
        margin = (net_price / spot_price * 100) if spot_price > 0 else 0
        st.metric("净价/现货比 / Net/Spot", f"{margin:.1f}%")

    st.caption(f"计算公式 / Formula: {payable_formula}")

    # Comparable transactions reference
    st.markdown("---")
    st.subheader("可比交易参考 / Comparable Transactions Reference")

    np.random.seed(88)
    tx_data = []
    for i in range(15):
        tx_data.append({
            "date": (datetime.now() - timedelta(days=np.random.randint(1, 180))).strftime("%Y-%m-%d"),
            "commodity": np.random.choice(["Copper Concentrate", "Iron Ore", "Gold Dore"]),
            "buyer": f"Refiner {i+1:02d}",
            "seller": f"Mine {i+1:02d}",
            "volume_mt": int(np.random.uniform(1000, 50000)),
            "price_usd": round(np.random.uniform(80, 9000), 2),
            "terms": np.random.choice(["FOB", "CIF", "CFR"]),
            "basis": np.random.choice(["LME", "Platts", "LBMA", "Spot"]),
        })
    tx_df = pd.DataFrame(tx_data)
    st.dataframe(tx_df, use_container_width=True, hide_index=True)


# Fix variable name typo



# ============== Page 4: Country Policy Lookup ==============

def country_policy() -> None:
    """Page 4: Country Policy Lookup - National TP regulation and treaty information."""
    st.markdown("<div class='main-header'>📋 各经济体转让定价政策查询 / Economy Policy Lookup</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='sub-header'>查询各国的转让定价法规、文档要求和条约网络 / National TP Regulations, Documentation & Treaty Networks</div>",
        unsafe_allow_html=True,
    )

    st.markdown("""
    <div class='info-bar'>
    💡 <b>本页功能说明 / Page Guide:</b><br>
    <b>中文：</b>查询各经济体转让定价政策的「百科全书」。<br>
    &nbsp;&nbsp;• <b>法规要求</b>：各国对转让定价文档的要求<br>
    &nbsp;&nbsp;• <b>争议解决机制</b>：APA、MAP、仲裁机制<br>
    &nbsp;&nbsp;• <b>安全港规则</b>：简化合规的安全港规定<br>
    <b>English:</b> Encyclopedia of national transfer pricing policies.<br>
    &nbsp;&nbsp;• <b>Regulations</b>: TP documentation requirements by economy<br>
    &nbsp;&nbsp;• <b>Dispute Resolution</b>: APA, MAP, and arbitration mechanisms<br>
    &nbsp;&nbsp;• <b>Safe Harbour</b>: Simplified compliance rules<br>
    <br>
    <i>依据 / Reference: 联合国国际税收合作框架公约第13条（争议预防与解决）及联合国税收协定范本<br>
    UN Framework Convention Art. 13 (Dispute Prevention & Resolution) & UN Model Tax Convention</i>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Help", key="help_policy", help="Guidance on economy policy lookup"):
        st.info("""
        **Economy Policy Lookup Help**

        Search for an economy to view its transfer pricing policy framework:
        - TP regulations and legislative basis
        - Documentation requirements (Local File, Master File, CbCR)
        - Safe harbor rules and thresholds
        - Tax treaty network size
        - MAP/APA availability and statistics
        - Information exchange mechanisms

        Use the **Compare** feature to view two countries side by side.
        """)

    df_policies = load_country_policies()

    # Search / select
    search_col, compare_toggle_col = st.columns([3, 1])
    with search_col:
        selected_country = st.selectbox(
            "选择经济体 / Select Economy",
            [""] + sorted(df_policies["country"].tolist()),
            help="选择一个经济体查看转让定价政策详情 / Choose an economy to view its TP policy details",
        )
    with compare_toggle_col:
        compare_mode = st.toggle("对比模式 / Compare Mode", help="并排对比两个经济体 / Enable side-by-side comparison of two economies")

    if compare_mode:
        col_a, col_b = st.columns(2)
        with col_a:
            country_a = st.selectbox("经济体A / Economy A", [""] + sorted(df_policies["country"].tolist()), key="country_a")
        with col_b:
            country_b = st.selectbox("经济体B / Economy B", [""] + sorted(df_policies["country"].tolist()), key="country_b")

        if country_a and country_b:
            row_a = df_policies[df_policies["country"] == country_a].iloc[0]
            row_b = df_policies[df_policies["country"] == country_b].iloc[0]

            for label, key in [
                ("TP法规 / TP Regulations", "has_tp_regulations"),
                ("文档要求 / Documentation", "documentation_req"),
                ("安全港 / Safe Harbor", "safe_harbor"),
                ("税收协定数 / Tax Treaties", "treaty_count"),
                ("MAP案件(2024) / MAP Cases", "map_cases_2024"),
                ("APA可用 / APA Available", "apa_available"),
                ("信息交换 / Exchange Mechanism", "exchange_mechanism"),
            ]:
                c1, c2 = st.columns(2)
                with c1:
                    val = row_a[key]
                    display = "Yes" if val is True else ("No" if val is False else str(val))
                    st.metric(f"{label} - {country_a}", display)
                with c2:
                    val = row_b[key]
                    display = "Yes" if val is True else ("No" if val is False else str(val))
                    st.metric(f"{label} - {country_b}", display)
                st.markdown("---")

            # Comparison chart
            fig_comp = go.Figure()
            fig_comp.add_trace(go.Bar(
                x=["Tax Treaties", "MAP Cases (2024)"],
                y=[row_a["treaty_count"], row_a["map_cases_2024"]],
                name=country_a,
                marker_color=PRIMARY_COLOR,
            ))
            fig_comp.add_trace(go.Bar(
                x=["Tax Treaties", "MAP Cases (2024)"],
                y=[row_b["treaty_count"], row_b["map_cases_2024"]],
                name=country_b,
                marker_color=ACCENT_GREEN,
            ))
            fig_comp.update_layout(
                barmode="group",
                title=f"对比: {country_a} vs {country_b} / Comparison",
                height=350,
            )
            st.plotly_chart(fig_comp, use_container_width=True)

    elif selected_country:
        row = df_policies[df_policies["country"] == selected_country].iloc[0]

        st.markdown(f"## {selected_country}")

        # Top metrics
        mcol1, mcol2, mcol3, mcol4 = st.columns(4)
        with mcol1:
            st.metric("TP法规 / TP Regs", "有/Yes" if row["has_tp_regulations"] else "无/No")
        with mcol2:
            st.metric("税收协定 / Tax Treaties", row["treaty_count"])
        with mcol3:
            st.metric("MAP案件(2024) / MAP Cases", row["map_cases_2024"])
        with mcol4:
            st.metric("APA可用性 / APA Available", "有/Yes" if row["apa_available"] else "无/No")

        st.markdown("---")

        # Detail sections
        dcol1, dcol2 = st.columns(2)
        with dcol1:
            st.markdown("#### 文档要求 / Documentation Requirements")
            st.info(row["documentation_req"])

            st.markdown("#### 安全港规则 / Safe Harbor Rules")
            st.markdown("""
            💡 **什么是安全港？** 安全港(Safe Harbour)是指税法规定的简化规则——如果关联交易满足特定条件
            （如金额低于门槛、利润率达标），企业无需进行完整的转让定价分析即可视为合规。
            安全港降低了发展中国家的合规成本和税务机关的审计成本。
            依据：UN TP Manual Ch.4 / OECD TPG Ch.4
            """)
            if row["safe_harbor"]:
                st.success("✅ 该经济体设有安全港规则。/ Safe harbor rules are available in this jurisdiction.")
            else:
                st.warning("⚠️ 该经济体目前无安全港规则。/ No safe harbor rules currently in effect.")

        with dcol2:
            st.markdown("#### 信息交换机制 / Information Exchange")
            st.info(row["exchange_mechanism"])

            st.markdown("#### MAP/APA 框架 / MAP/APA Framework")
            if row["apa_available"]:
                st.success("APA项目已运行。可通过税收协定启动MAP。/ APA program is operational. MAP is available under tax treaties.")
            else:
                st.warning("APA暂不可用。MAP可能通过税收协定网络进行。/ APA not yet available. MAP may be accessible through treaty network.")

        # Policy summary text
        st.markdown("---")
        st.markdown("#### 转让定价法规摘要 / TP Regulation Summary")
        summary_text = (
            f"**{selected_country}** has {'comprehensive' if row['has_tp_regulations'] else 'limited'} "
            f"transfer pricing regulations requiring **{row['documentation_req']}**. "
            f"The economy maintains **{row['treaty_count']} tax treaties** providing MAP access, "
            f"and {'operates' if row['apa_available'] else 'does not yet operate'} an APA program. "
            f"Information exchange is conducted through **{row['exchange_mechanism']}**. "
            f"Safe harbor provisions are {'available' if row['safe_harbor'] else 'not currently available'}."
        )
        st.markdown(summary_text)

        # Treaty network visualization (simplified)
        st.markdown("---")
        st.markdown("#### Tax Treaty Network (Illustrative)")
        treaty_count = row["treaty_count"]
        # Create a simple radial chart representation
        angles = np.linspace(0, 2 * np.pi, max(treaty_count, 10), endpoint=False)
        r_values = np.ones(len(angles))
        fig_treaty = go.Figure()
        fig_treaty.add_trace(go.Scatterpolar(
            r=r_values,
            theta=angles * 180 / np.pi,
            mode="markers",
            marker=dict(size=12, color=PRIMARY_COLOR),
            name="Treaty Partners",
        ))
        fig_treaty.add_trace(go.Scatterpolar(
            r=[0],
            theta=[0],
            mode="markers",
            marker=dict(size=20, color=ACCENT_RED),
            name=selected_country,
        ))
        fig_treaty.update_layout(
            polar=dict(radialaxis=dict(visible=False, range=[0, 1.5])),
            showlegend=True,
            title=f"Treaty Network: {treaty_count} agreements",
            height=400,
        )
        st.plotly_chart(fig_treaty, use_container_width=True)
    else:
        # Show all countries overview table
        st.info("请在上方选择一个经济体查看详细政策信息，或启用对比模式。/ Select an economy above to view detailed policy information, or enable Compare Mode.")
        st.subheader("各经济体总览 / Countries Overview")
        display_df = df_policies.copy()
        display_df["TP Regs"] = display_df["has_tp_regulations"].apply(lambda x: "Yes" if x else "No")
        display_df["Safe Harbor"] = display_df["safe_harbor"].apply(lambda x: "Yes" if x else "No")
        display_df["APA"] = display_df["apa_available"].apply(lambda x: "Yes" if x else "No")
        st.dataframe(
            display_df[["country", "TP Regs", "documentation_req", "Safe Harbor", "treaty_count", "map_cases_2024", "APA", "exchange_mechanism"]],
            use_container_width=True,
            hide_index=True,
        )

    # ===== Documentation Compliance Checker =====
    st.markdown("---")
    st.subheader("📝 三层文档合规检查器 / Three-Tier Documentation Compliance Checker")
    st.markdown("""
    💡 **功能说明 / Guide:** 输入跨国企业集团合并收入，系统自动判断该企业是否需要准备三层转让定价文档
    （主文件、本地文件、国别报告），并列出相关经济体的文档要求。

    依据：UN TP Manual Ch.4 + OECD TPG Ch.5 + BEPS Action 13

    ⚠️ **注意 / Note:** 不同经济体文档门槛各不相同。以下检查基于通用国际标准，
    具体门槛请以各经济体当地法规为准。
    """)

    # Jurisdiction selector for thresholds
    jur_options = ["通用国际标准 / General International", "中国 (42号公告) / China (STA No.42)", "欧盟 / EU", "印度 / India"]
    sel_jur = st.selectbox("选择适用经济体 / Select Jurisdiction", jur_options,
                           help="不同经济体的文档门槛不同")

    # Thresholds by jurisdiction
    jur_thresholds = {
        "通用国际标准 / General International": {"master": 750, "local": 50, "cbc": 750, "unit": "M EUR"},
        "中国 (42号公告) / China (STA No.42)": {"master": 5000, "local": 2000, "cbc": 5500, "unit": "M RMB (约€65M/€26M/€720M)"},
        "欧盟 / EU": {"master": 750, "local": 50, "cbc": 750, "unit": "M EUR"},
        "印度 / India": {"master": 0, "local": 0, "cbc": 750, "unit": "M EUR (CbC: ₹4000M)"},
    }
    thresholds = jur_thresholds[sel_jur]
    unit_display = thresholds["unit"]

    if "中国" in sel_jur:
        mne_revenue = st.number_input(
            "集团合并年收入（百万人民币）/ Group Revenue (M RMB)",
            min_value=0, max_value=1000000, value=5500, step=500,
            help="中国42号公告：国别报告门槛为年度合并收入≥55亿RMB",
        )
    else:
        mne_revenue = st.number_input(
            "集团合并年收入（百万欧元）/ Group Revenue (M EUR)",
            min_value=0, max_value=100000, value=850, step=50,
            help="BEPS第13项行动计划：年收入≥7.5亿欧元需准备国别报告",
        )

    check_col1, check_col2, check_col3 = st.columns(3)
    with check_col1:
        needs_master = mne_revenue >= thresholds["master"]
        icon_m = "✅" if needs_master else "❌"
        threshold_label_m = f"门槛: {thresholds['master']} ({unit_display.split('(')[0].strip()})"
        st.metric("主文件 / Master File", f"{'需要' if needs_master else '不需要'}", icon_m, help=threshold_label_m)
    with check_col2:
        needs_local = mne_revenue >= thresholds["local"]
        icon_l = "✅" if needs_local else "❌"
        threshold_label_l = f"门槛: {thresholds['local']} ({unit_display.split('(')[0].strip()})"
        st.metric("本地文件 / Local File", f"{'需要' if needs_local else '不需要'}", icon_l, help=threshold_label_l)
    with check_col3:
        needs_cbc = mne_revenue >= thresholds["cbc"]
        icon_c = "✅" if needs_cbc else "❌"
        threshold_label_c = f"门槛: {thresholds['cbc']} ({unit_display.split('(')[0].strip()})"
        st.metric("国别报告 / CbC Report", f"{'需要' if needs_cbc else '不需要'}", icon_c, help=threshold_label_c)

    if needs_cbc:
        st.success(f"📋 该企业需准备完整三层文档 / Full three-tier documentation required.")
        if "中国" in sel_jur:
            st.info("""
            **文档提交时间表 / Filing Timeline:**

            🇨🇳 **中国（国家税务总局公告2016年第42号）/ China (STA Announcement No. 42, 2016):**
            - 主文件 / Master File: **财年次年6月30日前** / By June 30 of the following year
            - 本地文件 / Local File: **财年次年6月30日前** / By June 30 of the following year
            - 国别报告 / CbC Report: **财年次年12月31日前** / By December 31 of the following year

            📌 **门槛 / Thresholds:**
            - 国别报告: 集团合并收入 ≥ **55亿RMB** (约€720M)
            - 主文件: 集团合并收入 ≥ **50亿RMB** (约€650M) 且有跨境关联交易
            - 本地文件: 集团合并收入 ≥ **2亿RMB** (约€26M) 且有关联交易

            🌍 **OECD BEPS Action 13 通用要求 / General OECD Requirement:**
            - Master File: Within 12 months of fiscal year-end
            - Local File: Within 12 months of fiscal year-end
            - CbC Report: Within 12 months of the reporting fiscal year-end (Ultimate Parent Entity)

            ⚠️ 各经济体可能有不同截止日期，请以当地法规为准。
            Filing deadlines vary by jurisdiction — always verify with local regulations.
            """)
        else:
            st.info("""
            **文档提交时间表 / Filing Timeline (OECD General):**
            - Master File: Within 12 months of fiscal year-end
            - Local File: Within 12 months of fiscal year-end
            - CbC Report: Within 12 months of the reporting fiscal year-end (Ultimate Parent Entity)

            ⚠️ 各经济体可能有不同截止日期和门槛，请以当地法规为准。
            """)
    elif needs_master or needs_local:
        st.warning("📋 该企业需准备主文件+本地文件，无需国别报告 / Master + Local required, CbC not required.")
    else:
        st.info("📋 该企业无需强制准备三层文档，但建议自愿准备以降低税务风险 / Not mandatory, but recommended.")


# ============== Page 5: Data Contribution Hub ==============

def data_contribution() -> None:
    """Page 5: Data Contribution Hub - Upload, template, tiered management."""
    st.markdown("<div class='main-header'>📤 数据贡献管理 / Data Contribution Hub</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='sub-header'>向全球转让定价数据库贡献可比公司数据 — 含导入模板与密级分层管理</div>",
        unsafe_allow_html=True,
    )

    st.markdown("""
    <div class='info-bar'>
    💡 <b>本页功能说明 / Page Guide:</b><br>
    <b>中文：</b>管理各经济体向数据库贡献数据的「工作台」，包含三大功能：<br>
    &nbsp;&nbsp;• <b>📥 导入模板下载</b>：提供标准化CSV模板，确保数据格式统一<br>
    &nbsp;&nbsp;• <b>🔐 密级分层管理</b>：Tier 1/2/3 三级数据分类，不同密级有不同访问权限和用途<br>
    &nbsp;&nbsp;• <b>📊 贡献统计</b>：查看各国的数据贡献排名、密级分布和进度<br>
    <b>English:</b> Workspace for economy data contributions with three core functions:<br>
    &nbsp;&nbsp;• <b>📥 Import Template</b>: Standardized CSV template for consistent data format<br>
    &nbsp;&nbsp;• <b>🔐 Tiered Classification</b>: Tier 1/2/3 data classification with differentiated access<br>
    &nbsp;&nbsp;• <b>📊 Contribution Stats</b>: Economy rankings, tier distribution, and progress<br>
    <br>
    <i>依据 / Reference: 联合国国际税收合作框架公约第11条（能力建设）<br>
    UN Framework Convention Art. 11 (Capacity Building) — Data sharing to support developing economies</i>
    </div>
    """, unsafe_allow_html=True)

    # ===== Tab 1: Import Template =====
    # ===== Tab 2: Upload =====
    # ===== Tab 3: Tier Management =====
    # ===== Tab 4: Stats =====
    contrib_tabs = st.tabs([
        "📥 导入模板 / Import Template",
        "📤 数据上传 / Upload",
        "🔐 密级管理 / Tier Management",
        "📊 贡献统计 / Statistics",
    ])

    # ========== Tab 1: Import Template ==========
    with contrib_tabs[0]:
        st.subheader("📥 导入模板下载 / Download Import Template")

        st.markdown("""
        **中文说明：** 请先下载标准化CSV模板，按模板格式填写数据后上传。
        模板包含字段说明（第二个Sheet），请仔细阅读每个字段的含义和要求。

        **English:** Download the standardized CSV template first. Fill in your data following
        the template format, then upload it. The template includes a field guide sheet explaining
        each field's meaning and requirements.
        """)

        # Generate template
        template_col1, template_col2 = st.columns([1, 2])
        with template_col1:
            template_format = st.radio(
                "选择格式 / Format",
                ["CSV", "Excel (.xlsx)"],
                help="CSV兼容性最好；Excel带字段说明"
            )

            # Always render download button directly (no nested button)
            template_data = generate_import_template()
            if template_format == "CSV":
                csv_str = template_data.to_csv(index=False)
                st.download_button(
                    label="📥 下载CSV模板 / Download CSV",
                    data=csv_str.encode('utf-8-sig'),
                    file_name="TP_Comparable_Data_Template.csv",
                    mime="text/csv",
                )
            else:
                import io
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    template_data.to_excel(writer, sheet_name='Data', index=False)
                    field_guide = pd.DataFrame({
                        '字段名 Field': ['company_id','company_name','country_code','industry_code_isic','revenue','net_worth','fiscal_year','data_source','data_tier'],
                        '说明 Description': ['公司ID / Company ID','公司名称 / Company name','经济体代码 / Economy code','ISIC行业代码 / ISIC code','年收入 / Revenue','净资产 / Net worth','财政年度 / Fiscal year','数据来源 / Data source','密级 / Tier'],
                        '必填 / Required': ['✅','✅','✅','✅','✅','✅','✅','✅','✅'],
                    })
                    field_guide.to_excel(writer, sheet_name='Field Guide', index=False)
                st.download_button(
                    label="📥 下载Excel模板 / Download Excel",
                    data=output.getvalue(),
                    file_name="TP_Comparable_Data_Template.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )
            st.success("✅ 点击上方按钮下载 / Click button above to download")

        with template_col2:
            st.markdown("##### 📋 模板字段预览 / Template Fields Preview")
            st.markdown("""
            | 字段 Field | 说明 Description | 必填 |
            |---|---|---|
            | company_id | 公司唯一ID | ✅ |
            | company_name | 公司名称（可匿名）| ✅ |
            | country_code | ISO 3国别代码 | ✅ |
            | industry_code_isic | ISIC行业代码 | ✅ |
            | revenue | 年收入（美元）| ✅ |
            | net_worth | 净资产（美元）| ✅ |
            | fiscal_year | 财政年度 | ✅ |
            | data_source | 数据来源 | ✅ |
            | data_tier | 数据密级 | ✅ |
            | cogs | 销售成本 | ❌ |
            | operating_margin | 营业利润率 | ❌ |
            | functions_performed | 履行功能 | ❌ |

            💡 **提示 / Tip:** 下载Excel版本可获得完整的字段说明（第二个Sheet）。
            """)

    # ========== Tab 2: Upload ==========
    with contrib_tabs[1]:
        st.subheader("📤 数据上传 / Upload Data")

        tier = st.selectbox(
            "选择数据密级 / Select Data Tier",
            [
                "Tier1 — 审计级数据 / Audited (Highest)",
                "Tier2 — 审阅级数据 / Reviewed",
                "Tier3 — 估算级数据 / Estimated",
            ],
            help="Tier1提供最详细可靠的数据；Tier3为汇总估算数据",
        )

        tier_info = {
            "Tier1 — 审计级数据 / Audited (Highest)": """
            **🔐 Tier 1 — 审计级数据 / Audited Data**
            - 来源：经审计的财务报表、公开上市年报
            - 可信度：最高 ✅✅✅
            - 允许用途：可比性分析、独立交易区间确定、税务审计参考
            - 访问权限：所有成员国税务机关
            - 来源 / Source: Audited financial statements, public filings
            - Reliability: Highest
            """,
            "Tier2 — 审阅级数据 / Reviewed": """
            **🟡 Tier 2 — 审阅级数据 / Reviewed Data**
            - 来源：税务机关评估报告、国别报告(CbC)、纳税申报数据
            - 可信度：较高 ✅✅
            - 允许用途：趋势分析、行业基准、初步筛选
            - 访问权限：签署信息交换协议的国家
            - 来源 / Source: Tax authority assessments, CbC reports
            - Reliability: Moderate-high
            """,
            "Tier3 — 估算级数据 / Estimated": """
            **🟠 Tier 3 — 估算级数据 / Estimated Data**
            - 来源：行业协会统计、第三方数据库估算、宏观经济推算
            - 可信度：一般 ✅
            - 允许用途：宏观分析、数据缺口填补、初步参考
            - 访问权限：仅限汇总统计，不披露单一公司
            - 来源 / Source: Industry statistics, third-party estimates
            - Reliability: Reference only
            """,
        }
        st.info(tier_info[tier])

        # File upload
        uploaded_file = st.file_uploader(
            "上传数据文件 / Upload Data File",
            type=["csv", "xlsx", "xls"],
            help="请使用模板格式上传 / Please use the template format",
        )

        # Data sovereignty declaration
        st.markdown("---")
        st.subheader("📋 数据主权声明 / Data Sovereignty Declaration")
        declaration = st.checkbox(
            "本人确认上传数据已匿名化处理，贡献国已授权在联合国国际税收合作框架公约下共享。"
            "数据不含个人身份信息或可识别单一纳税人的商业敏感信息。\n\n"
            "I confirm that the data is anonymized and the contributing economy has authorized "
            "its sharing under the UN Framework Convention. The data does not contain personally "
            "identifiable information or commercially sensitive information.",
            help="必填项 / Required",
        )

        if uploaded_file and declaration:
            st.success("✅ 文件上传成功！数据将进入审核流程。/ File uploaded! Data will enter the review process.")
            st.info(f"📄 文件 / File: **{uploaded_file.name}** | 大小 / Size: **{uploaded_file.size:,} bytes** | 密级 / Tier: **{tier[:6]}**")

            # Simulated validation
            st.markdown("##### 🔍 自动验证结果 / Automated Validation Results")
            val_col1, val_col2, val_col3 = st.columns(3)
            with val_col1:
                st.metric("必填字段完整率 / Required Fields", "100%", "✅")
            with val_col2:
                st.metric("经济体代码有效 / Valid Economy Codes", "98%", "✅")
            with val_col3:
                st.metric("数据合理性 / Reasonableness", "95%", "✅")

        elif uploaded_file and not declaration:
            st.warning("⚠️ 请先勾选数据主权声明。/ Please check the Data Sovereignty Declaration.")

    # ========== Tab 3: Tier Management ==========
    with contrib_tabs[2]:
        st.subheader("🔐 密级分层管理 / Tiered Data Classification Management")

        st.markdown("""
        <div class='info-bar'>
        <b>中文说明：</b>数据库采用三级密级分类体系，不同密级的数据有不同的来源、可信度和使用场景。<br>
        点击下方各密级标签查看详细说明和当前数据量。<br><br>
        <b>English:</b> The database uses a three-tier classification system. Each tier has different
        data sources, reliability levels, and permitted use cases.
        </div>
        """, unsafe_allow_html=True)

        # Load actual tier statistics from database
        tier_df = db_query("""
            SELECT
                data_tier,
                COUNT(*) as company_count,
                COUNT(DISTINCT country_code) as country_count,
                AVG(operating_margin) as avg_margin
            FROM comparable_companies
            GROUP BY data_tier
            ORDER BY data_tier
        """)

        if tier_df is None or tier_df.empty:
            tier_df = pd.DataFrame({
                'data_tier': ['Tier1', 'Tier2', 'Tier3'],
                'company_count': [495, 258, 281],
                'country_count': [85, 60, 45],
                'avg_margin': [12.5, 9.8, 7.2],
            })

        # Tier cards
        tier_col1, tier_col2, tier_col3 = st.columns(3)

        tier_colors = {'Tier1': '#28a745', 'Tier2': '#ffc107', 'Tier3': '#fd7e14'}
        tier_names = {
            'Tier1': 'Tier 1 — 审计级 / Audited',
            'Tier2': 'Tier 2 — 审阅级 / Reviewed',
            'Tier3': 'Tier 3 — 估算级 / Estimated',
        }
        tier_descs = {
            'Tier1': '经审计的财务报表\nAudited financial statements',
            'Tier2': '税务机关评估、CbC报告\nTax authority assessments',
            'Tier3': '行业估算、推算数据\nIndustry estimates',
        }
        tier_access = {
            'Tier1': '所有成员国 / All member states',
            'Tier2': '信息交换协议国 / Exchange agreement',
            'Tier3': '仅汇总统计 / Aggregated only',
        }

        for i, (col, tier_key) in enumerate(zip([tier_col1, tier_col2, tier_col3], ['Tier1', 'Tier2', 'Tier3'])):
            row = tier_df[tier_df['data_tier'] == tier_key]
            count = int(row['company_count'].iloc[0]) if not row.empty else 0
            countries = int(row['country_count'].iloc[0]) if not row.empty else 0
            with col:
                color = tier_colors[tier_key]
                st.markdown(f"""
                <div style='border: 2px solid {color}; border-radius: 10px; padding: 1rem; text-align: center;'>
                    <h3 style='color: {color}; margin: 0;'>{tier_names[tier_key]}</h3>
                    <p style='font-size: 0.85rem; color: #666;'>{tier_descs[tier_key]}</p>
                    <h2 style='margin: 0.5rem 0;'>{count:,}</h2>
                    <p style='font-size: 0.8rem;'>家公司 / companies</p>
                    <p style='font-size: 0.8rem;'>覆盖 {countries} 个国家 / {countries} countries</p>
                    <hr style='border: 0.5px solid #ddd;'>
                    <p style='font-size: 0.75rem;'>🔒 访问权限 / Access:<br>{tier_access[tier_key]}</p>
                </div>
                """, unsafe_allow_html=True)

        # Tier comparison table
        st.markdown("---")
        st.markdown("##### 📊 密级对比表 / Tier Comparison Table")
        comparison_data = pd.DataFrame({
            '密级 / Tier': ['Tier 1', 'Tier 2', 'Tier 3'],
            '来源 / Source': [
                '审计财务报表、公开年报\nAudited financials, public filings',
                '税务评估、CbC报告、纳税申报\nTax assessments, CbC reports',
                '行业统计、第三方估算\nIndustry stats, estimates',
            ],
            '可信度 / Reliability': ['⭐⭐⭐ 最高', '⭐⭐ 较高', '⭐ 一般'],
            '允许用途 / Permitted Use': [
                '可比性分析、独立交易区间确定\nFull comparability analysis',
                '趋势分析、行业基准、初步筛选\nTrend analysis, benchmarking',
                '宏观分析、缺口填补\nMacro analysis, gap filling',
            ],
            '访问权限 / Access': [
                '所有成员国 / All members',
                '信息交换协议国 / Exchange agreement',
                '仅汇总，不披露单一 / Aggregated only',
            ],
        })
        st.dataframe(comparison_data, use_container_width=True, hide_index=True)

        # Tier distribution by country
        st.markdown("---")
        st.markdown("##### 🌍 各密级数据经济体分布 / Tier Distribution by Economy")
        tier_country_df = db_query("""
            SELECT c.country_name, cc.data_tier, COUNT(*) as cnt
            FROM comparable_companies cc
            JOIN countries c ON c.country_code = cc.country_code
            GROUP BY c.country_name, cc.data_tier
            ORDER BY cnt DESC
            LIMIT 20
        """)
        if tier_country_df is not None and not tier_country_df.empty:
            fig_tier = px.bar(
                tier_country_df, x='country_name', y='cnt', color='data_tier',
                title='Top 20 Countries by Data Tier / 数据量前20国家的密级分布',
                labels={'country_name': '经济体 / Economy', 'cnt': '公司数 / Companies', 'data_tier': '密级 / Tier'},
                height=450,
                color_discrete_map={'Tier1': '#28a745', 'Tier2': '#ffc107', 'Tier3': '#fd7e14'},
            )
            st.plotly_chart(fig_tier, use_container_width=True)
        else:
            st.info("暂无数据 / No data available")

    # ========== Tab 4: Statistics ==========
    with contrib_tabs[3]:
        st.subheader("📊 贡献统计 / Contribution Statistics")

        stats_col1, stats_col2, stats_col3, stats_col4 = st.columns(4)
        with stats_col1:
            st.metric("贡献国家 / Contributing Countries", "28")
        with stats_col2:
            st.metric("总数据集 / Total Datasets", "1,247")
        with stats_col3:
            st.metric("待审核 / Pending Review", "23")
        with stats_col4:
            st.metric("本月批准 / Approved This Month", "18")

        # Contribution trend chart
        months = ["Jan", "Feb", "Mar", "Apr", "May", "Jun"]
        uploads = [15, 18, 22, 20, 25, 18]
        approvals = [12, 16, 20, 19, 22, 16]

        fig_trend = go.Figure()
        fig_trend.add_trace(go.Bar(
            x=months, y=uploads, name="上传 / Uploads", marker_color=PRIMARY_COLOR,
        ))
        fig_trend.add_trace(go.Bar(
            x=months, y=approvals, name="批准 / Approved", marker_color=ACCENT_GREEN,
        ))
        fig_trend.update_layout(
            barmode="group",
            title="月度数据贡献趋势 (2025) / Monthly Contribution Trend",
            xaxis_title="月份 / Month",
            yaxis_title="数据集数量 / Datasets",
            height=350,
        )
        st.plotly_chart(fig_trend, use_container_width=True)

        # Top contributors
        st.markdown("#### 🏆 贡献最多经济体 (2025) / Top Contributing Countries")
        contributors = pd.DataFrame({
            "经济体 / Economy": ["India", "South Africa", "Brazil", "China", "Kenya", "Nigeria", "Mexico", "Indonesia"],
            "数据集 / Datasets": [180, 145, 132, 128, 95, 88, 82, 75],
            "质量评分 / Quality Score": [95, 92, 90, 94, 88, 85, 91, 87],
            "Tier1占比 / Tier1 %": ["62%", "55%", "48%", "70%", "35%", "28%", "52%", "40%"],
        })
        st.dataframe(contributors, use_container_width=True, hide_index=True)


# ============== Page 6: MAP/APA Case Tracker ==============

def map_apa_tracker() -> None:
    """Page 6: MAP/APA Case Tracker - Monitor mutual agreement and advance pricing agreement cases."""
    st.markdown("<div class='main-header'>⚖️ 争议解决追踪器 / MAP & APA Tracker</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='sub-header'>追踪各经济体相互协商程序和预约定价安排 / Track MAP and APA Cases Across Jurisdictions</div>",
        unsafe_allow_html=True,
    )

    st.markdown("""
    <div class='info-bar'>
    💡 <b>本页功能说明 / Page Guide:</b><br>
    <b>中文：</b>追踪跨国税务争议解决案件的「监控台」。<br>
    &nbsp;&nbsp;• <b>MAP案例</b>：相互协商程序 — 两国税务部门协商解决双重征税<br>
    &nbsp;&nbsp;• <b>APA案例</b>：预约定价安排 — 事前约定关联交易定价方法<br>
    &nbsp;&nbsp;• <b>仲裁案例</b>：通过独立仲裁解决税务争议<br>
    <b>English:</b> Monitor cross-border tax dispute resolution cases.<br>
    &nbsp;&nbsp;• <b>MAP</b>: Mutual Agreement Procedure — resolving double taxation through bilateral negotiation<br>
    &nbsp;&nbsp;• <b>APA</b>: Advance Pricing Agreement — pre-agreeing TP methods with tax authorities<br>
    &nbsp;&nbsp;• <b>Arbitration</b>: Resolving disputes through independent arbitration<br>
    <br>
    <i>依据 / Reference: 联合国国际税收合作框架公约第13条（争议预防与解决）及联合国税收协定范本第25条<br>
    UN Framework Convention Art. 13 & UN Model Tax Convention Art. 25 (Mutual Agreement Procedure)<br>
    OECD BEPS Action 14 (Making Dispute Resolution More Effective) — 最低标准：MAP案件应在2年内解决，确保纳税人可有效进入MAP程序</i>
    </div>
    """, unsafe_allow_html=True)

    if st.button("❓ 帮助 / Help", key="help_map", help="使用指南 / Guide"):
        st.info("""
        **争议解决追踪器 — 使用指南 / MAP Tracker Guide**

        **中文：** 本页追踪各经济体相互协商程序(MAP)和预约定价安排(APA)案件。
        - **MAP**：两国税务部门协商解决双重征税争议
        - **APA**：事前与税务机关约定转让定价方法，预防争议
        - 颜色含义：🟡进行中 / 🟢已解决 / 🔵待处理 / 🔴未达成协议

        **English:** This tracker monitors MAP and APA cases across jurisdictions.
        """)

    cases_df = load_map_cases()

    # Filters
    fcol1, fcol2, fcol3 = st.columns(3)
    with fcol1:
        status_filter = st.multiselect("状态 / Status", cases_df["status"].unique().tolist(), default=[])
    with fcol2:
        type_filter = st.multiselect("案件类型 / Case Type", cases_df["case_type"].unique().tolist(), default=[])
    with fcol3:
        issue_filter = st.multiselect("争议事项 / Issue", cases_df["issue"].unique().tolist(), default=[])

    filtered = cases_df.copy()
    if status_filter:
        filtered = filtered[filtered["status"].isin(status_filter)]
    if type_filter:
        filtered = filtered[filtered["case_type"].isin(type_filter)]
    if issue_filter:
        filtered = filtered[filtered["issue"].isin(issue_filter)]

    # Statistics
    st.markdown("---")
    st.subheader("案件统计 / Case Statistics")

    total_cases = len(filtered)
    active_cases = len(filtered[filtered["status"] == "Active"])
    resolved_cases = len(filtered[filtered["status"] == "Resolved"])
    avg_duration = filtered["duration_days"].mean() if len(filtered) > 0 else 0
    total_disputed = filtered["amount_disputed_musd"].sum() if len(filtered) > 0 else 0

    stat_cols = st.columns(5)
    with stat_cols[0]:
        st.metric("总案件数 / Total", total_cases)
    with stat_cols[1]:
        st.metric("进行中 / Active", active_cases)
    with stat_cols[2]:
        st.metric("已解决 / Resolved", resolved_cases)
    with stat_cols[3]:
        st.metric("平均时长 / Avg Duration", f"{avg_duration:.0f} 天/days")
        if avg_duration > 0:
            st.caption(f"对标OECD MAP Statistics 2024: 全球平均约30个月(900天) / Benchmark: OECD MAP Stats ~30 months")
    with stat_cols[4]:
        st.metric("争议金额 / Disputed", f"${total_disputed:.1f}M")

    # Cases table
    st.markdown("---")
    st.subheader("案件列表 / Case List")

    if len(filtered) > 0:
        display_df = filtered[["case_id", "case_type", "taxpayer", "country_a", "country_b",
                               "issue", "status", "start_date", "duration_days", "amount_disputed_musd"]].copy()
        display_df.columns = ["案件号/ID", "类型/Type", "纳税人/Taxpayer", "经济体A/Economy A", "经济体B/Economy B",
                              "争议事项/Issue", "状态/Status", "起始日/Start", "时长(天)/Duration", "争议额($M)/Disputed"]

        st.dataframe(
            display_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Status": st.column_config.TextColumn("Status", help="Case status"),
            },
        )
    else:
        st.info("没有符合筛选条件的案件。/ No cases match the selected filters.")

    # Trend analysis
    st.markdown("---")
    st.subheader("趋势分析 / Trend Analysis")

    trend_col1, trend_col2 = st.columns(2)

    with trend_col1:
        # Status distribution
        status_counts = cases_df["status"].value_counts().reset_index()
        status_counts.columns = ["Status", "Count"]
        fig_pie = px.pie(
            status_counts, values="Count", names="Status",
            title="案件状态分布 / Case Status Distribution",
            color="Status",
            color_discrete_map={
                "Active": ACCENT_YELLOW,
                "Resolved": ACCENT_GREEN,
                "Pending": PRIMARY_COLOR,
                "Closed - No agreement": ACCENT_RED,
            },
            height=350,
        )
        st.plotly_chart(fig_pie, use_container_width=True)

    with trend_col2:
        # Duration by case type
        fig_box = px.box(
            cases_df, x="case_type", y="duration_days", color="status",
            title="各类型案件时长 / Case Duration by Type",
            labels={"duration_days": "时长(天) / Duration (days)", "case_type": "案件类型 / Case Type"},
            height=350,
            color_discrete_map={
                "Active": ACCENT_YELLOW,
                "Resolved": ACCENT_GREEN,
                "Pending": PRIMARY_COLOR,
                "Closed - No agreement": ACCENT_RED,
            },
        )
        st.plotly_chart(fig_box, use_container_width=True)

    # Efficiency metrics
    st.markdown("---")
    st.subheader("效率指标 / Efficiency Indicators")

    eff_col1, eff_col2, eff_col3 = st.columns(3)
    with eff_col1:
        resolution_rate = len(cases_df[cases_df["status"] == "Resolved"]) / max(len(cases_df), 1) * 100
        fig_gauge = go.Figure(go.Indicator(
            mode="gauge+number",
            value=round(resolution_rate, 1),
            title={"text": "解决率 / Resolution Rate (%)"},
            gauge={
                "axis": {"range": [0, 100]},
                "bar": {"color": ACCENT_GREEN},
                "steps": [
                    {"range": [0, 30], "color": "#ffcccc"},
                    {"range": [30, 70], "color": "#ffffcc"},
                    {"range": [70, 100], "color": "#ccffcc"},
                ],
                "threshold": {
                    "line": {"color": "red", "width": 4},
                    "thickness": 0.75,
                    "value": 50,
                },
            },
        ))
        fig_gauge.update_layout(height=250)
        st.plotly_chart(fig_gauge, use_container_width=True)

    with eff_col2:
        # MAP timeline trend (mock)
        quarters = ["Q1 2023", "Q2 2023", "Q3 2023", "Q4 2023", "Q1 2024", "Q2 2024", "Q3 2024", "Q4 2024"]
        new_cases = [8, 10, 7, 12, 9, 11, 8, 13]
        resolved_cases_trend = [5, 7, 8, 9, 7, 10, 9, 11]

        fig_line = go.Figure()
        fig_line.add_trace(go.Scatter(
            x=quarters, y=new_cases, mode="lines+markers",
            name="新增案件 / New Cases", line=dict(color=ACCENT_RED),
        ))
        fig_line.add_trace(go.Scatter(
            x=quarters, y=resolved_cases_trend, mode="lines+markers",
            name="已解决 / Resolved", line=dict(color=ACCENT_GREEN),
        ))
        fig_line.update_layout(
            title="MAP案件趋势 / MAP Case Trend",
            xaxis_title="季度 / Quarter",
            yaxis_title="案件数 / Cases",
            height=280,
            legend=dict(orientation="h", yanchor="bottom", y=-0.3, xanchor="center", x=0.5),
        )
        st.plotly_chart(fig_line, use_container_width=True)

    with eff_col3:
        # Average duration by issue type
        issue_duration = cases_df.groupby("issue")["duration_days"].mean().sort_values(ascending=True).reset_index()
        issue_duration.columns = ["Issue", "Avg Duration"]
        fig_bar = px.bar(
            issue_duration, y="Issue", x="Avg Duration",
            orientation="h", title="各争议类型平均时长 / Avg. Duration by Issue Type",
            color="Avg Duration",
            color_continuous_scale="RdYlGn_r",
            height=280,
        )
        st.plotly_chart(fig_bar, use_container_width=True)


# ============== Page 7: Digital Economy Dashboard ==============

def digital_economy() -> None:
    """Page 7: Digital Economy & Cross-Border Services Dashboard."""
    st.markdown("<div class='main-header'>💻 数字经济与跨境服务 / Digital Economy & Cross-Border Services</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='sub-header'>分析数字服务交易和跨境服务定价 / Digital Service Transactions & Cross-Border Pricing Analysis</div>",
        unsafe_allow_html=True,
    )

    st.markdown("""
    <div class='info-bar'>
    💡 <b>本页功能说明 / Page Guide:</b><br>
    <b>中文：</b>分析数字经济领域转让定价的「专业工具」。<br>
    &nbsp;&nbsp;• <b>数字经济可比数据筛选</b>：按国家、利润率、财政年度筛选数字服务公司<br>
    &nbsp;&nbsp;• <b>收入vs利润率分析</b>：气泡图展示不同平台类型的定价特征<br>
    &nbsp;&nbsp;• <b>数字强度指标</b>：衡量企业数字化程度的综合指标<br>
    <b>English:</b> Specialized tools for digital economy transfer pricing analysis.<br>
    &nbsp;&nbsp;• <b>Data Filtering</b>: Screen digital service companies by economy, margin, fiscal year<br>
    &nbsp;&nbsp;• <b>Revenue vs Margin</b>: Bubble chart showing pricing characteristics by platform type<br>
    &nbsp;&nbsp;• <b>Digital Intensity</b>: Composite metric for digitalization level<br>
    <br>
    <i>依据 / Reference: 联合国国际税收合作框架公约第5条及联合国数字经济税收挑战工作组报告<br>
    UN Framework Convention Art. 5 & UN Committee of Experts on International Cooperation in Tax Matters: Digital Economy Report</i>
    </div>
    """, unsafe_allow_html=True)

    if st.button("Help", key="help_digital", help="Guidance on digital economy analysis"):
        st.info("""
        **Digital Economy Dashboard Help**

        This dashboard provides tools for analyzing digital economy transactions:
        - **Service Type Selection**: Choose from cloud computing, SaaS, digital advertising, etc.
        - **Comparable Data Filter**: Screen digital economy comparables by region, size, and margins
        - **Cross-Border Pricing Analysis**: Analyze profitability indicators for cross-border services

        Workstream II of the Framework Convention addresses the tax challenges of the digitalized economy.
        """)

    # Service type selection
    st.markdown("---")
    st.subheader("数字服务类型选择 / Service Type Selection")

    service_type_en = st.selectbox(
        "选择数字服务类型 / Select Digital Service Type",
        ["All", "Cloud Computing", "SaaS", "Digital Advertising", "E-commerce Platform",
         "Data Analytics", "IT Outsourcing", "Digital Content", "Fintech Services"],
        format_func=lambda x: {
            "All": "全部 / All",
            "Cloud Computing": "云计算 / Cloud Computing",
            "SaaS": "软件即服务 / SaaS",
            "Digital Advertising": "数字广告 / Digital Advertising",
            "E-commerce Platform": "电商平台 / E-commerce",
            "Data Analytics": "数据分析 / Data Analytics",
            "IT Outsourcing": "IT外包 / IT Outsourcing",
            "Digital Content": "数字内容 / Digital Content",
            "Fintech Services": "金融科技 / Fintech",
        }.get(x, x),
        help="选择要分析的数字服务类型 / Select the type of digital service",
    )
    service_type = service_type_en

    # Comparable data filter
    st.markdown("---")
    st.subheader("数字经济可比数据筛选 / Digital Economy Comparable Data Filter")

    digital_df = load_digital_comparables()
    if service_type != "All":
        digital_df = digital_df[digital_df["service_type"] == service_type]

    # Filters
    fcol1, fcol2, fcol3 = st.columns(3)
    with fcol1:
        dig_countries = st.multiselect(
            "经济体 / Economies",
            sorted(digital_df["country"].unique().tolist()),
            default=[],
            help="Filter by economy of the comparable company",
        )
    with fcol2:
        margin_range = st.slider(
            "Operating Margin Range (%)",
            float(digital_df["operating_margin"].min() * 100),
            float(digital_df["operating_margin"].max() * 100),
            (0.0, 45.0),
            help="Filter by operating margin percentage",
        )
    with fcol3:
        fy_filter = st.multiselect(
            "Fiscal Year",
            sorted(digital_df["fiscal_year"].unique().tolist()),
            default=sorted(digital_df["fiscal_year"].unique().tolist()),
            help="Filter by fiscal year",
        )

    # Apply filters
    filtered_digital = digital_df.copy()
    if dig_countries:
        filtered_digital = filtered_digital[filtered_digital["country"].isin(dig_countries)]
    filtered_digital = filtered_digital[
        (filtered_digital["operating_margin"] * 100 >= margin_range[0]) &
        (filtered_digital["operating_margin"] * 100 <= margin_range[1])
    ]
    if fy_filter:
        filtered_digital = filtered_digital[filtered_digital["fiscal_year"].isin(fy_filter)]

    st.markdown(f"**筛选结果：{len(filtered_digital)} 家公司 / Filtered Results: {len(filtered_digital)} companies**")

    if len(filtered_digital) > 0:
        # Metrics
        mcol1, mcol2, mcol3, mcol4 = st.columns(4)
        with mcol1:
            st.metric("公司数 / Companies", len(filtered_digital))
        with mcol2:
            st.metric("平均利润率 / Avg Margin", f"{filtered_digital['operating_margin'].mean()*100:.2f}%")
        with mcol3:
            st.metric("中位收入 / Median Revenue", f"${filtered_digital['revenue_musd'].median():.1f}M")
        with mcol4:
            st.metric("平均数字强度 / Avg Digital Intensity", f"{filtered_digital['digital_intensity'].mean():.2f}")

        # Data table
        st.dataframe(
            filtered_digital[["company_name", "country", "service_type", "revenue_musd",
                             "operating_margin", "cost_plus_markup", "tnmm_profit_margin",
                             "digital_intensity", "fiscal_year"]],
            use_container_width=True,
            hide_index=True,
        )

        # Scatter plot - Revenue vs Operating Margin
        fig_scatter = px.scatter(
            filtered_digital,
            x="revenue_musd",
            y="operating_margin",
            color="service_type",
            size="digital_intensity",
            hover_data=["company_name", "country"],
            title="收入vs营业利润率（气泡大小=数字强度）/ Revenue vs Operating Margin",
            labels={"revenue_musd": "收入($M) / Revenue", "operating_margin": "营业利润率 / Operating Margin"},
            height=450,
        )
        fig_scatter.update_layout(margin=dict(l=0, r=0, t=40, b=0))
        st.plotly_chart(fig_scatter, use_container_width=True)
        st.info("💡 **图表说明 / Chart Guide:** 每个气泡代表一家公司。横轴是年收入，纵轴是营业利润率，气泡大小表示数字强度（越大越依赖数字业务）。颜色代表平台类型。/ Each bubble = a company. X-axis: annual revenue. Y-axis: operating margin. Bubble size: digital intensity. Color: platform type.")

        # Box plot by service type
        if service_type == "All":
            fig_box = px.box(
                filtered_digital,
                x="service_type",
                y="operating_margin",
                title="各平台类型利润率分布 / Operating Margin by Service Type",
                labels={"operating_margin": "营业利润率 / Operating Margin", "service_type": "平台类型 / Service Type"},
                height=400,
                color="service_type",
            )
            st.plotly_chart(fig_box, use_container_width=True)
            st.info("💡 **图表说明 / Chart Guide:** 箱线图展示每种平台类型的利润率分布。箱体中线=中位数，箱体上下边=四分位区间，圆点=异常值。/ Box plot shows margin distribution by platform type. Middle line: median. Box edges: Q1/Q3. Dots: outliers.")
    else:
        st.warning("没有符合筛选条件的公司，请调整筛选。/ No companies match the selected filters.")

    # Cross-Border Services Analysis
    st.markdown("---")
    st.subheader("跨境服务定价分析 / Cross-Border Services Pricing Analysis")

    cbs_col1, cbs_col2 = st.columns(2)
    with cbs_col1:
        st.markdown("#### TNMM 利润率指标 / TNMM Profitability Indicators")

        if len(filtered_digital) > 0:
            pli_options = ["Operating Margin", "Cost Plus Markup", "TNMM Profit Margin"]
            selected_pli = st.selectbox("利润水平指标 / Profit Level Indicator", pli_options)

            pli_col = {
                "Operating Margin": "operating_margin",
                "Cost Plus Markup": "cost_plus_markup",
                "TNMM Profit Margin": "tnmm_profit_margin",
            }[selected_pli]

            values = filtered_digital[pli_col].dropna() * 100
            fig_hist = px.histogram(
                values, nbins=20,
                title=f"{selected_pli} 分布 / Distribution",
                labels={"value": f"{selected_pli} (%)"},
                height=350,
                color_discrete_sequence=[PRIMARY_COLOR],
            )
            st.plotly_chart(fig_hist, use_container_width=True)

            st.markdown(f"**均值/Mean**: {values.mean():.2f}% | **中位数/Median**: {values.median():.2f}% | **标准差/Std Dev**: {values.std():.2f}%")

    with cbs_col2:
        st.markdown("#### 区域对比 / Regional Comparison")

        if len(filtered_digital) > 0:
            region_summary = filtered_digital.groupby("country").agg(
                companies=("company_name", "count"),
                avg_margin=("operating_margin", "mean"),
                avg_revenue=("revenue_musd", "mean"),
            ).reset_index()
            region_summary = region_summary.sort_values("avg_margin", ascending=True)

            fig_bar = px.bar(
                region_summary,
                y="country",
                x="avg_margin",
                orientation="h",
                title="各经济体平均营业利润率 / Average Operating Margin by Economy",
                labels={"avg_margin": "平均营业利润率 / Avg Operating Margin", "country": "经济体 / Economy"},
                height=350,
                color="avg_margin",
                color_continuous_scale="Blues",
            )
            st.plotly_chart(fig_bar, use_container_width=True)

    # Workstream II Knowledge
    st.markdown("---")
    st.subheader("工作流II：数字经济税收 / Workstream II: Digital Economy Taxation")

    with st.expander("了解联合国框架公约工作流II / Learn about Workstream II"):
        st.markdown("""
        **Workstream II** addresses the tax challenges arising from the digitalization of the economy.
        Key focus areas include:

        1. **Nexus Rules**: Determining when a jurisdiction has taxing rights over digital transactions
        2. **Profit Allocation**: How to allocate profits from digital services across jurisdictions
        3. **Administrative Cooperation**: Enhancing information exchange for digital economy transactions
        4. **Simplified Approaches**: Safe harbors and simplified methods for low-value digital transactions

        **Cross-Border Services Pricing Challenges:**
        - Difficulty in identifying comparable transactions for novel digital services
        - Rapidly evolving business models and value chains
        - Intangibles-heavy business models with limited physical presence
        - User participation and data as value creation factors

        **Recommended Approach:**
        The TNMM (Transactional Net Margin Method) using the operating margin or Berry ratio
        is often the most appropriate method for routine digital service providers.
        For highly integrated digital operations, the Profit Split Method may be more suitable.

        **相关国际框架 / Related International Frameworks:**

        📌 **OECD/G20 BEPS Pillar One (双支柱方案):**
        - **Amount A**: 对全球最大型MNE（收入≥€200亿且利润率>10%）的超额利润（超过10%利润率部分）的25%重新分配至市场管辖区
        - **Amount B**: 对基础营销和分销活动的标准化回报（简化应用独立交易原则）
        - 状态: 2024年签署多边公约(MAA)，部分管辖区已实施

        📌 **UN Model Tax Convention Article 12B (数字服务常设机构):**
        - 引入"显著经济存在"(Significant Economic Presence)概念
        - 允许来源国对数字服务收入征税，无需物理存在
        - 适用于自动数字服务(ADS)和其他数字服务

        📌 **数字服务税 (DST):**
        - 部分国家(英、法、意、印等)已实施数字服务税（通常2-3%）
        - 作为Pillar One实施前的过渡性措施
        """)


# ============== Page 8: Risk Assessment Center ==============

def risk_assessment() -> None:
    """Page 8: Transfer Pricing Risk Assessment Center — UN INC aligned."""
    st.markdown("<div class='main-header'>🚨 转让定价风险评估中心 / TP Risk Assessment Center</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='sub-header'>基于联合国国际税收合作框架公约的智能风险识别与评估</div>",
        unsafe_allow_html=True,
    )

    st.markdown("""
    <div class='info-bar'>
    💡 <b>本页功能说明 / Page Guide:</b><br>
    <b>中文：</b>基于联合国INC会议成果，对跨国企业转让定价风险进行智能评估。<br>
    &nbsp;&nbsp;• <b>风险雷达图</b>：从6个维度评估关联交易风险（利润率偏离、关联交易占比、避税港使用、无形资产迁移、债务杠杆、功能风险匹配）<br>
    &nbsp;&nbsp;• <b>风险等级</b>：🟢低风险 / 🟡中风险 / 🔴高风险 / ⚫极高风险<br>
    &nbsp;&nbsp;• <b>合规检查</b>：对照联合国框架公约各条款自动检查合规性<br>
    &nbsp;&nbsp;• <b>CbC国别报告分析</b>：分析跨国企业全球利润分配与经济活动实质的匹配度<br>
    <b>English:</b> Intelligent TP risk assessment aligned with UN INC outcomes.<br>
    &nbsp;&nbsp;• <b>Risk Radar</b>: 6-dimension risk scoring (profit deviation, related-party ratio, tax haven use, intangibles migration, debt leverage, function-risk match)<br>
    &nbsp;&nbsp;• <b>Risk Levels</b>: 🟢 Low / 🟡 Medium / 🔴 High / ⚫ Critical<br>
    &nbsp;&nbsp;• <b>Compliance Check</b>: Automated checks against UN Framework Convention articles<br>
    &nbsp;&nbsp;• <b>CbC Analysis</b>: Profit allocation vs. economic substance alignment<br>
    <br>
    <i>依据 / Reference: 联合国国际税收合作框架公约第5条、第6条、第13条及联合国INC会议成果文件<br>
    UN Framework Convention Arts. 5, 6, 13 & UN INC Conference Outcomes</i>
    </div>
    """, unsafe_allow_html=True)

    risk_tabs = st.tabs([
        "🎯 风险雷达 / Risk Radar",
        "📋 合规检查 / Compliance Check",
        "📊 CbC分析 / CbC Analysis",
    ])

    # ========== Tab 1: Risk Radar ==========
    with risk_tabs[0]:
        st.subheader("🎯 转让定价风险雷达图 / TP Risk Radar")

        st.markdown("""
        **中文说明：** 输入被评估企业的关键指标，系统从6个维度自动评分并生成风险雷达图。
        评分越高表示风险越大（0=无风险，100=极高风险）。

        **English:** Enter the assessed company's key metrics. The system scores 6 risk dimensions
        and generates a radar chart. Higher scores = higher risk (0=no risk, 100=critical).
        """)

        input_col1, input_col2 = st.columns(2)
        with input_col1:
            st.markdown("##### 企业信息 / Company Info")
            company_name = st.text_input("企业名称 / Company Name", value="MNE Group A", key="ra_company")
            country = st.selectbox("所在经济体 / Jurisdiction", ["China", "India", "Brazil", "South Africa", "Kenya", "Nigeria", "Other"], key="ra_country")
            industry = st.selectbox("行业 / Industry", ["Manufacturing", "Services", "Digital/Tech", "Pharmaceuticals", "Mining", "Financial Services"], key="ra_industry")

        with input_col2:
            st.markdown("##### 关键指标 / Key Metrics")
            operating_margin = st.slider("营业利润率 / Operating Margin (%)", -20.0, 50.0, 3.0, 0.5, key="ra_margin",
                                         help="被测试方的营业利润率，与行业基准对比")
            rp_ratio = st.slider("关联交易占比 / Related Party Ratio (%)", 0, 100, 40, 5, key="ra_rp",
                                 help="关联交易占总收入的比例")
            effective_tax = st.slider("实际税率 / Effective Tax Rate (%)", 0.0, 35.0, 15.0, 0.5, key="ra_tax",
                                      help="实际税率低于法定税率可能暗示利润转移")
            debt_to_equity = st.slider("债务权益比 / Debt-to-Equity Ratio", 0.0, 10.0, 3.0, 0.1, key="ra_de",
                                       help="过高的债务权益比可能暗示资本弱化")

        intangibles_transfer = st.checkbox("是否存在无形资产跨境转移 / Intangibles cross-border transfer", key="ra_intang",
                                           help="无形资产向低税地转移是高风险信号")
        tax_haven_use = st.checkbox("是否在低税管辖区设立子公司 / Subsidiaries in low-tax jurisdictions", key="ra_haven",
                                    help="在零税或极低税管辖区（如BVI、开曼、百慕大等）设立子公司是高风险信号。"
                                    "OECD有害税收实践论坛(FHTP)维护避税地清单。")

        # Industry benchmark margins (approximate median OM by industry)
        industry_benchmarks = {
            "Manufacturing": 8.0,
            "Services": 10.0,
            "Digital/Tech": 15.0,
            "Pharmaceuticals": 20.0,
            "Mining": 12.0,
            "Financial Services": 18.0,
        }
        industry_median_om = industry_benchmarks.get(industry, 10.0)

        # Calculate risk scores — relative to industry benchmark, not fixed 10%
        risk_dims = {
            '利润率偏离\nProfit Deviation': max(0, min(100, (industry_median_om + 5 - abs(operating_margin - industry_median_om)) * (100 / (industry_median_om + 5)))),
            '关联交易占比\nRelated Party': min(100, rp_ratio * 1.5),
            '低税管辖区使用\nTax Haven Use': (80 if tax_haven_use else 20) + (100 - effective_tax * 2),
            '无形资产迁移\nIntangibles Migration': (85 if intangibles_transfer else 15),
            '债务杠杆\nDebt Leverage': min(100, debt_to_equity * 15),
            '功能风险匹配\nFunction-Risk Match': 50,  # Default moderate
        }

        # Adjust for industry
        if industry == "Digital/Tech":
            risk_dims['无形资产迁移\nIntangibles Migration'] += 10
        if industry == "Pharmaceuticals":
            risk_dims['无形资产迁移\nIntangibles Migration'] += 15

        # Clamp
        risk_dims = {k: min(100, max(0, v)) for k, v in risk_dims.items()}

        st.caption(f"💡 利润率偏离基准: **{industry}行业中位OM = {industry_median_om:.1f}%** "
                   f"(不同行业利润率基准不同，本系统使用行业近似中位数)")

        # Overall risk
        overall_risk = sum(risk_dims.values()) / len(risk_dims)
        if overall_risk >= 70:
            risk_level = "⚫ 极高风险 / Critical"
            risk_color = "#000000"
            risk_advice = "强烈建议立即启动转让定价专项审计，并考虑启动相互协商程序(MAP)。"
        elif overall_risk >= 50:
            risk_level = "🔴 高风险 / High"
            risk_color = "#dc3545"
            risk_advice = "建议进行深入转让定价分析，准备完整的转让定价文档，考虑预约定价安排(APA)。"
        elif overall_risk >= 30:
            risk_level = "🟡 中风险 / Medium"
            risk_color = "#ffc107"
            risk_advice = "建议监控风险变化，完善转让定价文档，定期进行自查。"
        else:
            risk_level = "🟢 低风险 / Low"
            risk_color = "#28a745"
            risk_advice = "风险水平正常，保持常规合规监控即可。"

        # Display overall risk — modern gauge card
        st.markdown("---")
        risk_col1, risk_col2 = st.columns([1, 2])
        with risk_col1:
            st.markdown(f"""
            <div style='
                text-align: center;
                padding: 2rem 1.5rem;
                border: 2px solid {risk_color};
                border-radius: 16px;
                background: linear-gradient(135deg, {risk_color}11 0%, {risk_color}05 100%);
                box-shadow: 0 4px 20px {risk_color}22;
            '>
                <div style='font-size: 2rem; margin-bottom: 0.3rem;'>{'🟢' if overall_risk < 30 else '🟡' if overall_risk < 50 else '🔴' if overall_risk < 70 else '⚫'}</div>
                <h2 style='color: {risk_color}; margin: 0; font-size: 1.1rem;'>{risk_level}</h2>
                <div style='
                    font-size: 3rem; font-weight: 900; color: {risk_color};
                    line-height: 1.2; margin: 0.3rem 0;
                '>{overall_risk:.0f}<span style='font-size:1rem;font-weight:500;color:#7F8C8D'>/100</span></div>
                <p style='color: #7F8C8D; font-size:0.8rem; text-transform: uppercase; letter-spacing:1px; margin:0;'>综合风险评分 / Overall Risk Score</p>
            </div>
            """, unsafe_allow_html=True)

        with risk_col2:
            st.markdown("##### 💡 风险评估建议 / Risk Assessment Advice")
            st.markdown(f"<div style='padding: 1rem; border-left: 4px solid {risk_color}; background: #f8f9fa;'><b>中文：</b>{risk_advice}</div>", unsafe_allow_html=True)
            st.markdown(f"""
            <div style='padding: 1rem; border-left: 4px solid {risk_color}; background: #f8f9fa; margin-top: 0.5rem;'>
            <b>English:</b> {
                'Strongly recommend immediate TP audit and consider MAP.' if overall_risk >= 70 else
                'Recommend in-depth TP analysis, prepare full documentation, consider APA.' if overall_risk >= 50 else
                'Monitor risk changes, improve TP documentation, periodic self-review.' if overall_risk >= 30 else
                'Risk level normal. Maintain routine compliance monitoring.'
            }
            </div>
            """, unsafe_allow_html=True)

        # Radar chart
        st.markdown("---")
        fig_radar = go.Figure()
        fig_radar.add_trace(go.Scatterpolar(
            r=list(risk_dims.values()) + [list(risk_dims.values())[0]],
            theta=list(risk_dims.keys()) + [list(risk_dims.keys())[0]],
            fill='toself',
            fillcolor=f'rgba({int(risk_color[1:3], 16)}, {int(risk_color[3:5], 16)}, {int(risk_color[5:7], 16)}, 0.2)',
            line_color=risk_color,
            name='风险评分 / Risk Score',
        ))
        fig_radar.update_layout(
            polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
            title="6维风险雷达图 / 6-Dimension Risk Radar",
            height=500,
            showlegend=False,
        )
        st.plotly_chart(fig_radar, use_container_width=True)

    # ========== Tab 2: Compliance Check ==========
    with risk_tabs[1]:
        st.subheader("📋 联合国框架公约合规检查 / UN Convention Compliance Check")

        st.markdown("""
        **中文说明：** 对照联合国国际税收合作框架公约各核心条款，检查企业转让定价合规情况。
        每项检查都附有条款依据和详细说明。

        **English:** Check TP compliance against each core article of the UN Framework Convention.
        Each check includes the article reference and detailed explanation.
        """)

        compliance_items = [
            {
                'article': '第5条 / Art. 5',
                'title': '公平分配征税权 / Fair Distribution of Taxing Rights',
                'desc': '利润应在经济活动发生地和价值创造地之间合理分配\nProfit should be allocated where economic activity and value creation occur',
                'check': '关联交易定价是否反映了各方实际履行的功能和承担的风险',
                'risk': '高风险' if overall_risk >= 50 else '达标',
            },
            {
                'article': '第6条 / Art. 6',
                'title': '税基侵蚀 / Base Erosion',
                'desc': '防止通过关联交易将利润转移至低税管辖区\nPrevent profit shifting to low-tax jurisdictions through related-party transactions',
                'check': '是否存在向低税率地区的大额支付（特许权使用费、服务费、利息）',
                'risk': '高风险' if tax_haven_use or rp_ratio > 50 else '达标',
            },
            {
                'article': '第11条 / Art. 11',
                'title': '能力建设 / Capacity Building',
                'desc': '发展中国家应加强转让定价管理和审计能力\nDeveloping countries should strengthen TP management and audit capacity',
                'check': '是否保留了完整的转让定价文档（主文件、本地文件、国别报告）',
                'risk': '需完善',
            },
            {
                'article': '第13条 / Art. 13',
                'title': '争议预防与解决 / Dispute Prevention & Resolution',
                'desc': '通过MAP和APA预防和解决跨境税务争议\nPrevent and resolve cross-border tax disputes through MAP and APA',
                'check': '是否建立了争议预防机制（APA、预审沟通等）',
                'risk': '需完善',
            },
        ]

        for item in compliance_items:
            risk_color_val = '#dc3545' if '高风险' in item['risk'] else ('#ffc107' if '需完善' in item['risk'] else '#28a745')
            st.markdown(f"""
            <div style='border: 1px solid #ddd; border-left: 4px solid {risk_color_val};
                        border-radius: 8px; padding: 1rem; margin: 0.5rem 0;'>
                <h4 style='margin: 0; color: {PRIMARY_COLOR};'>{item['article']}: {item['title']}</h4>
                <p style='color: #666; font-size: 0.85rem; margin: 0.5rem 0;'>{item['desc']}</p>
                <p style='margin: 0.3rem 0;'><b>检查项 / Check:</b> {item['check']}</p>
                <span style='background-color: {risk_color_val}; color: white; padding: 0.2rem 0.8rem;
                            border-radius: 12px; font-size: 0.8rem;'>{item['risk']}</span>
            </div>
            """, unsafe_allow_html=True)

    # ========== Tab 3: CbC Analysis ==========
    with risk_tabs[2]:
        st.subheader("📊 国别报告(CbC)利润分配分析 / CbC Profit Allocation Analysis")

        st.markdown("""
        **中文说明：** 模拟分析跨国企业全球利润分配与经济活动实质（员工数、资产、收入）的匹配度。
        如果利润集中在低税地区而经济活动在发展中国家，可能存在利润转移风险。

        **English:** Analyze alignment between MNE global profit allocation and economic substance
        (employees, assets, revenue). Profit concentration in low-tax jurisdictions while economic
        activity occurs in developing economies may indicate profit shifting.
        """)

        # Simulated CbC data
        cbc_data = pd.DataFrame({
            '管辖区 / Jurisdiction': ['中国 / China', '印度 / India', '爱尔兰 / Ireland', '百慕大 / Bermuda', '新加坡 / Singapore', '美国 / USA'],
            '收入占比 / Revenue %': [35, 15, 8, 2, 10, 30],
            '利润占比 / Profit %': [15, 8, 30, 25, 12, 10],
            '员工占比 / Employee %': [40, 20, 3, 0, 5, 32],
            '资产占比 / Asset %': [30, 12, 15, 5, 8, 30],
            '实际税率 / ETR %': [25, 22, 12.5, 0, 17, 21],
        })

        st.markdown("##### 示例CbC数据 / CbC Sample Data")
        st.caption("💡 以下为演示用CbC数据，展示利润分配与经济实质的匹配度分析。实际CbC数据由各管辖区税务机关通过信息交换获取。")
        st.dataframe(cbc_data, use_container_width=True, hide_index=True)

        # Misalignment chart
        st.markdown("---")
        st.markdown("##### 📉 利润与经济实质匹配度 / Profit vs. Economic Substance Alignment")

        fig_cbc = go.Figure()
        jurisdictions = cbc_data['管辖区 / Jurisdiction']
        fig_cbc.add_trace(go.Bar(x=jurisdictions, y=cbc_data['收入占比 / Revenue %'], name='收入占比 / Revenue %', marker_color=PRIMARY_COLOR))
        fig_cbc.add_trace(go.Bar(x=jurisdictions, y=cbc_data['利润占比 / Profit %'], name='利润占比 / Profit %', marker_color=ACCENT_RED))
        fig_cbc.add_trace(go.Bar(x=jurisdictions, y=cbc_data['员工占比 / Employee %'], name='员工占比 / Employee %', marker_color=ACCENT_GREEN))
        fig_cbc.update_layout(
            barmode='group',
            title='收入/利润/员工分布对比 / Revenue vs Profit vs Employee Distribution',
            yaxis_title='占比 (%) / Share (%)',
            height=450,
        )
        st.plotly_chart(fig_cbc, use_container_width=True)

        # Interpretation
        st.markdown("---")
        st.markdown("##### 🤖 自动分析 / Automated Analysis")
        st.warning("""
        ⚠️ **风险信号 / Risk Signals:**

        **中文：**
        • 爱尔兰利润占比(30%)远超收入占比(8%)和员工占比(3%) → 可能存在无形资产相关利润转移
        • 百慕大利润占比(25%)但员工占比(0%)和收入占比(2%) → 典型的避税港利润集中
        • 中国收入占比(35%)和员工占比(40%)但利润占比仅(15%) → 利润可能与经济活动实质不匹配

        **English:**
        • Ireland: Profit share (30%) far exceeds revenue (8%) and employee (3%) → Possible intangibles-related profit shifting
        • Bermuda: Profit share (25%) but 0% employees and 2% revenue → Classic tax haven profit concentration
        • China: Revenue (35%) and employee (40%) but only 15% profit → Profit may not match economic substance

        💡 **建议 / Recommendation:** 依据联合国框架公约第5条，利润应在价值创造地分配。
        建议对爱尔兰和百慕大的利润来源进行深入转让定价分析。
        """)


# ============== Page 9: Industry Benchmark & Trend Analysis ==============

def industry_benchmark() -> None:
    """Page 9: Industry Benchmark & Multi-Year Trend Analysis."""
    st.markdown("<div class='main-header'>📈 行业基准与趋势分析 / Industry Benchmark & Trend Analysis</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='sub-header'>跨行业、跨国家、跨年度的利润率基准对比 — 对标BvD Orbis级别分析</div>",
        unsafe_allow_html=True,
    )

    st.markdown("""
    <div class='info-bar'>
    💡 <b>本页功能说明 / Page Guide:</b><br>
    <b>中文：</b>对标BvD Orbis等专业数据库的行业基准分析功能。<br>
    &nbsp;&nbsp;• <b>行业基准对比</b>：选择行业，查看各国该行业的利润率中位数、四分位区间<br>
    &nbsp;&nbsp;• <b>多年度趋势</b>：2022-2024年度利润率变化趋势，识别行业走势<br>
    &nbsp;&nbsp;• <b>PLI分布</b>：营业利润率(OM)、Berry Ratio、ROCE等标准TP利润水平指标<br>
    &nbsp;&nbsp;• <b>跨国对比</b>：同一行业在不同国家的盈利能力差异<br>
    <b>English:</b> BvD Orbis-level industry benchmark analysis.<br>
    &nbsp;&nbsp;• <b>Industry Benchmark</b>: Select an industry to view median margins, IQR by country<br>
    &nbsp;&nbsp;• <b>Multi-Year Trends</b>: 2022-2024 margin trends to identify industry trajectories<br>
    &nbsp;&nbsp;• <b>PLI Distribution</b>: Operating Margin, Berry Ratio, ROCE — standard TP profit level indicators<br>
    &nbsp;&nbsp;• <b>Cross-Country</b>: Same industry profitability differences across countries<br>
    <br>
    <i>依据 / Reference: 联合国转让定价实务手册第3章（可比性分析）<br>
    UN Practical Manual on Transfer Pricing, Ch. 3 (Comparability Analysis)</i>
    </div>
    """, unsafe_allow_html=True)

    companies_df = load_mock_companies()

    # Filters
    filter_col1, filter_col2, filter_col3 = st.columns(3)
    with filter_col1:
        # Group by ISIC section letter for cleaner categories
        companies_df['isic_section'] = companies_df['industry_code_isic'].astype(str).str[0]
        isic_sections = sorted(companies_df['isic_section'].dropna().unique().tolist())
        section_names = {
            'A': 'A-农林牧渔 / Agriculture',
            'B': 'B-采矿 / Mining',
            'C': 'C-制造 / Manufacturing',
            'D': 'D-电力燃气 / Energy',
            'E': 'E-水务 / Water',
            'F': 'F-建筑 / Construction',
            'G': 'G-批发零售 / Wholesale & Retail',
            'H': 'H-运输仓储 / Transport',
            'I': 'I-住宿餐饮 / Hospitality',
            'J': 'J-信息技术 / IT',
            'K': 'K-金融保险 / Financial',
            'L': 'L-房地产 / Real Estate',
            'M': 'M-专业服务 / Professional',
            'N': 'N-辅助服务 / Support',
        }
        section_labels = [section_names.get(s, f'{s}-其他') for s in isic_sections]
        sel_section_idx = st.selectbox(
            "行业大类 / Industry Sector (ISIC)",
            range(len(isic_sections)),
            format_func=lambda i: section_labels[i] if i < len(section_labels) else 'All',
        )
        sel_section = isic_sections[sel_section_idx]

    with filter_col2:
        # Year selection
        available_years = sorted(companies_df['fiscal_year'].dropna().unique().tolist())
        sel_years = st.multiselect(
            "年度 / Fiscal Years",
            available_years,
            default=available_years,
            help="选择要分析的年度，可多选"
        )

    with filter_col3:
        sel_tier = st.multiselect(
            "数据密级 / Data Tier",
            ['Tier1', 'Tier2', 'Tier3'],
            default=['Tier1', 'Tier2'],
            help="Tier1最可靠，建议优先使用"
        )

    # Apply filters
    filtered = companies_df.copy()
    if sel_section and sel_section != 'All':
        filtered = filtered[filtered['isic_section'] == sel_section]
    if sel_years:
        filtered = filtered[filtered['fiscal_year'].isin(sel_years)]
    if 'data_tier' in filtered.columns and sel_tier:
        filtered = filtered[filtered['data_tier'].isin(sel_tier)]

    if filtered.empty:
        st.warning("⚠️ 无符合条件的数据，请调整筛选。/ No data matches the selected criteria.")
        return

    st.markdown(f"**符合条件：{len(filtered)} 家公司 / {len(filtered['country'].unique())} 个经济体**")

    # ===== Section 1: Industry Benchmark by Economy =====
    st.markdown("---")
    st.subheader("📊 各经济体利润率基准 / Profitability Benchmark by Economy")

    country_benchmark = filtered.groupby('country').agg(
        companies=('company_id', 'count'),
        median_margin=('operating_margin', 'median'),
        q1_margin=('operating_margin', lambda x: np.percentile(x.dropna(), 25)),
        q3_margin=('operating_margin', lambda x: np.percentile(x.dropna(), 75)),
        median_roe=('berry_ratio', 'median'),
    ).round(4).reset_index().sort_values('companies', ascending=False).head(20)

    # Box plot by country
    fig_box = px.box(
        filtered[filtered['country'].isin(country_benchmark['country'].head(15))],
        x='country', y='operating_margin',
        color='country',
        title=f'{section_labels[sel_section_idx]} — 各经济体营业利润率分布 / Operating Margin by Economy',
        labels={'operating_margin': '营业利润率 / Operating Margin', 'country': '经济体 / Economy'},
        height=450,
    )
    fig_box.update_layout(showlegend=False, margin=dict(l=0, r=0, t=40, b=0))
    st.plotly_chart(fig_box, use_container_width=True)

    # Benchmark table with export
    st.markdown("##### 📋 基准数据表 / Benchmark Table")
    display_benchmark = country_benchmark.copy()
    display_benchmark.columns = ['经济体/Economy', '公司数/Companies', '中位利润率/Median OM',
                                 'Q1', 'Q3', '中位Berry Ratio/Median Berry']
    display_benchmark['中位利润率/Median OM'] = (display_benchmark['中位利润率/Median OM'] * 100).round(2).astype(str) + '%'
    display_benchmark['Q1'] = (display_benchmark['Q1'] * 100).round(2).astype(str) + '%'
    display_benchmark['Q3'] = (display_benchmark['Q3'] * 100).round(2).astype(str) + '%'
    display_benchmark['中位Berry Ratio/Median Berry'] = display_benchmark['中位Berry Ratio/Median Berry'].round(3).astype(str)
    st.dataframe(display_benchmark, use_container_width=True, hide_index=True)

    # Export button
    export_col1, export_col2 = st.columns([1, 4])
    with export_col1:
        csv_data = country_benchmark.to_csv(index=False).encode('utf-8-sig')
        st.download_button(
            label="📥 导出CSV / Export CSV",
            data=csv_data,
            file_name=f"industry_benchmark_{sel_section}_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
        )

    # ===== Section 2: Multi-Year Trend =====
    st.markdown("---")
    st.subheader("📅 多年度趋势分析 / Multi-Year Trend Analysis")

    if len(sel_years) >= 2:
        trend_data = filtered.groupby(['fiscal_year', 'isic_section']).agg(
            median_margin=('operating_margin', 'median'),
            q1=('operating_margin', lambda x: np.percentile(x.dropna(), 25)),
            q3=('operating_margin', lambda x: np.percentile(x.dropna(), 75)),
            companies=('company_id', 'count'),
        ).reset_index()

        fig_trend = go.Figure()
        fig_trend.add_trace(go.Scatter(
            x=trend_data['fiscal_year'], y=trend_data['q3'],
            name='Q3 (75th percentile)', mode='lines',
            line=dict(width=0), showlegend=False,
        ))
        fig_trend.add_trace(go.Scatter(
            x=trend_data['fiscal_year'], y=trend_data['q1'],
            name='IQR Range', mode='lines',
            line=dict(width=0), fill='tonexty',
            fillcolor='rgba(0, 86, 149, 0.15)', showlegend=True,
        ))
        fig_trend.add_trace(go.Scatter(
            x=trend_data['fiscal_year'], y=trend_data['median_margin'],
            name='中位数 / Median', mode='lines+markers',
            line=dict(color=PRIMARY_COLOR, width=3),
            marker=dict(size=8),
        ))
        fig_trend.update_layout(
            title=f'{section_labels[sel_section_idx]} — 年度利润率趋势 / Annual Margin Trend',
            xaxis_title='年度 / Fiscal Year',
            yaxis_title='营业利润率 / Operating Margin',
            height=400,
            margin=dict(l=0, r=0, t=40, b=0),
        )
        st.plotly_chart(fig_trend, use_container_width=True)

        # Trend interpretation
        if len(trend_data) >= 2:
            first_median = trend_data['median_margin'].iloc[0]
            last_median = trend_data['median_margin'].iloc[-1]
            change = last_median - first_median
            if change > 0.01:
                trend_msg = f"📈 利润率从 {first_median*100:.1f}% 上升至 {last_median*100:.1f}%，行业盈利能力在改善。/ Margins improving."
            elif change < -0.01:
                trend_msg = f"📉 利润率从 {first_median*100:.1f}% 下降至 {last_median*100:.1f}%，行业盈利能力在下降。/ Margins declining."
            else:
                trend_msg = f"➡️ 利润率保持稳定在 {last_median*100:.1f}% 左右。/ Margins stable."
            st.info(trend_msg)
    else:
        st.info("💡 选择2个或以上年度以查看趋势。/ Select 2+ years to view trends.")

    # ===== Section 3: PLI Distribution =====
    st.markdown("---")
    st.subheader("📊 PLI分布密度图 / PLI Distribution Density")

    pli_col1, pli_col2, pli_col3 = st.columns(3)
    with pli_col1:
        st.markdown("##### 营业利润率 / Operating Margin (OM)")
        fig_m = px.histogram(filtered, x='operating_margin', nbins=30, marginal='box',
                             color_discrete_sequence=[PRIMARY_COLOR], height=300)
        fig_m.update_layout(margin=dict(l=0, r=0, t=10, b=0), showlegend=False)
        st.plotly_chart(fig_m, use_container_width=True)

    with pli_col2:
        st.markdown("##### Berry Ratio (毛利/营业费用率)")
        if 'berry_ratio' in filtered.columns and filtered['berry_ratio'].notna().any():
            fig_r = px.histogram(filtered.dropna(subset=['berry_ratio']), x='berry_ratio', nbins=30, marginal='box',
                                color_discrete_sequence=[ACCENT_GREEN], height=300)
            fig_r.update_layout(margin=dict(l=0, r=0, t=10, b=0), showlegend=False)
            st.plotly_chart(fig_r, use_container_width=True)
            st.caption("💡 Berry Ratio = 毛利润 / 营业费用(SG&A) (OECD TPG Ch.6.28-6.33)\n"
                       "适用于分销商/服务提供商：>1.0表示有毛利润覆盖营业费用，通常1.05-1.20为正常区间")
        else:
            st.info("Berry Ratio数据不可用 / Berry Ratio data not available")

    with pli_col3:
        st.markdown("##### ROCE (已动用资本回报率)")
        if 'roce' in filtered.columns and filtered['roce'].notna().any():
            fig_a = px.histogram(filtered.dropna(subset=['roce']), x='roce', nbins=30, marginal='box',
                                color_discrete_sequence=[ACCENT_YELLOW], height=300)
            fig_a.update_layout(margin=dict(l=0, r=0, t=10, b=0), showlegend=False)
            st.plotly_chart(fig_a, use_container_width=True)
            st.caption("💡 ROCE适用于资本密集型行业：衡量资本使用效率，通常5%-25%为正常区间")
        else:
            st.info("ROCE数据不可用 / ROCE data not available")

    # ===== Section 4: Company Listing with Export =====
    st.markdown("---")
    st.subheader("📋 公司明细 / Company Listing")
    display_cols = ['company_name', 'country', 'industry_code_isic', 'revenue_musd',
                    'operating_margin', 'berry_ratio', 'roce', 'related_party_pct',
                    'fiscal_year', 'data_tier']
    available_cols = [c for c in display_cols if c in filtered.columns]
    st.dataframe(filtered[available_cols], use_container_width=True, hide_index=True)

    export_col1, export_col2 = st.columns([1, 4])
    with export_col1:
        csv_full = filtered[available_cols].to_csv(index=False).encode('utf-8-sig')
        st.download_button(
            label="📥 导出全部CSV / Export All",
            data=csv_full,
            file_name=f"company_listing_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",
        )


# ============== Helper: Country Detail Panel ==============

def render_country_detail(country_name: str, country_code: str) -> None:
    """Render a country detail panel with companies, policies, and disputes."""
    st.markdown(f"### 🌍 {country_name} / {country_code}")

    companies_df = load_mock_companies()
    # Match by country name — also check if "China" should match "Taiwan (China)" entries
    if country_name == "China":
        country_data = companies_df[companies_df['country'].isin(["China", "Taiwan (China)"])]
    else:
        country_data = companies_df[companies_df['country'] == country_name]

    detail_col1, detail_col2, detail_col3, detail_col4 = st.columns(4)
    with detail_col1:
        st.metric("可比公司 / Companies", len(country_data))
    with detail_col2:
        st.metric("中位利润率 / Median Margin",
                  f"{country_data['operating_margin'].median()*100:.1f}%" if not country_data.empty else "N/A")
    with detail_col3:
        st.metric("中位收入 / Median Revenue (M$)",
                  f"${country_data['revenue_musd'].median():.1f}" if not country_data.empty else "N/A")
    with detail_col4:
        if 'data_tier' in country_data.columns:
            tier_counts = country_data['data_tier'].value_counts()
            st.metric("主要密级 / Primary Tier", tier_counts.index[0] if not tier_counts.empty else "N/A")
        else:
            st.metric("密级 / Tier", "N/A")

    if not country_data.empty:
        st.markdown("##### 公司列表 / Company List")
        display_cols = ['company_name', 'industry_code_isic', 'revenue_musd',
                        'operating_margin', 'berry_ratio', 'fiscal_year']
        available = [c for c in display_cols if c in country_data.columns]
        st.dataframe(country_data[available].head(20), use_container_width=True, hide_index=True)
    else:
        st.info(f"暂无 {country_name} 的可比公司数据。/ No comparable company data for {country_name}.")


# ============== Helper: Global Search ==============

def global_search(query: str) -> list:
    """Search across companies, countries, and industries. Returns list of (type, label, detail) tuples."""
    results = []
    if not query or len(query.strip()) < 2:
        return results

    companies_df = load_mock_companies()
    countries_df = load_countries_data_v2()

    # Search companies
    company_matches = companies_df[
        companies_df['company_name'].str.contains(query, case=False, na=False)
    ].head(10)
    for _, row in company_matches.iterrows():
        results.append(('company', row['company_name'],
                         f"{row.get('country', '')} | {row.get('industry_code_isic', '')} | FY{row.get('fiscal_year', '')}"))

    # Search countries — also search the full countries dataset
    country_matches = companies_df[
        companies_df['country'].str.contains(query, case=False, na=False)
    ]['country'].unique()[:10]
    for c in country_matches:
        count = len(companies_df[companies_df['country'] == c])
        results.append(('country', c, f"{count} 家公司 / {count} companies"))

    # Also search the full UN economies list (even those without companies)
    if 'country' in countries_df.columns:
        econ_matches = countries_df[
            countries_df['country'].str.contains(query, case=False, na=False)
        ].head(10)
        for _, row in econ_matches.iterrows():
            c = row['country']
            if not any(r[1] == c for r in results):  # Avoid duplicates
                comp_count = int(row.get('comparables_count', 0))
                results.append(('country', c, f"{comp_count} 家可比 / {comp_count} comparables"))

    # Search industries
    if 'industry_code_isic' in companies_df.columns:
        industry_matches = companies_df[
            companies_df['industry_code_isic'].astype(str).str.contains(query, case=False, na=False)
        ]['industry_code_isic'].unique()[:5]
        for ind in industry_matches:
            count = len(companies_df[companies_df['industry_code_isic'] == ind])
            results.append(('industry', f"ISIC {ind}", f"{count} 家公司 / {count} companies"))

    return results


# ============== Page 10: Customs Valuation Reference ==============

def customs_valuation() -> None:
    """Page 10: Customs Valuation Data as TP auxiliary reference."""
    st.markdown("<div class='main-header'>🛃 海关估价参考 / Customs Valuation Reference</div>", unsafe_allow_html=True)
    st.markdown(
        "<div class='sub-header'>海关申报数据辅助转让定价分析 — WTO海关估价协定 × 联合国转让定价实务手册</div>",
        unsafe_allow_html=True,
    )

    st.markdown("""
    <div class='info-bar'>
    💡 <b>本页功能说明 / Page Guide:</b><br>
    <b>中文：</b>海关估价数据是转让定价分析的「重要辅助参考」。根据WTO海关估价协定和联合国转让定价实务手册，海关进口申报价格可作为可比非受控价格(CUP)的参考依据。<br>
    &nbsp;&nbsp;• <b>海关vs转让定价对比</b>：将关联交易的海关申报价与可比价格对比，识别异常定价<br>
    &nbsp;&nbsp;• <b>关联交易筛查</b>：标记海关系统中申报的关联进出口交易<br>
    &nbsp;&nbsp;• <b>价格偏离预警</b>：关联交易价格与独立交易价格偏离>15%自动预警<br>
    &nbsp;&nbsp;• <b>HS编码查询</b>：按商品编码查看进出口价格分布<br>
    <b>English:</b> Customs valuation data serves as an important auxiliary reference for TP analysis. Per WTO Customs Valuation Agreement and UN Practical Manual, customs declared prices can support CUP method.<br>
    &nbsp;&nbsp;• <b>Customs vs TP Comparison</b>: Compare related-party customs prices with comparable prices<br>
    &nbsp;&nbsp;• <b>Related-Party Screening</b>: Flag related-party import/export transactions<br>
    &nbsp;&nbsp;• <b>Price Deviation Alert</b>: Auto-flag >15% price deviations from arm's length<br>
    &nbsp;&nbsp;• <b>HS Code Lookup</b>: View price distribution by HS commodity code<br>
    <br>
    <i>依据 / Reference: WTO海关估价协定 / WTO Customs Valuation Agreement; 联合国转让定价实务手册第4章 / UN Practical Manual Ch.4; OECD转让定价指南第2章（海关与转让定价协同）/ OECD TP Guidelines Ch.2</i>
    </div>
    """, unsafe_allow_html=True)

    customs_df = db_query("""
        SELECT cv.*, 
               rc.country_name AS reporting_country_name,
               pc.country_name AS partner_country_name
        FROM customs_valuations cv
        LEFT JOIN countries rc ON rc.country_code = cv.reporting_country
        LEFT JOIN countries pc ON pc.country_code = cv.partner_country
        ORDER BY cv.valuation_date DESC
    """)

    if customs_df is None or customs_df.empty:
        # Fallback: use comprehensive UN member states data
        from un_countries_data import generate_customs_data
        customs_df = generate_customs_data(500)

    customs_tabs = st.tabs([
        "🔍 海关数据查询 / Data Search",
        "⚖️ 关联交易对比 / Related-Party Comparison",
        "📊 统计分析 / Statistics",
        "📖 使用指引 / Usage Guide",
    ])

    # ===== Tab 1: Data Search =====
    with customs_tabs[0]:
        st.subheader("🔍 海关数据查询 / Customs Data Search")

        filter_c1, filter_c2, filter_c3, filter_c4 = st.columns(4)
        with filter_c1:
            sel_direction = st.selectbox(
                "贸易方向 / Trade Direction",
                ["全部 / All"] + customs_df["trade_direction"].unique().tolist(),
                key="customs_dir"
            )
        with filter_c2:
            sel_category = st.selectbox(
                "商品类别 / Category",
                ["全部 / All"] + sorted(customs_df["commodity_category"].unique().tolist()),
                key="customs_cat"
            )
        with filter_c3:
            sel_related = st.selectbox(
                "关联交易 / Related Party",
                ["全部 / All", "仅关联 / Related only", "仅独立 / Unrelated only"],
                key="customs_rel"
            )
        with filter_c4:
            sel_year = st.multiselect(
                "年度 / Fiscal Year",
                sorted(customs_df["fiscal_year"].unique().tolist()),
                default=sorted(customs_df["fiscal_year"].unique().tolist()),
                key="customs_year"
            )

        filtered = customs_df.copy()
        if sel_direction != "全部 / All":
            filtered = filtered[filtered["trade_direction"] == sel_direction]
        if sel_category != "全部 / All":
            filtered = filtered[filtered["commodity_category"] == sel_category]
        if sel_related == "仅关联 / Related only":
            filtered = filtered[filtered["related_party_flag"] == True]
        elif sel_related == "仅独立 / Unrelated only":
            filtered = filtered[filtered["related_party_flag"] == False]
        if sel_year:
            filtered = filtered[filtered["fiscal_year"].isin(sel_year)]

        st.markdown(f"**筛选结果：{len(filtered)} 条记录 / {len(filtered)} records**")

        if len(filtered) > 0:
            display_cols = ["customs_id", "trade_direction", "reporting_country_name",
                           "partner_country_name", "hs_code", "product_description",
                           "declared_value_usd", "unit_price_usd", "related_party_flag",
                           "price_difference_pct", "verification_status", "valuation_date"]
            available = [c for c in display_cols if c in filtered.columns]
            display = filtered[available].copy()
            display.columns = [c.replace("_name", "").replace("_usd", " ($)") for c in available]
            st.dataframe(display, use_container_width=True, hide_index=True)

            # Export
            exp_c1, _ = st.columns([1, 4])
            with exp_c1:
                csv_customs = filtered.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    label="📥 导出CSV / Export",
                    data=csv_customs,
                    file_name=f"customs_data_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv",
                )

    # ===== Tab 2: Related-Party Comparison =====
    with customs_tabs[1]:
        st.subheader("⚖️ 关联交易海关vs转让定价对比 / Related-Party Customs vs TP Comparison")

        related_df = customs_df[customs_df["related_party_flag"] == True].copy()

        if related_df.empty:
            st.info("暂无关联交易数据。/ No related-party transactions found.")
        else:
            st.markdown(f"**关联交易记录：{len(related_df)} 条 / {len(related_df)} related-party records**")

            # Price deviation chart
            if "price_difference_pct" in related_df.columns:
                related_df["deviation_abs"] = related_df["price_difference_pct"].abs()

                fig_dev = px.scatter(
                    related_df.dropna(subset=["price_difference_pct"]),
                    x="declared_value_usd",
                    y="price_difference_pct",
                    color="commodity_category",
                    hover_data=["hs_code", "product_description", "reporting_country_name"],
                    title="关联交易价格偏离度散点图 / Related-Party Price Deviation",
                    labels={
                        "declared_value_usd": "申报价值($) / Declared Value",
                        "price_difference_pct": "价格偏离(%) / Price Deviation %",
                        "commodity_category": "类别 / Category",
                    },
                    height=450,
                )
                # Add ±15% threshold lines
                fig_dev.add_hline(y=15, line_dash="dash", line_color="red", annotation_text="+15% 预警线")
                fig_dev.add_hline(y=-15, line_dash="dash", line_color="red", annotation_text="-15% 预警线")
                st.plotly_chart(fig_dev, use_container_width=True)
                st.info("💡 **图表说明 / Chart Guide:** 每个点是一条关联交易。纵轴是海关申报价与转让定价可比价的偏离百分比，红色虚线为±15%预警线，超出预警线的交易需要进一步审查。/ Each dot = a related-party transaction. Y-axis: price deviation %. Red dashed lines: ±15% alert threshold.")

            # Flagged transactions
            st.markdown("---")
            st.markdown("##### 🚨 价格偏离预警 / Price Deviation Alerts")
            flagged = related_df[related_df["price_adjustment_flag"] == True]
            st.metric("需调整的交易数 / Flagged Transactions", len(flagged))
            if not flagged.empty:
                flag_cols = ["customs_id", "product_description", "reporting_country_name",
                            "partner_country_name", "declared_value_usd", "price_difference_pct",
                            "adjustment_reason"]
                available_flag = [c for c in flag_cols if c in flagged.columns]
                st.dataframe(flagged[available_flag], use_container_width=True, hide_index=True)

    # ===== Tab 3: Statistics =====
    with customs_tabs[2]:
        st.subheader("📊 海关数据统计分析 / Customs Statistics")

        stat_c1, stat_c2, stat_c3, stat_c4 = st.columns(4)
        with stat_c1:
            st.metric("总记录数 / Total Records", len(customs_df))
        with stat_c2:
            st.metric("关联交易 / Related Party", len(customs_df[customs_df["related_party_flag"] == True]))
        with stat_c3:
            flagged_count = len(customs_df[customs_df["price_adjustment_flag"] == True])
            st.metric("预警交易 / Flagged", flagged_count)
        with stat_c4:
            avg_dev = customs_df["price_difference_pct"].dropna().mean() if "price_difference_pct" in customs_df.columns else 0
            st.metric("平均偏离度 / Avg Deviation", f"{avg_dev:.1f}%" if avg_dev else "N/A")

        # Category distribution
        st.markdown("---")
        st.markdown("##### 📦 商品类别分布 / Category Distribution")
        cat_counts = customs_df.groupby("commodity_category").agg(
            records=("customs_id", "count"),
            related=("related_party_flag", "sum"),
            avg_value=("declared_value_usd", "mean"),
        ).reset_index().sort_values("records", ascending=False)
        fig_cat = px.bar(
            cat_counts, x="commodity_category", y="records",
            color="related", title="各类别海关记录数 / Records by Category",
            labels={"commodity_category": "类别 / Category", "records": "记录数 / Records", "related": "关联 / Related"},
            height=400,
        )
        st.plotly_chart(fig_cat, use_container_width=True)

        # Country distribution
        st.markdown("---")
        st.markdown("##### 🌍 申报经济体分布 / Reporting Economy Distribution")
        country_counts = customs_df.groupby("reporting_country_name").agg(
            records=("customs_id", "count"),
            related=("related_party_flag", "sum"),
        ).reset_index().sort_values("records", ascending=False).head(15)
        fig_country = px.bar(
            country_counts, x="reporting_country_name", y="records",
            color="related", title="Top 15 申报经济体 / Top 15 Reporting Economies",
            labels={"reporting_country_name": "经济体 / Economy", "records": "记录数 / Records", "related": "关联 / Related"},
            height=400,
        )
        st.plotly_chart(fig_country, use_container_width=True)

    # ===== Tab 4: Usage Guide =====
    with customs_tabs[3]:
        st.subheader("📖 海关估价在转让定价中的使用指引 / Usage Guide")

        st.markdown("""
        ### 🔗 海关估价与转让定价的关系
        ### Customs Valuation & Transfer Pricing: The Link

        **中文：**

        海关估价和转让定价是同一问题的两个视角：
        - **海关估价**关注进口货物的申报价格是否合理（征收关税的依据）
        - **转让定价**关注关联企业之间的交易价格是否符合独立交易原则（征收所得税的依据）

        两者都遵循「独立交易原则」(Arm's Length Principle)，但适用场景不同：
        - 海关：进口时一次性估价
        - 转让定价：年度整体利润水平分析

        **如何使用海关数据辅助转让定价分析：**

        1. **CUP法参考**：如果海关有同类商品的独立交易进口价格，可作为CUP法的可比价格
        2. **异常筛查**：如果关联进口价格显著低于独立进口价格，可能存在利润转移
        3. **事后验证**：对已完成的海关申报进行转让定价审查
        4. **风险预警**：海关价格偏离>15%的交易应重点审查

        ---

        **English:**

        Customs valuation and transfer pricing are two perspectives on the same issue:
        - **Customs Valuation**: Whether the declared import price is reasonable (basis for tariffs)
        - **Transfer Pricing**: Whether related-party prices meet the arm's length principle (basis for income tax)

        Both follow the **Arm's Length Principle**, but apply in different contexts:
        - Customs: One-time valuation at import
        - Transfer Pricing: Annual overall profit analysis

        **How to use customs data for TP analysis:**

        1. **CUP Method Reference**: Independent import prices for similar goods can serve as CUP comparables
        2. **Anomaly Screening**: Related-party import prices significantly below independent prices may indicate profit shifting
        3. **Post-Import Verification**: Conduct TP review of completed customs declarations
        4. **Risk Alert**: Transactions with >15% price deviation should be prioritized for review

        ---

        ### 📋 依据文件 / References

        | 文件 / Document | 相关条款 / Relevant Articles |
        |---|---|
        | WTO海关估价协定 / WTO Customs Valuation Agreement | 第1-7条：估价方法 / Articles 1-7: Valuation methods |
        | 联合国转让定价实务手册 / UN Practical Manual | 第4章：海关与转让定价协同 / Ch.4: Customs-TP synergy |
        | OECD转让定价指南 / OECD TP Guidelines | 第2章：可比性分析 / Ch.2: Comparability analysis |
        | 联合国国际税收合作框架公约 / UN Framework Convention | 第5条、第6条 / Art. 5, Art. 6 |
        """)


# ============== Sidebar ==============

def render_sidebar() -> str:
    """Render the sidebar navigation and return the selected page name."""
    with st.sidebar:
        st.markdown(
            "<div class='sidebar-title'>🌍 国际转让定价对比数据库<br><span style='font-size:0.85rem;font-weight:500'>ITP Comparability Database</span></div>",
            unsafe_allow_html=True,
        )
        st.caption("联合国国际税收合作框架公约 / UN International Tax Cooperation Framework Convention")
        st.markdown("---")

        # Global Search
        st.markdown("#### 🔎 全局搜索 / Global Search")
        search_query = st.text_input(
            "搜索公司/经济体/行业 / Search",
            placeholder="输入关键词... / Enter keyword...",
            key="global_search_input",
            help="搜索公司名称、经济体名、ISIC行业代码 / Search companies, economies, ISIC codes",
            label_visibility="collapsed",
        )
        if search_query and len(search_query.strip()) >= 2:
            results = global_search(search_query)
            if results:
                st.markdown(f"**找到 {len(results)} 条结果 / {len(results)} results:**")
                for r_type, r_label, r_detail in results[:15]:
                    icon = {"company": "🏢", "country": "🌐", "industry": "🏭"}.get(r_type, "•")
                    st.markdown(f"<small>{icon} <b>{r_label}</b> — {r_detail}</small>", unsafe_allow_html=True)
                st.markdown("---")
            else:
                st.markdown("<small>无匹配结果 / No results found</small>")
                st.markdown("---")

        # Navigation — categorized with bilingual labels and informative tooltips
        st.markdown("#### 📑 功能导航 / Navigation")

        nav_categories = {
            "整体 / Overview": [
                ("Global Data Portal", "🌐 全球数据总览 / Global Portal", "世界地图+KPI+经济体详情 / World map, KPIs & economy detail"),
                ("Industry Benchmark", "📈 行业基准与趋势 / Industry Benchmark", "按行业/经济体/年度查看利润率基准 / View margin benchmarks by industry/economy/year"),
            ],
            "行业 / Industry": [
                ("Extractive Pricing", "⛏️ 采掘业定价 / Extractive Pricing", "CUP计算器+交易所参考价 / CUP calculator + exchange reference prices"),
                ("Digital Economy", "💻 数字经济分析 / Digital Economy", "数字服务定价+平台可比分析 / Digital services pricing & platform comparables"),
                ("Customs Valuation", "🛃 海关估价参考 / Customs Valuation", "海关数据辅助转让定价分析 / Customs data as TP auxiliary reference"),
            ],
            "案件 / Cases": [
                ("Comparability Wizard", "🔍 可比性分析向导 / Wizard", "5步引导完成完整TP分析 / 5-step guided TP analysis workflow"),
                ("MAP/APA Tracker", "⚖️ 争议解决追踪 / MAP-APA Tracker", "MAP/APA案件状态+趋势分析 / MAP/APA case status & trend analysis"),
                ("Risk Assessment", "🚨 风险评估中心 / Risk Center", "6维风险雷达+合规检查+CbC / 6-dim risk radar + compliance + CbC"),
            ],
            "政策 / Policy": [
                ("Country Policy", "📋 各经济体政策 / Country Policy", "TP法规+安全港+APA+MAP框架 / TP regs, safe harbor, APA & MAP framework"),
                ("Data Contribution", "📤 数据贡献管理 / Data Contribution", "上传+模板+验证+审批流程 / Upload, template, validation & approval"),
            ],
        }

        page_labels_short = {}
        for cat_items in nav_categories.values():
            for key, label, _tooltip in cat_items:
                page_labels_short[key] = label

        selected = None
        for cat_name, items in nav_categories.items():
            # Category header — styled as a subtle section label
            st.markdown(
                f"<div style='font-size:0.7rem;font-weight:700;color:#7F8C8D;"
                f"text-transform:uppercase;letter-spacing:1px;margin:0.8rem 0 0.3rem 0;'>{cat_name}</div>",
                unsafe_allow_html=True,
            )
            for page_key, page_label, page_tooltip in items:
                if st.button(page_label, key=f"nav_{page_key}", use_container_width=True,
                             help=page_tooltip):
                    selected = page_key
            st.markdown("<div style='height:0.2rem'></div>", unsafe_allow_html=True)

        # Fallback: if nothing selected yet, default to first page
        if selected is None:
            selected = st.session_state.get("_last_selected_page", "Global Data Portal")
        st.session_state["_last_selected_page"] = selected

        st.markdown("---")
        with st.expander("ℹ️ 关于本系统 / About"):
            st.markdown("""
            本系统支持**联合国国际税收合作框架公约**的实施，
            帮助发展中国家税务官员获取可比数据进行转让定价分析。

            This system supports the implementation of the
            **UN International Tax Cooperation Framework Convention**,
            helping developing economies access comparable data
            for transfer pricing analyses.

            **支持条款 / Articles Supported:**
            - 第5条 / Art. 5: 公平分配征税权 / Fair Distribution of Taxing Rights
            - 第11条 / Art. 11: 能力建设 / Capacity Building
            - 第13条 / Art. 13: 争议预防与解决 / Dispute Prevention & Resolution

            **方法论参考 / Methodology References:**
            - 联合国转让定价实务手册 / UN Practical Manual on TP
            - OECD转让定价指南 / OECD TPG
            (以UN体系为主，OECD体系为辅 / UN-primary, OECD-supplementary)
            """)
            if 'db_error' in st.session_state:
                st.warning(f"⚠️ 数据库连接失败，使用模拟数据。错误: {st.session_state['db_error'][:100]}")

        st.markdown(f"<small style='color:{TEXT_MUTED}'>版本 v4.0 | {datetime.now().year} | Powered by UN Framework Convention</small>", unsafe_allow_html=True)

        # Glossary
        with st.expander("📖 术语表 / Glossary"):
            st.markdown("""
            **转让定价核心术语 / Key TP Terms:**

            | 术语 | 说明 |
            |------|------|
            | **ALP** | 独立交易原则 Arm's Length Principle |
            | **FAR** | 功能分析 Functions-Assets-Risks |
            | **CUP** | 可比非受控价格法 Comparable Uncontrolled Price |
            | **RPM** | 再销售价格法 Resale Price Method |
            | **CPM** | 成本加成法 Cost Plus Method |
            | **TNMM** | 交易净利润法 Transactional Net Margin Method |
            | **PSM** | 利润分割法 Profit Split Method |
            | **PLI** | 利润水平指标 Profit Level Indicator |
            | **OM** | 营业利润率 Operating Margin |
            | **Berry Ratio** | 毛利/营业费用率 (>1.0有毛利) (OECD TPG Ch.6.28) |
            | **ROCE** | 已动用资本回报率 Return on Capital Employed |
            | **WCA** | 营运资金调整 Working Capital Adjustment |
            | **CbC** | 国别报告 Country-by-Country Report |
            | **MAP** | 相互协商程序 Mutual Agreement Procedure |
            | **APA** | 预约定价安排 Advance Pricing Agreement |
            | **IQR** | 四分位距 Interquartile Range (Q1-Q3) |
            """)

        return selected


# ============== Main ==============

def main() -> None:
    """Main entry point: render sidebar navigation and the selected page."""
    selected_page = render_sidebar()

    # Page router
    if selected_page == "Global Data Portal":
        global_data_portal()
    elif selected_page == "Industry Benchmark":
        industry_benchmark()
    elif selected_page == "Comparability Wizard":
        comparability_wizard()
    elif selected_page == "Extractive Pricing":
        extractive_pricing()
    elif selected_page == "Country Policy":
        country_policy()
    elif selected_page == "Data Contribution":
        data_contribution()
    elif selected_page == "MAP/APA Tracker":
        map_apa_tracker()
    elif selected_page == "Digital Economy":
        digital_economy()
    elif selected_page == "Risk Assessment":
        risk_assessment()
    elif selected_page == "Customs Valuation":
        customs_valuation()


if __name__ == "__main__":
    main()
